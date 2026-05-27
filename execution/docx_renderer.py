"""Word document renderer with a Chongqing University professional-master profile."""

from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from core.utils import setup_logging

logger = setup_logging("docx_renderer")

CQU_PROFESSIONAL_MASTER = "cqu_professional_master"


def render_docx(ast: dict, output_path: str = "") -> str:
    """Render one paper AST into a DOCX deliverable."""
    doc = Document()
    language = str(ast.get("language") or "中文")
    profile = str(ast.get("document_profile") or "")
    if profile == CQU_PROFESSIONAL_MASTER:
        _render_cqu_professional_master(doc, ast, language)
    else:
        _render_legacy(doc, ast, language)

    if not output_path:
        output_dir = ast.get("_output_dir", "data/outputs")
        os.makedirs(output_dir, exist_ok=True)
        title = ast.get("title", "未命名论文")
        safe_title = "".join(c for c in title if c.isalnum() or c in " _-").strip() or "paper"
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        job_suffix = re.sub(r"[^A-Za-z0-9_-]+", "", str(ast.get("_job_id", "")))[:24]
        suffix = f"_{job_suffix}" if job_suffix else ""
        output_path = os.path.join(output_dir, f"{safe_title}_{timestamp}{suffix}.docx")

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    finalize_fields = bool(ast.pop("_finalize_word_fields", False))
    if finalize_fields and profile == CQU_PROFESSIONAL_MASTER:
        warnings = _finalize_word_toc(output_path)
        if warnings:
            front = ast.setdefault("front_matter", {})
            front["layout_warnings"] = list(dict.fromkeys(
                list(front.get("layout_warnings") or []) + warnings
            ))
    logger.info("Word 文档已保存: %s", output_path)
    return output_path


def _render_cqu_professional_master(doc: Document, ast: dict, language: str) -> None:
    _setup_cqu_styles(doc)
    title = str(ast.get("title") or "未命名论文")
    front = dict(ast.get("front_matter") or {})
    registry = ast.get("entity_registry") or {}
    sections = list(ast.get("sections") or [])
    front, body_sections = _resolved_front_matter(front, sections, title)

    _setup_cqu_section(doc.sections[0], front_matter=True)
    _add_cqu_cover(doc, title, front, english=False)
    doc.add_page_break()
    _add_cqu_cover(doc, title, front, english=True)

    front_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _setup_cqu_section(front_section, front_matter=True, start_page=1, page_format="upperRoman")
    _add_abstract(doc, front.get("abstract_cn", "摘要待补充"), front.get("keywords_cn", "关键词待补充"), chinese=True)
    doc.add_page_break()
    _add_abstract(doc, front.get("abstract_en", "English abstract to be completed."), front.get("keywords_en", "Keywords to be completed"), chinese=False)
    doc.add_page_break()
    _add_toc(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _setup_cqu_section(body_section, front_matter=False, start_page=1, page_format="decimal")
    images = list(registry.get("images", [])) + list(registry.get("figures", []))
    tables = list(registry.get("tables", []))
    chapter_index = 0
    for section in body_sections:
        chapter_index += 1
        sec_title = _strip_heading_number(str(section.get("title") or f"第{chapter_index}章"))
        _add_heading(doc, f"{chapter_index} {sec_title}", 1)
        _render_content(
            doc,
            str(section.get("content") or ""),
            language=language,
            cqu=True,
            parent_title=sec_title,
        )
        _render_section_images(doc, images, section, chapter_no=chapter_index, language=language, cqu=True)
        _render_section_tables(doc, tables, section, chapter_no=chapter_index, language=language, cqu=True)

    _add_references(doc, ast.get("references", []), language=language, cqu=True)
    if ast.get("generation_notes"):
        _add_generation_notes(doc, ast["generation_notes"], language=language)
    _add_acknowledgements(doc)


def _render_legacy(doc: Document, ast: dict, language: str) -> None:
    _setup_legacy_styles(doc)
    _add_title_page(doc, ast.get("title", "未命名论文"))
    registry = ast.get("entity_registry") or {}
    images = list(registry.get("images", [])) + list(registry.get("figures", []))
    tables = list(registry.get("tables", []))
    for position, section in enumerate(ast.get("sections", []), start=1):
        doc.add_heading(section.get("title", ""), level=1)
        _render_content(doc, section.get("content", ""), language=language)
        _render_section_images(doc, images, section, chapter_no=position, language=language)
        _render_section_tables(doc, tables, section, chapter_no=position, language=language)
    _add_references(doc, ast.get("references", []), language=language)
    if ast.get("generation_notes"):
        _add_generation_notes(doc, ast["generation_notes"], language=language)


def _resolved_front_matter(front: dict, sections: list[dict], title: str) -> tuple[dict, list[dict]]:
    body = []
    for section in sections:
        section_title = str(section.get("title") or "").strip().lower()
        content = str(section.get("content") or "")
        if section_title in {"摘要", "abstract", "中文摘要"} and not front.get("abstract_cn"):
            front["abstract_cn"] = _extract_abstract_text(content)
            continue
        if section_title in {"英文摘要", "english abstract"} and not front.get("abstract_en"):
            front["abstract_en"] = _strip_markdown_heading(content)
            continue
        body.append(section)
    front.setdefault("title_en", "English Title To Be Completed")
    front.setdefault("abstract_cn", f"本文围绕“{title}”开展研究，摘要内容待根据正文进一步校核。")
    front.setdefault("abstract_en", "English abstract to be completed after the Chinese manuscript is finalized.")
    front.setdefault("keywords_cn", "待填写；待填写；待填写")
    front.setdefault("keywords_en", "To be completed; To be completed; To be completed")
    return front, body


def _add_cqu_cover(doc: Document, title: str, front: dict, *, english: bool) -> None:
    page_title = (
        front.get("title_en") or "English Title To Be Completed"
        if english else title
    )
    school_label = "A Thesis Submitted to Chongqing University" if english else "重庆大学硕士学位论文"
    degree_label = (
        "for the Professional Degree of\n[To Be Completed]"
        if english else "（专业学位）"
    )
    fields = (
        [
            ("By", "[To Be Completed]"),
            ("Supervised by", "[To Be Completed]"),
            ("Professional Degree Category", "[To Be Completed]"),
            ("Research Field", "[To Be Completed]"),
            ("Degree Date", "[To Be Completed]"),
        ]
        if english else
        [
            ("学生姓名", "待填写"),
            ("指导教师", "待填写"),
            ("专业学位类别", "待填写"),
            ("研究方向", "待填写"),
            ("授位时间", "待填写"),
        ]
    )
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font_run(p.add_run(page_title), size=22, bold=True, east_asia="黑体")
    p.paragraph_format.space_after = Pt(34)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font_run(p.add_run(school_label), size=18, bold=True, east_asia="黑体")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font_run(p.add_run(degree_label), size=15, east_asia="宋体")
    p.paragraph_format.space_after = Pt(70)
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font_run(p.add_run(f"{label}：{value}" if not english else f"{label}: {value}"), size=12, east_asia="宋体")
        p.paragraph_format.space_after = Pt(10)


def _add_abstract(doc: Document, text: str, keywords: str, *, chinese: bool) -> None:
    _add_heading(doc, "摘要" if chinese else "Abstract", 1)
    for paragraph in re.split(r"\n\s*\n", _strip_markdown_heading(text)):
        if paragraph.strip():
            _add_body_paragraph(doc, paragraph.strip(), first_indent=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    prefix = "关键词：" if chinese else "Keywords: "
    _font_run(p.add_run(prefix), size=12, bold=True, east_asia="黑体" if chinese else "宋体")
    _font_run(p.add_run(str(keywords)), size=12, east_asia="宋体")


def _add_toc(doc: Document) -> None:
    # Keep the visible title out of Heading 1 so it cannot index itself.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(8)
    _font_run(title.add_run("目录"), size=16, bold=True, east_asia="黑体")

    field_paragraph = doc.add_paragraph()
    field_paragraph.paragraph_format.first_line_indent = None
    field_paragraph.paragraph_format.line_spacing = Pt(20)
    _append_field(
        field_paragraph,
        r'TOC \o "1-3" \h \z \u',
        placeholder="目录待生成。",
        dirty=True,
    )


def _setup_cqu_section(section, *, front_matter: bool, start_page: int | None = None, page_format: str = "") -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.6)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not front_matter:
        _font_run(header.add_run("重庆大学硕士学位论文"), size=9, east_asia="宋体")
    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if start_page is not None:
        _append_field(footer, "PAGE")
    sect_pr = section._sectPr
    if start_page is not None:
        pg_num = sect_pr.find(qn("w:pgNumType"))
        if pg_num is None:
            pg_num = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num)
        pg_num.set(qn("w:start"), str(start_page))
        if page_format:
            pg_num.set(qn("w:fmt"), page_format)


def _append_field(paragraph, instruction: str, *, placeholder: str = "1", dirty: bool = False) -> None:
    run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    if dirty:
        begin.set(qn("w:dirty"), "true")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, code, sep, text, end):
        run.append(node)
    paragraph._p.append(run)


def _finalize_word_toc(output_path: str) -> list[str]:
    """Populate TOC result text before delivery without forcing update-on-open."""
    word = None
    document = None
    pythoncom = None
    try:
        import pythoncom
        import win32com.client

        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(Path(output_path).resolve()),
            ConfirmConversions=False,
            ReadOnly=False,
            AddToRecentFiles=False,
        )
        for toc in document.TablesOfContents:
            toc.Update()
        document.Fields.Update()
        document.Save()
        return []
    except Exception as exc:
        logger.warning("Word 目录预生成失败，将保留可手动更新目录域: %s", exc)
        return [f"目录自动预生成失败，请在 Word 中手动更新目录: {exc}"]
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if pythoncom is not None:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass


def _add_title_page(doc: Document, title: str) -> None:
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font_run(title_para.add_run(title), size=22, bold=True)
    doc.add_paragraph()


def _render_content(
    doc: Document,
    content: str,
    language: str = "中文",
    cqu: bool = False,
    parent_title: str = "",
) -> None:
    content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
    lines = content.split("\n")
    buffer: list[str] = []
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        match = re.match(r"^(#{2,4})[ \u3000]+(.+)$", stripped)
        if match:
            heading_text = match.group(2).strip()
            if (
                cqu
                and _strip_heading_number(heading_text) == _strip_heading_number(parent_title)
            ):
                index += 1
                continue
            _flush_text(doc, buffer, cqu=cqu)
            level = min(len(match.group(1)), 3) if cqu else len(match.group(1)) - 1
            _add_heading(doc, heading_text, level) if cqu else doc.add_heading(heading_text, level=level)
            buffer = []
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and _is_markdown_separator(lines[index + 1]):
            _flush_text(doc, buffer, cqu=cqu)
            table_lines = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            _insert_markdown_table(doc, table_lines)
            buffer = []
            continue
        if not stripped and buffer:
            if buffer[-1] != "":
                buffer.append("")
            index += 1
            continue
        fig_match = re.match(r"\[Figure:\s*(.+?)\]", stripped)
        if fig_match:
            _flush_text(doc, buffer, cqu=cqu)
            _insert_figure(doc, fig_match.group(1).strip(), language=language)
            buffer = []
            index += 1
            continue
        buffer.append(line)
        index += 1
    _flush_text(doc, buffer, cqu=cqu)


def _flush_text(doc: Document, buffer: list[str], cqu: bool = False) -> None:
    text = "\n".join(buffer).strip()
    if not text:
        return
    for para_text in re.split(r"\n+", text):
        para_text = re.sub(r"\*\*", "", para_text.strip())
        if para_text:
            _add_body_paragraph(doc, para_text, first_indent=True if cqu else True)


def _is_markdown_separator(line: str) -> bool:
    cells = _markdown_cells(line)
    return len(cells) >= 2 and all(
        bool(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")))
        for cell in cells
    )


def _markdown_cells(line: str) -> list[str]:
    return [cell.strip() for cell in str(line).strip().strip("|").split("|")]


def _insert_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = [_markdown_cells(line) for line in lines if not _is_markdown_separator(line)]
    if len(rows) < 2 or not rows[0]:
        for line in lines:
            _add_body_paragraph(doc, line, first_indent=False)
        return
    column_count = len(rows[0])
    table = doc.add_table(rows=1, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    width = Cm(15.0 / column_count)
    for cell, value in zip(table.rows[0].cells, rows[0]):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_text(cell, value, bold=True)
    for row in rows[1:]:
        values = (row + [""] * column_count)[:column_count]
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_text(cell, value)


def _add_body_paragraph(doc: Document, text: str, *, first_indent: bool) -> None:
    para = doc.add_paragraph()
    _font_run(para.add_run(text), size=12, east_asia="宋体")
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Cm(0.74) if first_indent else None
    para.paragraph_format.line_spacing = Pt(20)
    para.paragraph_format.space_after = Pt(0)


def _insert_figure(doc: Document, fig_path: str, language: str = "中文") -> None:
    if not os.path.exists(fig_path):
        logger.warning("图片不存在，跳过: %s", fig_path)
        return
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_visual_paragraph_format(para)
        para.add_run().add_picture(fig_path, width=Cm(14))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_visual_paragraph_format(caption)
        label = "图" if _is_chinese(language) else "Figure"
        _font_run(caption.add_run(f"{label}: {os.path.basename(fig_path)}"), size=10, east_asia="宋体")
    except Exception as exc:
        logger.warning("图片插入失败: %s", exc)


def _render_section_images(doc, images: list[dict], section: dict, chapter_no: int, language: str = "中文", cqu: bool = False) -> None:
    aliases = _section_aliases(section, chapter_no, include_position=not cqu)
    count = 0
    for image in images:
        if not image.get("approved"):
            continue
        image_section = str(image.get("section_id", "")).strip().lower()
        if image_section and image_section not in aliases:
            continue
        local_path = image.get("local_path", "")
        if not local_path or not os.path.exists(local_path):
            continue
        count += 1
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_visual_paragraph_format(para)
        para.paragraph_format.keep_with_next = True
        para.add_run().add_picture(local_path, width=Cm(14))
        caption_text = image.get("caption") or image.get("title") or os.path.basename(local_path)
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_visual_paragraph_format(caption)
        number = f"{chapter_no}.{count}"
        primary = f"图{number} {caption_text}" if cqu else f"{'图' if _is_chinese(language) else 'Figure'}: {caption_text}"
        _font_run(caption.add_run(primary), size=10.5 if cqu else 10, east_asia="宋体")
        english_caption = image.get("english_caption")
        if cqu and english_caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_visual_paragraph_format(p)
            _font_run(p.add_run(f"Fig.{number} {english_caption}"), size=10.5)
        source_bits = [image.get("attribution", ""), image.get("license", ""), image.get("source_url", "")]
        source = " | ".join(str(bit) for bit in source_bits if bit)
        if source:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_visual_paragraph_format(p)
            run = p.add_run(f"{'来源' if _is_chinese(language) else 'Source'}: {source}")
            _font_run(run, size=9, east_asia="宋体")
            run.font.color.rgb = RGBColor(96, 96, 96)


def _render_section_tables(doc, tables: list[dict], section: dict, chapter_no: int, language: str = "中文", cqu: bool = False) -> None:
    aliases = _section_aliases(section, chapter_no, include_position=not cqu)
    count = 0
    for table_spec in tables:
        table_section = str(table_spec.get("section_id", "")).strip().lower()
        if table_section and table_section not in aliases:
            continue
        headers = list(table_spec.get("headers") or [])
        rows = list(table_spec.get("rows") or [])
        if not headers or not rows:
            continue
        count += 1
        number = f"{chapter_no}.{count}"
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_visual_paragraph_format(caption)
        text = table_spec.get("caption", "")
        primary = f"表{number} {text}" if cqu else f"{'表' if _is_chinese(language) else 'Table'}: {text}"
        _font_run(caption.add_run(primary), size=10.5 if cqu else 10, bold=True, east_asia="宋体")
        if cqu and table_spec.get("english_caption"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_visual_paragraph_format(p)
            _font_run(p.add_run(f"Table {number} {table_spec['english_caption']}"), size=10.5)
        word_table = doc.add_table(rows=1, cols=len(headers))
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        word_table.style = "Table Grid"
        word_table.autofit = False
        width = Cm(15.0 / len(headers))
        for cell, value in zip(word_table.rows[0].cells, headers):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_text(cell, value, bold=True)
        for row in rows:
            cells = word_table.add_row().cells
            for cell, value in zip(cells, row):
                cell.width = width
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _set_cell_text(cell, value)
        source = table_spec.get("source", "")
        if source:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_visual_paragraph_format(p)
            run = p.add_run(f"{'来源' if _is_chinese(language) else 'Source'}: {source}")
            _font_run(run, size=9, east_asia="宋体")
            run.font.color.rgb = RGBColor(96, 96, 96)


def _set_cell_text(cell, value, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font_run(p.add_run(str(value)), size=10, bold=bold, east_asia="宋体")


def _set_visual_paragraph_format(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(4)


def _section_aliases(section: dict, position: int, *, include_position: bool = True) -> set[str]:
    section_id = section.get("section_id") or section.get("id") or section.get("title", "")
    aliases = {str(section_id).strip().lower(), str(section.get("title", "")).strip().lower()}
    if include_position or (not section.get("section_id") and not section.get("id")):
        aliases.add(f"sec{position}")
    return aliases


def _add_references(doc: Document, references: list[dict], language: str = "中文", cqu: bool = False) -> None:
    if cqu:
        _add_heading(doc, "参考文献", 1)
    else:
        doc.add_heading("参考文献" if _is_chinese(language) else "References", level=1)
    seen_titles = set()
    unique_references = []
    for ref in references:
        title = str(ref.get("title", "")).strip()
        if not title or title.lower() in seen_titles:
            continue
        seen_titles.add(title.lower())
        unique_references.append(ref)
    for index, ref in enumerate(unique_references, 1):
        title = str(ref.get("title", "")).strip()
        authors = ", ".join(ref.get("authors", [])[:3])
        if len(ref.get("authors", [])) > 3:
            authors += ", et al"
        year = ref.get("year", "")
        venue = str(ref.get("venue", "") or "").strip()
        citation = f"[{index}] {authors}. {title}[J]."
        if venue:
            citation += f" {venue},"
        if year:
            citation += f" {year}."
        if ref.get("url"):
            citation += f" {ref['url']}."
        p = doc.add_paragraph()
        _font_run(p.add_run(citation), size=10.5, east_asia="宋体")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.line_spacing = Pt(18)


def _add_generation_notes(doc: Document, notes: dict, language: str = "中文") -> None:
    quality = list(notes.get("quality_warnings") or [])
    evidence = list(notes.get("evidence_warnings") or [])
    assets = list(notes.get("asset_warnings") or [])
    layout = list(notes.get("layout_warnings") or [])
    if notes.get("generation_mode") == "standard" and not (quality or evidence or assets or layout):
        return
    chinese = _is_chinese(language)
    _add_heading(doc, "生成说明与待核验项" if chinese else "Generation Notes and Items Requiring Verification", 1)
    _add_body_paragraph(
        doc,
        "本文档基于已提供材料、检索文献与自动写作流程生成。标记为待核验的内容应在正式提交前复核。"
        if chinese else
        "This document was generated from supplied materials, retrieved literature and an automated writing workflow. Marked items require review before formal use.",
        first_indent=False,
    )
    for label, items in (
        ("证据" if chinese else "Evidence", evidence),
        ("质量" if chinese else "Quality", quality),
        ("图像与图表" if chinese else "Images and Figures", assets),
        ("版式" if chinese else "Layout", layout),
    ):
        if not items:
            continue
        _add_heading(doc, label, 2)
        for item in items[:12]:
            doc.add_paragraph(str(item), style="List Bullet")


def _add_acknowledgements(doc: Document) -> None:
    _add_heading(doc, "致谢", 1)
    _add_body_paragraph(doc, "致谢内容待填写。", first_indent=True)


def _setup_cqu_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _style_font(normal, size=12, east_asia="宋体")
    normal.paragraph_format.line_spacing = Pt(20)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_after = Pt(0)
    for name, size, east_asia, alignment in (
        ("Heading 1", 16, "黑体", WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 15, "黑体", WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 12, "黑体", WD_ALIGN_PARAGRAPH.LEFT),
    ):
        style = doc.styles[name]
        _style_font(style, size=size, east_asia=east_asia, bold=True)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(8)


def _setup_legacy_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    _style_font(style, size=12, east_asia="宋体")


def _add_heading(doc: Document, text: str, level: int):
    return doc.add_paragraph(text, style=f"Heading {max(1, min(level, 3))}")


def _style_font(style, *, size: float, east_asia: str, bold: bool = False) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    fonts = rpr.rFonts
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), east_asia)


def _font_run(run, *, size: float, bold: bool = False, east_asia: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), east_asia)


def _strip_heading_number(text: str) -> str:
    return re.sub(r"^\s*(?:第?[一二三四五六七八九十\d]+章?[.、\s]*)", "", text).strip() or text


def _strip_markdown_heading(text: str) -> str:
    return re.sub(r"(?m)^\s*#{1,6}\s*", "", str(text or "")).strip()


def _extract_abstract_text(text: str) -> str:
    abstract = _strip_markdown_heading(text)
    abstract = re.sub(r"^\s*(?:摘要|中文摘要)\s*(?:\r?\n)+", "", abstract, count=1)
    return re.split(r"(?m)^\s*参考文献\s*$", abstract, maxsplit=1)[0].strip()


def _is_chinese(language: str) -> bool:
    return str(language or "").strip().lower() not in {"english", "en", "英文"}
