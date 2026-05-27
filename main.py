#!/usr/bin/env python3
"""
智能论文写作 Agent v6.0 — 聚合 MCP Server（主从解耦版）

架构: 本地控制面 + 本地执行面 + Nexent 推理面
Nexent 仅作为 3 个原子 API (Generator_Coder, Generator_Writer, Discriminator) 的 LLM 推理端点。
所有编排逻辑、状态管理、代码执行、渲染输出全部在本地 Python 进程中完成。

对外暴露 8 个 Tool:
  1. literature_search   — 动态检索真实文献
  2. generate_outline    — 生成论文大纲（含 nature 结构校验）
  3. section_writer      — 基于知识库写单章（含 Discriminator 审查）
  4. generate_figure     — 图表代码生成 + 沙盒执行 + Debug 循环
  5. write_paper         — 全自动流水线（逐章写作 + 引用解析 + 渲染）
  6. render_paper        — 输出渲染入口
  7. edit_document       — Word 后处理
  8. table_operation     — Word 表格操作
"""

import asyncio, os, sys, json, re, shutil, tempfile, socket, time
from pathlib import Path
from urllib.parse import urlparse, unquote

import requests
import yaml
from fastmcp import FastMCP, Context

_root = Path(__file__).parent
if str(_root) not in sys.path:
    sys.path.insert(0, str(_root))

from core.literature_search import search_paper_pool
from core.section_writer import write_section as write_section_local
from core.llm_client import get_llm_client
from core.utils import (
    read_document, save_json, load_json,
    get_previous_context, setup_logging,
)
from orchestrator.state_manager import StateManager, PaperProject
from orchestrator.workflow_engine import WorkflowEngine, WorkflowState
from orchestrator.context_builder import build_section_context
from orchestrator.asset_manager import (
    ImageAssetStore,
    search_commons_candidates,
    search_planned_commons_candidates,
)
from orchestrator.material_registry import MaterialRegistry
from orchestrator.material_sufficiency import (
    assess_material_sufficiency,
    build_asset_plan,
)
from orchestrator.write_jobs import WriteJobProgressContext, WriteJobStore
from orchestrator.pending_outlines import PendingOutlineStore, request_fingerprint
from orchestrator.quality_audit import audit_manuscript
from orchestrator.evidence_registry import build_evidence_registry
from inference.nexent_client import NexentClient
from inference.discriminator import review_section
from inference.generator_coder import generate_code
from inference.generator_writer import write_section as write_section_nexent
from execution.code_sandbox import CodeSandbox
from execution.figure_generator import generate_figure
from execution.concept_diagram import render_research_framework_diagram, render_system_workflow_diagram
from execution.simulated_results import build_simulated_result_assets
from execution.citation_resolver import CitationResolver
from execution.docx_renderer import CQU_PROFESSIONAL_MASTER, render_docx
from execution.render_validator import validate_renderable_ast
from patterns.move_sequences import format_moves_prompt
from patterns.article_architecture import validate_outline

logger = setup_logging("orchestrator")
mcp = FastMCP("paper_agent_v6")

_config_path = Path(__file__).parent / "config.yaml"
_APP_CONFIG = yaml.safe_load(_config_path.read_text("utf-8")) if _config_path.exists() else {}
_FILE_GATEWAY = _APP_CONFIG.get("nexent", {}).get("file_gateway_url", "") or ""
_MINIO_PUBLIC_URL = _APP_CONFIG.get("nexent", {}).get("minio_public_url", "") or ""

# ============================================================
# 全局单例
# ============================================================
_state_manager = StateManager(
    checkpoint_dir=_APP_CONFIG.get("paper", {}).get("checkpoint_dir", "data/checkpoints")
)
_nexent_client = NexentClient()
_sandbox = CodeSandbox(
    timeout=_APP_CONFIG.get("figure", {}).get("sandbox_timeout", 30)
)
_citation_resolver = CitationResolver()
_image_assets = ImageAssetStore("data/checkpoints/image_assets.json")
_materials = MaterialRegistry("data/checkpoints/materials.json")
_write_jobs = WriteJobStore("data/checkpoints/write_jobs.json")
_pending_outlines = PendingOutlineStore("data/checkpoints/pending_outlines.json")
_write_job_tasks: dict[str, asyncio.Task] = {}
_workflow_engine = None

def _get_advertised_host() -> str:
    """Return the host embedded in browser-facing document download links."""
    custom = os.getenv("PAPER_AGENT_ADVERTISED_HOST") or _APP_CONFIG.get("server", {}).get("advertised_host", "")
    return custom if custom else "localhost"


def _get_server_port() -> int:
    """Return the externally published MCP/download port."""
    return int(os.getenv("PAPER_AGENT_PORT") or _APP_CONFIG.get("server", {}).get("port", 8001))


def _get_listen_host() -> str:
    """Return the interface on which this process accepts connections."""
    return os.getenv("PAPER_AGENT_HOST") or _APP_CONFIG.get("server", {}).get("host", "0.0.0.0")


# 内存动态知识库（兼容旧版简化架构）
DYNAMIC_KNOWLEDGE_BASE: dict = {
    "current_references": [],
    "written_chapters": {},
    "current_outline": None,
    "current_ast": None,
    "source_material": "",
    "result_evidence": "",
    "asset_plan": [],
    "material_sufficiency": None,
    "active_material_ids": [],
    "active_material_request_fingerprint": "",
    "material_policy": {
        "allow_degraded_writing": False,
        "user_acknowledgement": "",
    },
}

# ============================================================
# 内部工具函数
# ============================================================

def _build_citation_registry(all_papers):
    """构建全局引用注册表（兼容旧版）"""
    registry, mappings = {}, []
    for sp in all_papers:
        m = {}
        for p in sp:
            t = (p.get("title") or "").strip().lower()
            if not t:
                continue
            if t not in registry:
                registry[t] = {**p, "_global_idx": len(registry)}
            gk = f"ref{registry[t]['_global_idx']}"
            m[p.get("citation_key", "")] = gk
        mappings.append(m)
    return registry, mappings

def _replace_citations(content, mapping):
    for lk, gk in mapping.items():
        content = content.replace(f"[{lk}]", f"[{gk}]")
    return content

def _extract_filename(path):
    try:
        parsed = urlparse(path)
        hint = os.path.basename(parsed.path) if parsed.path else ""
        if not hint or hint in ("fetch", ""):
            qs = parsed.query if parsed.query else path
            m = re.search(r"presigned_url=([^&\s]+)", qs)
            if m:
                inner = unquote(m.group(1))
                hint = os.path.basename(inner.split("?")[0]) or ""
        for marker in ("=X-Amz-", "?X-Amz-", "&X-Amz-"):
            if marker in hint:
                hint = hint.split(marker)[0]
        return hint or "uploaded_file"
    except Exception:
        return os.path.basename(path) or "uploaded_file"

def _download_from_nexent(s3_path):
    logger.info(f"检测到 Nexent 云端路径，尝试下载: {s3_path[:120]}...")
    urls_to_try = _build_nexent_download_urls(s3_path)
    if not urls_to_try:
        logger.warning("  无法提取下载地址")
        return None

    logger.info(f"  下载地址候选: {urls_to_try[0][:100]}...")

    # 绕过系统代理（Clash/V2Ray），直连 Docker 暴露端口
    sess = requests.Session()
    sess.trust_env = False

    for url in urls_to_try:
        try:
            resp = sess.get(url, timeout=30, allow_redirects=True)
            resp.raise_for_status()
            raw_url = str(resp.url)
            clean = raw_url.split("?")[0]
            suffix = Path(clean).suffix or ""
            if not suffix or len(suffix) > 10 or any(c in suffix for c in "%?&"):
                ct = resp.headers.get("content-type", "")
                if "pdf" in ct: suffix = ".pdf"
                elif "text/plain" in ct: suffix = ".txt"
                elif "openxml" in ct or "word" in ct: suffix = ".docx"
                else: suffix = ".docx"
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tmp.write(resp.content); tmp.close()
            logger.info(f"  下载成功: {tmp.name} ({len(resp.content)} bytes)")
            return tmp.name
        except Exception as e:
            logger.warning(f"  URL {url[:80]}... 下载失败: {e}")
            continue
    logger.error("  所有下载方式均失败")
    return None


def _build_nexent_download_urls(raw_path: str) -> list[str]:
    urls = []
    if raw_path.startswith("http://") or raw_path.startswith("https://"):
        urls.append(raw_path)
    elif raw_path.startswith("s3://"):
        parsed = urlparse(raw_path)
        bucket = parsed.netloc
        object_path = parsed.path.lstrip("/")
        if _MINIO_PUBLIC_URL and bucket and object_path:
            urls.append(f"{_MINIO_PUBLIC_URL.rstrip('/')}/{bucket}/{object_path}")
        if _FILE_GATEWAY and object_path:
            urls.append(_FILE_GATEWAY.rstrip("/") + "/" + object_path)
        # Keep legacy fallbacks last for deployments exposing bucket DNS names.
        urls.append(raw_path.replace("s3://", "https://", 1))
    elif "presigned_url=" in raw_path:
        match = re.search(r"presigned_url=([^&\s]+)", raw_path)
        if match:
            urls.append(unquote(match.group(1)))

    expanded = []
    for url in urls:
        if url not in expanded:
            expanded.append(url)
        if url.startswith("https://"):
            http_url = url.replace("https://", "http://", 1)
            if http_url not in expanded:
                expanded.append(http_url)
        if "nexent" in url and "5013" not in url and "9010" not in url:
            parsed = urlparse(url)
            query = f"?{parsed.query}" if parsed.query else ""
            proxy_url = f"http://localhost:5013{parsed.path}{query}"
            if proxy_url not in expanded:
                expanded.append(proxy_url)
    return expanded


def _active_materials() -> list[dict]:
    materials = _materials.all()
    active_material_ids = DYNAMIC_KNOWLEDGE_BASE.get("active_material_ids")
    if active_material_ids is None:
        return materials
    active_ids = set(active_material_ids)
    return [material for material in materials if material.get("material_id") in active_ids]


def _set_active_materials(materials: list[dict], *, replace: bool = False) -> None:
    current = [] if replace else list(DYNAMIC_KNOWLEDGE_BASE.get("active_material_ids") or [])
    ids = current + [material.get("material_id", "") for material in materials if material.get("material_id")]
    DYNAMIC_KNOWLEDGE_BASE["active_material_ids"] = list(dict.fromkeys(ids))


def _material_text() -> str:
    text = "\n\n".join(
        material.get("extracted_text", "")
        for material in _active_materials()
        if material.get("parse_status") == "parsed"
        and material.get("material_role") in {"source_text", "reference_material"}
        and material.get("extracted_text")
    )
    return text


def _result_evidence() -> str:
    registry_summary = "\n\n".join(
        material.get("data_summary", "")
        for material in _active_materials()
        if material.get("parse_status") == "parsed"
        and material.get("material_role") == "result_dataset"
        and material.get("data_summary")
    )
    explicit = DYNAMIC_KNOWLEDGE_BASE.get("result_evidence", "")
    return "\n\n".join(part for part in (explicit, registry_summary) if part)


def _material_sufficiency(topic: str = "", material_text: str = "", result_evidence: str = "") -> dict:
    policy = _material_policy()
    return assess_material_sufficiency(
        topic=topic,
        material_text=material_text or _material_text(),
        result_evidence=result_evidence or _result_evidence(),
        references=DYNAMIC_KNOWLEDGE_BASE.get("current_references", []),
        # Visual assets are acquired/generated by this service during writing;
        # old approved assets must never make a new paper appear sufficient.
        approved_assets=[],
        materials=_active_materials(),
        allow_degraded_writing=policy["allow_degraded_writing"],
        degraded_reason=policy["user_acknowledgement"],
    )


def _material_policy() -> dict:
    policy = DYNAMIC_KNOWLEDGE_BASE.get("material_policy") or {}
    return {
        "allow_degraded_writing": bool(policy.get("allow_degraded_writing")),
        "user_acknowledgement": str(policy.get("user_acknowledgement", "")).strip(),
    }


def _strict_material_policy() -> dict:
    return {
        "allow_degraded_writing": False,
        "user_acknowledgement": "",
    }


def _prepare_outline_request_scope(source_request: dict) -> str:
    """Discard state bound to another request while retaining freshly staged uploads."""
    fingerprint = request_fingerprint(source_request)
    binding = str(DYNAMIC_KNOWLEDGE_BASE.get("active_material_request_fingerprint") or "")
    if binding not in {"staged", fingerprint}:
        DYNAMIC_KNOWLEDGE_BASE["active_material_ids"] = []
        DYNAMIC_KNOWLEDGE_BASE["source_material"] = ""
        DYNAMIC_KNOWLEDGE_BASE["result_evidence"] = ""
    if binding != fingerprint:
        DYNAMIC_KNOWLEDGE_BASE["material_policy"] = _strict_material_policy()
    return fingerprint


def _resolve_language(language: str = "") -> str:
    requested = str(language or "").strip().lower()
    if requested in {"中文", "汉语", "chinese", "zh", "zh-cn", "简体中文"}:
        return "中文"
    if requested in {"english", "en", "英文"}:
        return "English"
    configured = str(_APP_CONFIG.get("output", {}).get("default_language", "中文")).strip()
    return "English" if configured.lower() in {"english", "en", "英文"} else "中文"


def _language_instruction(language: str) -> str:
    if _resolve_language(language) == "中文":
        return (
            "\n\n## 语言约束（必须遵守）\n"
            "整篇论文必须使用中文撰写，包括章节标题、二三级标题、正文、图表说明与结论。"
            "仅专业缩写、算法名和参考文献原题允许保留英文。"
        )
    return (
        "\n\n## Language Requirement (mandatory)\n"
        "Write the entire manuscript in English, including headings, body text, "
        "figure/table captions and conclusions. Keep non-English source titles only when citing them."
    )


def _localize_standard_outline(outline: dict, language: str) -> dict:
    """Keep the standard paper skeleton aligned with the frozen task language."""
    target = _resolve_language(language)
    title_map = {
        "abstract": ("摘要", "Abstract"),
        "introduction": ("引言", "Introduction"),
        "methods": ("方法", "Methods"),
        "method": ("方法", "Methods"),
        "results": ("结果", "Results"),
        "result": ("结果", "Results"),
        "discussion": ("讨论", "Discussion"),
        "conclusion": ("结论", "Conclusion"),
    }
    for section in outline.get("sections", []):
        identifier = str(section.get("id") or "").strip().lower()
        title = str(section.get("title") or "").strip().lower()
        for key, labels in title_map.items():
            if identifier == key or title == key:
                section["title"] = labels[0] if target == "中文" else labels[1]
                break
    outline["language"] = target
    return outline


def _material_report_allows_writing(report: dict) -> bool:
    if report.get("status") in {"sufficient", "DEGRADED_WRITING_ALLOWED"}:
        return True
    system_resolvable = {"references", "figure_or_image_plan", "unbound_uploaded_image_asset"}
    missing = set(report.get("missing_materials") or [])
    return bool(missing) and not report.get("blocking_missing_materials") and missing <= system_resolvable


def _material_report_blocks_output(report: dict) -> bool:
    """Output-first writing blocks only when there is no usable paper subject."""
    return bool(report.get("blocking_missing_materials"))


def _warnings_from_material_report(report: dict) -> list[str]:
    missing = list(report.get("missing_materials") or [])
    if not missing:
        return []
    return [f"材料缺口: {item}" for item in missing]


def _audit_warning_messages(audit: dict) -> list[str]:
    return [
        f"{issue.get('code', 'quality_issue')}: {issue.get('message', '')}".strip()
        for issue in audit.get("issues", [])
    ]


def _warning_groups(
    *,
    material_report: dict,
    audit: dict,
    validation: dict,
    citation_warnings: list[str] | None = None,
    asset_warnings: list[str] | None = None,
) -> dict:
    evidence = _warnings_from_material_report(material_report)
    evidence.extend(citation_warnings or [])
    asset_validation = []
    for warning in validation.get("warnings") or []:
        lowered = str(warning).lower()
        if "reference" in lowered or "citation" in lowered or "source material" in lowered:
            evidence.append(str(warning))
        else:
            asset_validation.append(str(warning))
    return {
        "quality_warnings": _dedupe_warnings(_audit_warning_messages(audit)),
        "evidence_warnings": _dedupe_warnings(evidence),
        "asset_warnings": _dedupe_warnings(
            asset_validation + list(asset_warnings or [])
        ),
    }


def _generation_notes(warnings: dict, generation_mode: str) -> dict:
    if not _APP_CONFIG.get("output", {}).get("include_generation_notes", True):
        return {}
    return {**warnings, "generation_mode": generation_mode}


def _document_profile() -> str:
    return CQU_PROFESSIONAL_MASTER


def _front_matter(title: str) -> dict:
    return {
        "title_en": "English Title To Be Completed",
        "abstract_en": "English abstract to be completed after the Chinese manuscript is finalized.",
        "keywords_cn": "待填写；待填写；待填写",
        "keywords_en": "To be completed; To be completed; To be completed",
        "cover_fields_mode": "placeholder",
        "layout_warnings": ["英文题目、英文摘要及封面身份信息为待填写占位内容，正式提交前需补全。"],
    }


def _complete_english_front_matter(ast: dict) -> None:
    """Best-effort English front matter generation from the completed Chinese abstract."""
    front = dict(ast.get("front_matter") or {})
    placeholder = "English abstract to be completed after the Chinese manuscript is finalized."
    if str(front.get("abstract_en") or "").strip() not in {"", placeholder}:
        return
    abstract_cn = ""
    for section in ast.get("sections") or []:
        if str(section.get("title") or "").strip().lower() in {"摘要", "中文摘要"}:
            abstract_cn = str(section.get("content") or "").strip()
            break
    if not abstract_cn:
        return
    abstract_cn = re.split(r"(?m)^\s*参考文献\s*$", abstract_cn, maxsplit=1)[0].strip()
    base_prompt = (
        "Translate and condense the following Chinese master's thesis abstract into formal academic English. "
        "Return JSON with keys title_en, abstract_en, keywords_en only. "
        "Do not introduce unverified results and preserve any simulated-data limitation.\n\n"
        f"Chinese title: {ast.get('title', '')}\nChinese abstract:\n{abstract_cn[:5000]}"
    )
    errors = []
    for attempt in range(2):
        try:
            prompt = base_prompt
            if attempt:
                prompt = (
                    "The prior response could not be parsed as JSON. Return a shorter valid JSON object only, "
                    "with no markdown or explanation.\n\n" + base_prompt
                )
            response = get_llm_client().call(
                prompt,
                system="You prepare faithful English front matter for a university thesis.",
                response_format="json",
            )
            payload = json.loads(re.sub(r"^```(?:json)?\s*|\s*```$", "", response.strip()))
            if len(str(payload.get("abstract_en") or "").strip()) < 80:
                raise ValueError("generated English abstract is too short")
            keywords_en = payload.get("keywords_en") or front.get("keywords_en") or ""
            if isinstance(keywords_en, list):
                keywords_en = "; ".join(str(keyword).strip() for keyword in keywords_en if str(keyword).strip())
            front["title_en"] = str(payload.get("title_en") or front.get("title_en") or "").strip()
            front["abstract_en"] = str(payload["abstract_en"]).strip()
            front["keywords_en"] = str(keywords_en).strip()
            front["layout_warnings"] = ["封面身份信息为待填写占位内容，正式提交前需补全。"]
            ast["front_matter"] = front
            return
        except Exception as exc:
            errors.append(str(exc))
    if errors:
        warnings = list(front.get("layout_warnings") or [])
        warnings.append(f"英文摘要自动生成失败，当前保留待补充标记: {' | '.join(errors)}")
        front["layout_warnings"] = _dedupe_warnings(warnings)
    ast["front_matter"] = front


def _task_asset_root(job_id: str) -> str:
    return str((Path("data/assets") / (job_id or "paper")).resolve())


def _renderable_asset_path(asset: dict) -> str:
    path = str(asset.get("local_path") or asset.get("path") or "").strip()
    return path if path and Path(path).is_file() else ""


def _asset_delivery_metrics(ast: dict) -> dict:
    registry = ast.get("entity_registry") or {}
    images = [asset for asset in (registry.get("images") or []) if _renderable_asset_path(asset)]
    figures = [asset for asset in (registry.get("figures") or []) if _renderable_asset_path(asset)]
    tables = list(registry.get("tables") or [])
    all_assets = images + figures
    return {
        "image_count": len(images),
        "figure_count": len(figures),
        "table_count": len(tables),
        "network_image_count": len([asset for asset in images if asset.get("source") in {"wikimedia_commons", "openverse"}]),
        "uploaded_image_count": len([asset for asset in images if asset.get("source") == "user_upload"]),
        "simulated_asset_count": len([asset for asset in all_assets + tables if asset.get("simulated")]),
    }


def _dedupe_warnings(items: list[str]) -> list[str]:
    seen = set()
    deduped = []
    for item in items:
        text = str(item or "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        deduped.append(text)
    return deduped


def _job_snapshot(
    confirmed_outline,
    source_request: dict | None = None,
    outline_id: str = "",
) -> dict:
    source_request = dict(source_request or {})
    title, section_count = _outline_job_summary(confirmed_outline)
    if source_request.get("topic"):
        title = str(source_request["topic"])
    return {
        "title": title,
        "section_count": section_count,
        "outline": confirmed_outline,
        "source_request": source_request,
        "outline_id": outline_id,
        "language": _resolve_language(
            source_request.get("language", "")
            or (confirmed_outline.get("language", "") if isinstance(confirmed_outline, dict) else "")
        ),
        "reference_titles": [
            ref.get("title", "") for ref in DYNAMIC_KNOWLEDGE_BASE.get("current_references", [])
        ],
        "references": list(DYNAMIC_KNOWLEDGE_BASE.get("current_references", [])),
        "material_ids": [material.get("material_id", "") for material in _active_materials()],
        "materials": list(_active_materials()),
        "source_material": str(DYNAMIC_KNOWLEDGE_BASE.get("source_material", "") or ""),
        "result_evidence": str(DYNAMIC_KNOWLEDGE_BASE.get("result_evidence", "") or ""),
        "material_policy": _material_policy(),
        "asset_plan": list(DYNAMIC_KNOWLEDGE_BASE.get("asset_plan", [])),
        "document_profile": _document_profile(),
        "front_matter": _front_matter(title),
        "asset_root_template": "data/assets/{job_id}",
        "caption_numbering": "chapter",
        "bibliography_style": "sequential_numeric",
    }


def _activate_pending_outline(pending: dict) -> dict | None:
    """Restore the selected outline's scoped state before starting its paper job."""
    staged_material_ids = []
    if DYNAMIC_KNOWLEDGE_BASE.get("active_material_request_fingerprint") == "staged":
        staged_material_ids = list(DYNAMIC_KNOWLEDGE_BASE.get("active_material_ids") or [])
    outline_result = dict(pending.get("outline_result") or {})
    outline = outline_result.get("outline")
    if not isinstance(outline, dict):
        return None
    DYNAMIC_KNOWLEDGE_BASE["current_outline"] = outline
    DYNAMIC_KNOWLEDGE_BASE["asset_plan"] = list(
        outline_result.get("asset_plan") or build_asset_plan(
            outline.get("sections", []), topic=str(outline.get("title", ""))
        )
    )
    material_ids = outline_result.get("active_material_ids")
    if isinstance(material_ids, list):
        DYNAMIC_KNOWLEDGE_BASE["active_material_ids"] = list(dict.fromkeys(list(material_ids) + staged_material_ids))
        DYNAMIC_KNOWLEDGE_BASE["source_material"] = _material_text()
        DYNAMIC_KNOWLEDGE_BASE["result_evidence"] = _result_evidence()
    DYNAMIC_KNOWLEDGE_BASE["active_material_request_fingerprint"] = str(
        pending.get("request_fingerprint")
        or request_fingerprint(dict(pending.get("source_request") or {}))
    )
    references = outline_result.get("references")
    if isinstance(references, list):
        DYNAMIC_KNOWLEDGE_BASE["current_references"] = list(references)
        DYNAMIC_KNOWLEDGE_BASE["current_reference_topic"] = str(outline.get("title", ""))
    sufficiency = outline_result.get("material_sufficiency")
    if isinstance(sufficiency, dict):
        DYNAMIC_KNOWLEDGE_BASE["material_sufficiency"] = dict(sufficiency)
    save_json(DYNAMIC_KNOWLEDGE_BASE, "data/checkpoints/knowledge_base.json")
    return outline


def _set_material_policy_from_write_request(
    allow_degraded_writing: bool = False,
    user_acknowledgement: str = "",
) -> dict | None:
    acknowledgement = user_acknowledgement.strip()
    if allow_degraded_writing and not acknowledgement:
        return {
            "status": "error",
            "message": "user_acknowledgement is required when allow_degraded_writing is true",
        }
    DYNAMIC_KNOWLEDGE_BASE["material_policy"] = {
        "allow_degraded_writing": bool(allow_degraded_writing),
        "user_acknowledgement": acknowledgement if allow_degraded_writing else "",
    }
    save_json(DYNAMIC_KNOWLEDGE_BASE, "data/checkpoints/knowledge_base.json")
    return None


def _material_waiting_message(action: str) -> str:
    return (
        f"资料不足，{action}。可补充真实数据继续；如果用户已明确确认无法提供缺失材料"
        "并接受论文质量与真实性风险，请以 allow_degraded_writing=true 且填写 "
        "user_acknowledgement 后重试当前工具。"
    )


def _planned_section_titles(sections: list[dict] | None = None) -> dict[str, str]:
    section_titles = {}
    for section in sections or []:
        section_id = str(section.get("id") or section.get("section_id") or section.get("title", ""))
        section_titles[section_id] = str(section.get("title") or section_id)
    return section_titles


def _prepare_planned_image_candidates(
    *,
    title: str,
    asset_plan: list[dict] | None = None,
    sections: list[dict] | None = None,
    limit_per_requirement: int = 3,
) -> dict:
    plan = list(asset_plan or DYNAMIC_KNOWLEDGE_BASE.get("asset_plan", []))
    searched_requirements = [
        item for item in plan
        if "searched_image" in str(item.get("asset_type", ""))
    ]
    if not searched_requirements:
        return {"status": "not_required", "candidates": [], "count": 0}

    # Never display cached candidates from an older manuscript as candidates
    # for the current outline.
    registered = _image_assets.register_candidates(
        search_planned_commons_candidates(
            title,
            searched_requirements,
            section_titles=_planned_section_titles(sections),
            limit_per_requirement=limit_per_requirement,
        )
    )
    candidates = registered
    return {
        "status": "candidates_ready" if candidates else "search_empty",
        "approval_required": bool(candidates),
        "requirements": searched_requirements,
        "new_candidates": registered,
        "candidates": candidates,
        "count": len(candidates),
        "message": (
            "Image candidates were searched from the asset plan. Approve selected asset_ids before final rendering."
            if candidates
            else "Image assets are planned but no open-source candidate was found."
        ),
    }


def _auto_attach_open_image_assets(
    *,
    title: str,
    asset_plan: list[dict],
    sections: list[dict],
    job_id: str = "",
    source_fingerprint: str = "",
) -> tuple[list[dict], list[str]]:
    """Best-effort Commons asset search for the output-first whole-paper path."""
    warnings = []
    searched_requirements = [
        requirement for requirement in asset_plan
        if "searched_image" in str(requirement.get("asset_type", ""))
    ]
    if not searched_requirements:
        return [], warnings
    try:
        candidates = search_planned_commons_candidates(
            title,
            searched_requirements,
            section_titles=_planned_section_titles(sections),
            limit_per_requirement=2,
        )
    except Exception as exc:
        return [], [f"开放源图片检索失败: {exc}"]

    assets = []
    chosen_keys = set()
    unresolved_candidate_warnings: dict[tuple, list[str]] = {}
    for candidate in candidates:
        key = (candidate.get("section_id"), candidate.get("purpose"))
        if key in chosen_keys:
            continue
        candidate["approved"] = True
        candidate["job_id"] = job_id
        candidate["source_fingerprint"] = source_fingerprint
        candidate["retrieved_at"] = time.strftime("%Y-%m-%dT%H:%M:%S")
        selected_id = ""
        try:
            registered = _image_assets.register_candidates([candidate])
            selected_id = registered[0].get("asset_id")
            _image_assets.download_approved(
                Path(_task_asset_root(job_id)) / "images",
                asset_ids=[selected_id],
            )
            downloaded = next(
                (
                    asset for asset in _image_assets.approved_assets()
                    if asset.get("asset_id") == selected_id and _renderable_asset_path(asset)
                ),
                None,
            )
        except Exception as exc:
            unresolved_candidate_warnings.setdefault(key, []).append(f"开放源图片缓存失败: {exc}")
            downloaded = None
        if downloaded:
            assets.append(downloaded)
            chosen_keys.add(key)
            unresolved_candidate_warnings.pop(key, None)
            continue
        stored = next(
            (
                asset for asset in _image_assets.approved_assets()
                if selected_id and asset.get("asset_id") == selected_id
            ),
            {},
        )
        if stored.get("download_error"):
            unresolved_candidate_warnings.setdefault(key, []).append(
                f"开放源图片下载失败 ({stored.get('source', '')}, {stored.get('query', '')}): "
                f"{stored['download_error']}"
            )
    for key, key_warnings in unresolved_candidate_warnings.items():
        if key not in chosen_keys:
            warnings.extend(key_warnings)
    if not candidates:
        return [], ["开放源图片检索未返回可插入候选。"]
    if not assets:
        warnings.append("开放源图片未形成可渲染本地资产。")
    return assets, warnings


def _current_job_image_assets(
    job_id: str,
    auto_assets: list[dict] | None = None,
    material_ids: list[str] | None = None,
) -> list[dict]:
    """Keep whole-paper rendering from inheriting image assets from older papers."""
    auto_ids = {asset.get("asset_id") for asset in (auto_assets or [])}
    active_material_ids = set(material_ids or [])
    return [
        asset for asset in _image_assets.approved_assets()
        if _renderable_asset_path(asset)
        and (
            asset.get("asset_id") in auto_ids
            or (job_id and asset.get("job_id") == job_id)
            or (asset.get("material_id") in active_material_ids)
        )
    ]


def _auto_bind_active_material_images(
    *,
    materials: list[dict],
    sections: list[dict],
    asset_plan: list[dict],
    job_id: str,
) -> tuple[list[dict], list[str]]:
    """Bind parsed user images into this task instead of silently dropping them."""
    warnings = []
    assets = []
    methods_section = next(
        (
            str(item.get("section_id") or "")
            for item in asset_plan
            if str(item.get("purpose") or "") == "system_or_device_context"
        ),
        "",
    )
    results_section = next(
        (
            str(item.get("section_id") or "")
            for item in asset_plan
            if str(item.get("asset_type") or "") in {"data_figure", "data_table"}
        ),
        "",
    )
    fallback_section = str((sections[0] if sections else {}).get("id") or "")
    image_index = 0
    for material in materials:
        if material.get("parse_status") != "parsed" or material.get("material_role") != "image_asset":
            continue
        if not _renderable_asset_path(material):
            continue
        image_index += 1
        embedded = str(material.get("metadata", {}).get("source_kind") or "") == "docx_embedded_image"
        section_id = (methods_section or fallback_section) if embedded else (results_section or methods_section or fallback_section)
        if not section_id:
            warnings.append(f"用户图片无法自动定位章节，未插入: {material.get('material_id', '')}")
            continue
        filename = str(material.get("metadata", {}).get("filename") or material.get("material_id") or "image")
        if embedded:
            caption = f"用户材料内嵌图像 {image_index}（题注与章节归属待核验）"
            purpose = "source_document_illustration"
            warnings.append(f"用户材料内嵌图片已插入，题注与章节归属需核验: {filename}")
        else:
            caption = f"用户上传结果图：{filename}"
            purpose = "user_provided_result_image"
        asset = _image_assets.bind_uploaded_material(
            material,
            section_id=section_id,
            purpose=purpose,
            caption=caption,
            job_id=job_id,
            output_dir=Path(_task_asset_root(job_id)) / "images",
        )
        assets.append(asset)
    return assets, warnings


def _auto_generate_result_figures(
    *,
    title: str,
    asset_plan: list[dict],
    job_id: str = "",
    language: str = "中文",
    result_evidence: str | None = None,
    material_policy: dict | None = None,
) -> tuple[list[dict], list[dict], list[str], str]:
    """Generate real-data charts or explicitly labelled simulated result assets."""
    evidence = _result_evidence().strip() if result_evidence is None else str(result_evidence).strip()
    policy = _material_policy() if material_policy is None else dict(material_policy)
    requirements = [
        item for item in asset_plan
        if str(item.get("asset_type", "")) == "data_figure"
    ]
    if not requirements:
        return [], [], [], ""
    if not evidence:
        if policy.get("allow_degraded_writing"):
            figures, tables, summary = build_simulated_result_assets(
                title,
                requirements[0].get("section_id", "results"),
                job_id=job_id or "paper",
                output_dir=Path(_task_asset_root(job_id)) / "figures",
                language=language,
            )
            return (
                figures,
                tables,
                ["真实结果数据缺失：已插入明确标注的模拟结果图表，正式使用前必须替换为真实实验数据。"],
                summary,
            )
        return [], [], ["未提供真实结果数据，且未确认降级写作，未生成性能图表。"], ""
    requirement = requirements[0]
    try:
        generated = generate_figure(
            (
                f"论文题目: {title}\n"
                f"绘图目的: {requirement.get('purpose', 'result_evidence')}\n"
                "仅使用下述已摄取结果证据绘制统计图，不得补造数值：\n"
                f"{evidence}"
            ),
            nexent_client=_nexent_client,
            sandbox=_sandbox,
        )
    except Exception as exc:
        return [], [], [f"结果图表生成失败: {exc}"], evidence
    if not generated.success or not generated.figure_paths:
        return [], [], ["结果图表生成失败，已跳过图表插入。"], evidence
    result_folder = Path(_task_asset_root(job_id)) / "figures"
    result_folder.mkdir(parents=True, exist_ok=True)
    local_figure_paths = []
    for idx, source_path in enumerate(generated.figure_paths, start=1):
        suffix = Path(source_path).suffix or ".png"
        target = result_folder / f"{job_id or 'paper'}_result_{idx}{suffix}"
        shutil.copy2(source_path, target)
        local_figure_paths.append(str(target.resolve()))
    figures = [
        {
            "asset_id": f"figure-{job_id or 'paper'}-{idx}",
            "approved": True,
            "local_path": path,
            "section_id": requirement.get("section_id", ""),
            "purpose": requirement.get("purpose", "result_evidence"),
            "source": "generated_from_result_dataset",
            "attribution": "根据已上传结果数据生成",
            "caption": "基于上传结果数据生成的性能比较图",
            "english_caption": "Performance comparison generated from uploaded result data",
            "job_id": job_id,
        }
        for idx, path in enumerate(local_figure_paths, start=1)
    ]
    return figures, [], [], evidence


def _results_section(title: str) -> bool:
    lowered = str(title or "").strip().lower()
    return "结果" in lowered or "result" in lowered


def _enforce_simulated_results_disclosure(
    result: dict,
    language: str,
    *,
    include_section_notice: bool = True,
) -> dict:
    content = str(result.get("content") or "")
    marker = "模拟数据说明" if _resolve_language(language) == "中文" else "Simulated Data Notice"
    if _resolve_language(language) == "中文":
        replacements = {
            "真实空气质量监测数据集": "模拟空气质量占位数据集",
            "真实空气质量监测数据": "模拟空气质量占位数据",
            "真实实验观测": "模拟占位观测",
            "真实观测数据": "模拟占位数据",
            "真实数据集": "模拟占位数据集",
            "在真实": "在模拟",
            "非模拟占位数据": "属于模拟占位数据",
            "真实值": "模拟目标值",
            "本文实验采用": "为展示成稿结构，本文模拟实验采用",
            "所有结果均为测试集上的五次独立实验取平均": "以下结果均为模拟测试集上的占位数值，不代表已完成真实实验",
            "进行了迁移测试": "设计了模拟迁移测试",
        }
        for original, replacement in replacements.items():
            content = content.replace(original, replacement)
    else:
        replacements = {
            "real-world observations": "simulated placeholder observations",
            "real observations": "simulated placeholder observations",
            "real dataset": "simulated placeholder dataset",
            "real data": "simulated placeholder data",
            "non-simulated placeholder data": "simulated placeholder data",
            "predicted and actual values": "predicted and simulated target values",
        }
        for original, replacement in replacements.items():
            content = content.replace(original, replacement)
    if not include_section_notice or marker in content:
        updated = dict(result)
        updated["content"] = content
        return updated
    notice = (
        "## 模拟数据说明\n\n本节中的数值、表格与图形为材料不足条件下生成的模拟占位数据，"
        "不属于真实实验观测结果，正式提交或据此形成结论前必须替换为真实数据并重新分析。\n\n"
        if _resolve_language(language) == "中文"
        else "## Simulated Data Notice\n\nThe numerical results, tables, and figures in this section are "
        "simulated placeholders generated because empirical data were unavailable. They must be replaced "
        "with real observations and re-analysed before formal use.\n\n"
    )
    updated = dict(result)
    updated["content"] = notice + content
    return updated


def _auto_generate_method_figures(
    *,
    title: str = "",
    asset_plan: list[dict],
    language: str,
    job_id: str = "",
    output_dir: str | Path | None = None,
) -> tuple[list[dict], list[str]]:
    """Always provide a truthful conceptual workflow visual when methods require one."""
    requirements = [
        item for item in asset_plan
        if str(item.get("purpose", "")) == "system_or_device_context"
    ]
    if not requirements:
        return [], []
    requirement = requirements[0]
    folder = Path(output_dir) if output_dir is not None else Path(_task_asset_root(job_id)) / "figures"
    output_path = folder / f"{job_id or 'paper'}_system_workflow.png"
    try:
        figure_path = render_system_workflow_diagram(output_path, language=language, topic=title)
    except Exception as exc:
        return [], [f"系统流程示意图生成失败: {exc}"]
    lowered_title = str(title or "").lower()
    chinese = _resolve_language(language) == "中文"
    if any(token in lowered_title for token in ("空气质量", "air quality", "pm2.5")):
        caption = (
            "空气质量站点数据、时空图网络与预测输出流程示意图（概念图，不表示实验结果）"
            if chinese else
            "Conceptual workflow for station data, spatiotemporal graph modelling and air-quality forecasting (not an experimental result)"
        )
    elif any(token in lowered_title for token in ("脑电", "脑机", "eeg", "bci", "运动想象", "机械臂")):
        caption = (
            "基于运动想象脑电的机械臂控制与触觉反馈闭环流程示意图（概念图，不表示实验结果）"
            if chinese else
            "Conceptual closed-loop workflow for motor-imagery EEG robotic-arm control and tactile feedback (not an experimental result)"
        )
    else:
        caption = (
            "输入数据、模型推理与输出评估流程示意图（概念图，不表示实验结果）"
            if chinese else
            "Conceptual workflow for input data, model inference and output evaluation (not an experimental result)"
        )
    workflow = {
        "asset_id": f"diagram-{job_id or 'paper'}-system",
        "approved": True,
        "local_path": figure_path,
        "section_id": requirement.get("section_id", "methods"),
        "purpose": "system_or_device_context",
        "source": "generated_conceptual_diagram",
        "source_url": "generated_locally",
        "license": "generated_by_system",
        "attribution": "paper_agent_v6",
        "caption": caption,
        "english_caption": "Conceptual research workflow (not an experimental result)",
        "job_id": job_id,
    }
    intro_requirement = next(
        (item for item in asset_plan if str(item.get("purpose", "")) == "application_context"),
        requirement,
    )
    framework_path = folder / f"{job_id or 'paper'}_research_framework.png"
    try:
        framework_image = render_research_framework_diagram(framework_path, language=language, topic=title)
    except Exception as exc:
        return [workflow], [f"研究框架图生成失败: {exc}"]
    framework = {
        "asset_id": f"diagram-{job_id or 'paper'}-framework",
        "approved": True,
        "local_path": framework_image,
        "section_id": intro_requirement.get("section_id", requirement.get("section_id", "methods")),
        "purpose": "research_framework",
        "source": "generated_conceptual_diagram",
        "source_url": "generated_locally",
        "license": "generated_by_system",
        "attribution": "paper_agent_v6",
        "caption": "研究框架与关键要素关系图（概念图，不表示实验结果）" if chinese else "Research framework and key elements (conceptual diagram, not an experimental result)",
        "english_caption": "Research framework and key elements (conceptual diagram, not an experimental result)",
        "job_id": job_id,
    }
    return [workflow, framework], []


def _ingest_material_input(raw_path: str) -> dict:
    local_path = raw_path
    downloaded = False
    if not os.path.exists(raw_path) and (
        raw_path.startswith("s3://") or "presigned_url=" in raw_path or raw_path.startswith("http")
    ):
        local_path = _download_from_nexent(raw_path)
        downloaded = bool(local_path)
    if not local_path:
        return {
            "materials": [],
            "failed": [{"source_path": raw_path, "error": "download_failed"}],
            "unsupported": [],
        }
    report = _materials.ingest_paths([local_path])
    for material in report["materials"]:
        logger.info(
            "material ingested: role=%s type=%s path=%s",
            material.get("material_role"),
            material.get("file_type"),
            material.get("local_path"),
        )
    for failure in report["failed"]:
        logger.warning("material ingest failed: %s", failure)
    for unsupported in report["unsupported"]:
        logger.warning("material unsupported: %s", unsupported.get("source_path"))
    for material in report["materials"]:
        if downloaded:
            material.setdefault("metadata", {})["nexent_source_path"] = raw_path
    return report


async def _notify_progress(
    ctx: Context | None,
    *,
    progress: float,
    total: float | None,
    message: str,
) -> None:
    """Best-effort MCP progress + log notification for Nexent long tasks."""
    if not ctx:
        return
    try:
        await ctx.report_progress(progress=progress, total=total, message=message)
    except Exception as exc:
        logger.debug(f"progress notification failed: {exc}")
    try:
        await ctx.info(message)
    except Exception as exc:
        logger.debug(f"info notification failed: {exc}")


async def _run_with_heartbeat(
    func,
    *args,
    ctx: Context | None = None,
    progress: float = 0,
    total: float | None = None,
    message: str = "处理中",
    heartbeat_seconds: int = 20,
    **kwargs,
):
    """Run blocking local/LLM work while keeping the MCP request active."""
    if not ctx:
        return await asyncio.to_thread(func, *args, **kwargs)

    stop = asyncio.Event()

    async def heartbeat():
        tick = 0
        while not stop.is_set():
            suffix = "" if tick == 0 else f"（持续处理中 {tick * heartbeat_seconds}s）"
            await _notify_progress(
                ctx,
                progress=progress,
                total=total,
                message=f"{message}{suffix}",
            )
            tick += 1
            try:
                await asyncio.wait_for(stop.wait(), timeout=heartbeat_seconds)
            except asyncio.TimeoutError:
                continue

    task = asyncio.create_task(heartbeat())
    try:
        return await asyncio.to_thread(func, *args, **kwargs)
    finally:
        stop.set()
        await task


def _verify_docx_output(output_path: str, expected_sections: int) -> tuple[bool, str]:
    """Reject missing or title-only DOCX artifacts before reporting completion."""
    if not output_path or not os.path.exists(output_path):
        return False, "Word 输出文件不存在"
    if os.path.getsize(output_path) < 1024:
        return False, "Word 输出文件过小"
    try:
        from docx import Document
        doc = Document(output_path)
        text_paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    except Exception as exc:
        return False, f"Word 输出自检失败: {exc}"
    minimum_blocks = max(3, min(expected_sections + 1, 8))
    if len(text_paragraphs) < minimum_blocks:
        return False, f"Word 输出正文块不足: {len(text_paragraphs)}"
    return True, ""


def _literature_queries(title: str, sections: list[dict] | None = None) -> list[str]:
    """Construct focused queries so the reference pool is topical rather than merely large."""
    title_text = str(title or "").strip()
    lowered = title_text.lower()
    queries = [title_text]
    has_graph = any(token in lowered for token in ("图神经", "graph neural"))
    has_air_quality = any(token in lowered for token in ("空气质量", "air quality", "pm2.5"))
    if any(token in lowered for token in ("脑电", "脑机", "eeg", "bci", "运动想象")):
        queries.extend([
            "motor imagery EEG brain-computer interface robotic arm control",
            "noninvasive EEG brain-computer interface tactile feedback rehabilitation",
            "motor imagery EEG classification CSP SVM deep learning",
        ])
    if has_graph and has_air_quality:
        queries.extend([
            "spatiotemporal graph neural network air quality forecasting",
            "graph neural network PM2.5 prediction monitoring stations",
        ])
    elif has_graph or any(token in lowered for token in ("分子", "molecular")):
        queries.extend([
            "graph neural network molecular property prediction",
            "message passing neural network molecular representation learning",
        ])
    if has_air_quality:
        queries.extend([
            "air quality prediction machine learning PM2.5",
            "deep learning air pollution time series forecasting",
        ])
    for section in (sections or []):
        section_title = str(section.get("title") or "")
        if _results_section(section_title) or "方法" in section_title or "method" in section_title.lower():
            queries.append(f"{title_text} {section_title}")
    return queries


def _build_reference_pool(title: str, sections: list[dict] | None = None, *, force: bool = False) -> list[dict]:
    frozen_topic = str(DYNAMIC_KNOWLEDGE_BASE.get("current_reference_topic") or "").strip()
    frozen_references = list(DYNAMIC_KNOWLEDGE_BASE.get("current_references") or [])
    if not force and frozen_topic == str(title or "").strip() and frozen_references:
        return frozen_references
    literature_config = _APP_CONFIG.get("literature_search", {})
    papers = search_paper_pool(
        _literature_queries(title, sections),
        target_count=int(literature_config.get("default_limit", 15)),
        per_query_limit=int(literature_config.get("per_query_limit", 8)),
    )
    DYNAMIC_KNOWLEDGE_BASE["current_references"] = papers
    DYNAMIC_KNOWLEDGE_BASE["current_reference_topic"] = str(title or "").strip()
    return papers


# ============================================================
# MCP Tool 1: 文献检索
# ============================================================
@mcp.tool()
def literature_search(query: str, limit: int = 15) -> dict:
    """根据课题检索真实学术文献，灌入知识库"""
    papers = search_paper_pool(
        _literature_queries(query),
        target_count=max(8, min(int(limit or 15), 30)),
        per_query_limit=int(_APP_CONFIG.get("literature_search", {}).get("per_query_limit", 8)),
    )
    DYNAMIC_KNOWLEDGE_BASE["current_references"] = papers
    DYNAMIC_KNOWLEDGE_BASE["current_reference_topic"] = query
    save_json(DYNAMIC_KNOWLEDGE_BASE, "data/checkpoints/knowledge_base.json")
    return {
        "status": "success", "query": query, "count": len(papers),
        "message": f"已检索到 {len(papers)} 篇真实文献。",
        "papers": [{"citation_key": p.get("citation_key"), "title": p.get("title"),
                     "authors": p.get("authors"), "year": p.get("year")} for p in papers],
    }


@mcp.tool()
def ingest_materials(material_paths_json: str = "", material_path: str = "") -> dict:
    """Ingest paper source text, result datasets, ZIP bundles, or uploaded image materials."""
    raw_paths = []
    if material_paths_json:
        try:
            parsed = json.loads(material_paths_json)
        except json.JSONDecodeError:
            return {"status": "error", "message": "material_paths_json is not valid JSON"}
        raw_paths = parsed if isinstance(parsed, list) else [str(parsed)]
    if material_path:
        raw_paths.append(material_path)
    if not raw_paths:
        return {"status": "error", "message": "provide material_path or material_paths_json"}

    materials, failed, unsupported = [], [], []
    for raw_path in raw_paths:
        report = _ingest_material_input(str(raw_path))
        materials.extend(report["materials"])
        failed.extend(report["failed"])
        unsupported.extend(report["unsupported"])
    _set_active_materials(
        materials,
        replace=DYNAMIC_KNOWLEDGE_BASE.get("active_material_request_fingerprint") != "staged",
    )
    DYNAMIC_KNOWLEDGE_BASE["active_material_request_fingerprint"] = "staged"
    DYNAMIC_KNOWLEDGE_BASE["source_material"] = _material_text()
    DYNAMIC_KNOWLEDGE_BASE["result_evidence"] = _result_evidence()
    DYNAMIC_KNOWLEDGE_BASE["material_sufficiency"] = _material_sufficiency()
    save_json(DYNAMIC_KNOWLEDGE_BASE, "data/checkpoints/knowledge_base.json")
    return {
        "status": "ingested" if materials else "ingest_failed",
        "materials": materials,
        "failed": failed,
        "unsupported": unsupported,
        "material_summaries": _materials.summaries(),
        "material_sufficiency": DYNAMIC_KNOWLEDGE_BASE["material_sufficiency"],
    }


@mcp.tool()
def list_materials() -> dict:
    """List parsed paper materials and their roles without exposing full extracted text."""
    return {
        "status": "success",
        "count": len(_materials.all()),
        "materials": _materials.summaries(),
    }


@mcp.tool()
def bind_uploaded_image_asset(
    material_id: str,
    section_id: str,
    purpose: str = "",
    caption: str = "",
) -> dict:
    """Bind a user-uploaded image material to a paper section for final Word rendering."""
    material = _materials.get(material_id)
    if not material or material.get("material_role") != "image_asset":
        return {"status": "error", "message": "material_id is not an uploaded image material"}
    try:
        asset = _image_assets.bind_uploaded_material(
            material, section_id=section_id, purpose=purpose, caption=caption
        )
    except ValueError as exc:
        return {"status": "error", "message": str(exc)}
    return {"status": "bound", "asset": asset}


@mcp.tool()
def check_material_sufficiency(
    topic: str = "",
    material_text: str = "",
    result_evidence: str = "",
    references_json: str = "",
    allow_degraded_writing: bool = False,
    user_acknowledgement: str = "",
) -> dict:
    """Block manuscript writing when the available material cannot support a full paper."""
    policy_error = _set_material_policy_from_write_request(
        allow_degraded_writing=allow_degraded_writing,
        user_acknowledgement=user_acknowledgement,
    )
    if policy_error:
        return policy_error
    references = DYNAMIC_KNOWLEDGE_BASE.get("current_references", [])
    if references_json:
        try:
            references = json.loads(references_json)
        except json.JSONDecodeError:
            return {"status": "error", "message": "references_json is not valid JSON"}
    policy = _material_policy()
    report = assess_material_sufficiency(
        topic=topic,
        material_text=material_text or _material_text(),
        result_evidence=result_evidence or _result_evidence(),
        references=references,
        approved_assets=_image_assets.approved_assets(),
        materials=_materials.all(),
        allow_degraded_writing=policy["allow_degraded_writing"],
        degraded_reason=policy["user_acknowledgement"],
    )
    DYNAMIC_KNOWLEDGE_BASE["source_material"] = material_text or _material_text()
    DYNAMIC_KNOWLEDGE_BASE["result_evidence"] = result_evidence or _result_evidence()
    DYNAMIC_KNOWLEDGE_BASE["material_sufficiency"] = report
    save_json(DYNAMIC_KNOWLEDGE_BASE, "data/checkpoints/knowledge_base.json")
    return report


@mcp.tool()
def set_material_policy(
    allow_degraded_writing: bool = False,
    user_acknowledgement: str = "",
) -> dict:
    """
    Store the user's evidence policy for this paper.

    Set allow_degraded_writing only after the user confirms missing materials
    cannot be provided and accepts a lower-confidence manuscript based on the
    available materials and literature.
    """
    acknowledgement = user_acknowledgement.strip()
    if allow_degraded_writing and not acknowledgement:
        return {
            "status": "error",
            "message": "user_acknowledgement is required before degraded writing can be enabled",
        }
    DYNAMIC_KNOWLEDGE_BASE["material_policy"] = {
        "allow_degraded_writing": bool(allow_degraded_writing),
        "user_acknowledgement": acknowledgement,
    }
    report = _material_sufficiency()
    DYNAMIC_KNOWLEDGE_BASE["material_sufficiency"] = report
    save_json(DYNAMIC_KNOWLEDGE_BASE, "data/checkpoints/knowledge_base.json")
    return {
        "status": "updated",
        "material_policy": _material_policy(),
        "material_sufficiency": report,
        "message": (
            "Degraded writing is enabled for this paper with a quality-risk disclosure."
            if allow_degraded_writing
            else "Strict evidence gating is enabled for this paper."
        ),
    }


@mcp.tool()
def search_image_candidates(
    query: str,
    purpose: str = "",
    section_id: str = "",
    limit: int = 5,
) -> dict:
    """Search traceable Wikimedia Commons image candidates without inserting them."""
    try:
        candidates = _image_assets.register_candidates(
            search_commons_candidates(query, purpose=purpose, section_id=section_id, limit=limit)
        )
    except Exception as e:
        return {"status": "error", "message": f"image search failed: {e}", "candidates": []}
    return {
        "status": "candidates_ready",
        "approval_required": True,
        "count": len(candidates),
        "candidates": candidates,
        "message": "Review candidate source and license metadata, then approve selected asset_ids.",
    }


@mcp.tool()
def approve_image_assets(asset_ids: str) -> dict:
    """Approve searched image candidates and cache them locally for Word rendering."""
    ids = []
    try:
        parsed = json.loads(asset_ids)
        ids = parsed if isinstance(parsed, list) else [str(parsed)]
    except json.JSONDecodeError:
        ids = [item.strip() for item in asset_ids.split(",") if item.strip()]
    approved = _image_assets.approve(ids)
    try:
        approved = _image_assets.download_approved()
    except Exception as e:
        return {
            "status": "approved_download_failed",
            "approved_assets": approved,
            "message": f"assets approved but local download failed: {e}",
        }
    return {
        "status": "approved",
        "count": len(approved),
        "approved_assets": approved,
    }


@mcp.tool()
def prepare_planned_image_assets(limit_per_requirement: int = 3) -> dict:
    """Search open-source image candidates required by the current outline asset plan."""
    outline = DYNAMIC_KNOWLEDGE_BASE.get("current_outline") or {}
    title = str(outline.get("title") or "")
    if not title:
        return {"status": "error", "message": "generate an outline before preparing planned images"}
    try:
        return _prepare_planned_image_candidates(
            title=title,
            asset_plan=list(DYNAMIC_KNOWLEDGE_BASE.get("asset_plan", [])),
            sections=list(outline.get("sections", [])),
            limit_per_requirement=limit_per_requirement,
        )
    except Exception as exc:
        return {"status": "search_failed", "message": f"planned image search failed: {exc}"}


# ============================================================
# MCP Tool 2: 生成大纲
# ============================================================
@mcp.tool()
def generate_outline(
    document_path: str = "", requirements: str = "", topic: str = "", language: str = "",
    force_regenerate: bool = False,
) -> dict:
    """分析文档/主题，生成论文大纲（含 Nature 结构校验）"""
    target_language = _resolve_language(language)
    source_request = {
        "document_path": document_path.strip(),
        "topic": topic.strip(),
        "requirements": requirements.strip(),
        "language": target_language,
    }
    if not force_regenerate:
        pending = _pending_outlines.find_open(source_request)
        if pending:
            _activate_pending_outline(pending)
            cached = dict(pending.get("outline_result") or {})
            return {
                **cached,
                "status": "outline_confirmation_required",
                "outline_id": pending["outline_id"],
                "confirmation_tool": "confirm_outline_and_start_writing",
                "confirmation_without_id_supported": True,
                "reused_pending_outline": True,
                "message": (
                    "该请求已存在待确认大纲，已返回原大纲而未重新生成。"
                    "用户确认后调用 confirm_outline_and_start_writing(outline_id=...) 启动整篇写作；"
                    "如后续轮次无法取得 outline_id，直接省略该参数以确认最近待处理大纲，"
                    "不得为了取回 ID 再调用 generate_outline。"
                ),
            }
    retain_staged_materials = (
        DYNAMIC_KNOWLEDGE_BASE.get("active_material_request_fingerprint") == "staged"
        and bool(DYNAMIC_KNOWLEDGE_BASE.get("active_material_ids"))
    )
    source_fingerprint = _prepare_outline_request_scope(source_request)
    doc_content = _material_text(); file_hint = ""; _temp_files = []
    if document_path:
        ingest_report = _ingest_material_input(document_path)
        if ingest_report["failed"] or (
            not ingest_report["materials"] and ingest_report["unsupported"]
        ):
            return {
                "status": "material_error",
                "message": f"cannot read uploaded material {_extract_filename(document_path)}",
                "failed": ingest_report["failed"],
                "unsupported": ingest_report["unsupported"],
            }
        _set_active_materials(ingest_report["materials"], replace=not retain_staged_materials)
        doc_content = _material_text()
        if os.path.exists(document_path) and not doc_content:
            try:
                doc_content = read_document(document_path)
            except Exception as e:
                logger.warning(f"读取文档失败: {e}")
                file_hint = _extract_filename(document_path)
        elif not doc_content and (document_path.startswith("s3://") or "presigned_url=" in document_path):
            local_path = _download_from_nexent(document_path)
            if local_path:
                _temp_files.append(local_path)
                try:
                    doc_content = read_document(local_path)
                except Exception:
                    file_hint = _extract_filename(document_path)
            else:
                file_hint = _extract_filename(document_path)
        else:
            file_hint = _extract_filename(document_path)
    for t in _temp_files:
        try: os.unlink(t)
        except: pass
    if doc_content: pass
    elif topic: doc_content = f"论文主题：{topic}"
    elif document_path and not doc_content:
        # 文件下载失败且没有 topic，明确告知用户
        return {"status": "error", "message": f"无法读取文件「{_extract_filename(document_path)}」，请确认文件已上传成功。如有必要，直接提供论文主题。"}
    elif requirements: doc_content = f"论文主题：{requirements}"
    else: doc_content = "通用学术论文"

    system_prompt = f"""你是一个学术论文写作助手。遵循 Nature 期刊标准生成论文大纲。
大纲应包含: 摘要(Abstract)、引言(Introduction)、方法(Methods)、结果(Results)、讨论(Discussion)、结论(Conclusion)
遵循 hourglass 结构：宽 -> 窄 -> 宽。
目标写作语言：{target_language}。章节标题与 key_points 必须使用目标写作语言，除专业缩写外不得混用其他语言。"""
    user_prompt = f"""生成论文大纲：

{'_'*40}
{doc_content}
{'_'*40}

要求：{requirements or '生成 5-8 个章节'}
目标语言：{target_language}
输出 JSON 格式：
{{"title": "论文标题", "sections": [{{"id": "sec1", "title": "引言", "key_points": [...]}}]}}"""
    try:
        client = get_llm_client()
        raw = client.call(user_prompt, system=system_prompt, response_format="json")
        raw = raw.strip(); raw = re.sub(r"^```(?:json)?\s*", "", raw); raw = re.sub(r"\s*```$", "", raw)
        outline = _localize_standard_outline(json.loads(raw), target_language)
    except Exception as e:
        return {"status": "error", "message": f"大纲生成失败: {e}"}

    validation = validate_outline(outline.get("sections", []))
    DYNAMIC_KNOWLEDGE_BASE["current_outline"] = outline
    DYNAMIC_KNOWLEDGE_BASE["active_material_request_fingerprint"] = source_fingerprint
    DYNAMIC_KNOWLEDGE_BASE["source_material"] = _material_text() or doc_content
    DYNAMIC_KNOWLEDGE_BASE["result_evidence"] = _result_evidence()
    DYNAMIC_KNOWLEDGE_BASE["asset_plan"] = build_asset_plan(
        outline.get("sections", []), topic=str(outline.get("title", ""))
    )
    try:
        _build_reference_pool(
            outline.get("title", ""),
            outline.get("sections", []),
            force=True,
        )
    except Exception as exc:
        DYNAMIC_KNOWLEDGE_BASE["current_references"] = []
        DYNAMIC_KNOWLEDGE_BASE["current_reference_topic"] = outline.get("title", "")
        logger.warning("大纲阶段文献池建立失败，将在正文阶段重试: %s", exc)
    try:
        planned_images = _prepare_planned_image_candidates(
            title=outline.get("title", ""),
            asset_plan=DYNAMIC_KNOWLEDGE_BASE["asset_plan"],
            sections=outline.get("sections", []),
        )
    except Exception as exc:
        planned_images = {"status": "search_failed", "message": f"planned image search failed: {exc}"}
    DYNAMIC_KNOWLEDGE_BASE["material_sufficiency"] = _material_sufficiency(
        topic=outline.get("title", ""),
        material_text=doc_content,
        result_evidence=_result_evidence(),
    )
    save_json(DYNAMIC_KNOWLEDGE_BASE, "data/checkpoints/knowledge_base.json")
    save_json(outline, "data/checkpoints/outline.json")
    result = {
        "status": (
            "waiting_confirmation"
            if _material_report_allows_writing(DYNAMIC_KNOWLEDGE_BASE["material_sufficiency"])
            else "waiting_materials"
        ), "outline": outline,
        "section_count": len(outline.get("sections", [])),
        "validation": validation,
        "asset_plan": DYNAMIC_KNOWLEDGE_BASE["asset_plan"],
        "planned_image_assets": planned_images,
        "material_sufficiency": DYNAMIC_KNOWLEDGE_BASE["material_sufficiency"],
        "active_material_ids": list(DYNAMIC_KNOWLEDGE_BASE.get("active_material_ids") or []),
        "references": list(DYNAMIC_KNOWLEDGE_BASE.get("current_references") or []),
        "message": (
            "大纲已生成。当前材料已允许进入写作。"
            if _material_report_allows_writing(DYNAMIC_KNOWLEDGE_BASE["material_sufficiency"])
            else "大纲已生成。写作前需补齐 material_sufficiency 中列出的真实材料，或在用户明确接受质量风险后启用降级写作。"
        ),
    }
    pending = _pending_outlines.create(source_request, result)
    return {
        **result,
        "outline_id": pending["outline_id"],
        "confirmation_tool": "confirm_outline_and_start_writing",
        "confirmation_without_id_supported": True,
        "confirmation_instruction": (
            "用户确认此大纲后调用 confirm_outline_and_start_writing。"
            "如果下一轮无法读取 outline_id，省略 outline_id 即可确认最近待处理大纲；"
            "不得重新调用 generate_outline 以获取 ID。"
        ),
    }


# ============================================================
# MCP Tool 3: 章节写作（含 Discriminator 审查）
# ============================================================
@mcp.tool()
def section_writer(
    section_id: str = "", section_title: str = "", key_points: str = "",
    skip_review: bool = False,
    allow_degraded_writing: bool = False,
    user_acknowledgement: str = "",
) -> dict:
    """基于知识库写作单个章节，自动审查（Discriminator），写完回写知识库"""
    policy_error = _set_material_policy_from_write_request(
        allow_degraded_writing=allow_degraded_writing,
        user_acknowledgement=user_acknowledgement,
    )
    if policy_error:
        return policy_error
    papers = DYNAMIC_KNOWLEDGE_BASE.get("current_references", [])
    if not papers:
        current_outline = DYNAMIC_KNOWLEDGE_BASE.get("current_outline") or {}
        papers = _build_reference_pool(
            current_outline.get("title", section_title),
            current_outline.get("sections", []),
        )
    material_report = _material_sufficiency(
        topic=section_title,
        material_text=_material_text() or key_points,
        result_evidence=_result_evidence(),
    )
    if not _material_report_allows_writing(material_report):
        return {
            **material_report,
            "section_id": section_id,
            "section_title": section_title,
            "message": _material_waiting_message("禁止开始章节正文写作"),
            "content": "",
        }

    written = DYNAMIC_KNOWLEDGE_BASE.get("written_chapters", {})
    ctx_sections = list(written.values()) if written else []
    context = get_previous_context(ctx_sections, max_count=3)

    # 获取 nature move 序列
    move_prompt = format_moves_prompt(section_title)

    outline = {"id": section_id or f"sec_{len(written)+1}", "title": section_title,
               "key_points": [k.strip() for k in key_points.split(",") if k.strip()]}

    # 使用增强版 write_section（含 nature move 序列）
    result = write_section_local(outline, papers, context=context, move_sequence=move_prompt)

    # Discriminator 审查
    if not skip_review:
        max_review = _APP_CONFIG.get("discriminator", {}).get("max_review_attempts", 3)
        for attempt in range(max_review):
            review = review_section(result.get("content", ""), section_title)
            if review.get("status") == "PASS":
                logger.info(f"章节「{section_title}」通过 Discriminator 审查")
                break
            if attempt < max_review - 1:
                logger.warning(f"审查未通过，重写 (第 {attempt+1}/{max_review} 次): {review.get('reason', '')[:100]}")
                result = write_section_local(outline, papers, context=context, move_sequence=move_prompt)
        else:
            logger.warning(f"章节「{section_title}」审查超限，强制通过")

    DYNAMIC_KNOWLEDGE_BASE["written_chapters"][section_title] = result
    save_json(DYNAMIC_KNOWLEDGE_BASE, "data/checkpoints/knowledge_base.json")
    save_json(result, f"data/checkpoints/{outline['id']}.json")

    return {
        "status": "success", "section_id": outline["id"], "section_title": section_title,
        "content": result.get("content", ""), "references_used": result.get("references_used", []),
        "total_chapters_written": len(DYNAMIC_KNOWLEDGE_BASE["written_chapters"]),
        "message": f"「{section_title}」写作完成，已存入本地知识库。",
    }


# ============================================================
# MCP Tool 4: 图表生成（含 Debug 循环）
# ============================================================
@mcp.tool()
def generate_figure_tool(context: str = "", section_title: str = "", data_summary: str = "") -> dict:
    """调用 Generator_Coder 生成绘图代码 -> 沙盒执行 -> Debug 循环（最多 3 次）"""
    evidence = data_summary or _result_evidence()
    if not evidence.strip():
        if _material_policy().get("allow_degraded_writing"):
            figures, tables, summary = build_simulated_result_assets(
                DYNAMIC_KNOWLEDGE_BASE.get("current_outline", {}).get("title", "degraded-paper"),
                section_title or "results",
                job_id="manual-figure",
                language=_resolve_language(),
            )
            return {
                "status": "success_simulated",
                "success": True,
                "simulated": True,
                "figure_paths": [figure["local_path"] for figure in figures],
                "tables": tables,
                "data_summary": summary,
                "message": "真实结果数据不足，已生成带明确标识的模拟图表；正式使用前必须替换为真实数据。",
            }
        return {
            "status": "WAITING_REQUIRED_USER_MATERIALS",
            "missing_materials": ["result_evidence"],
            "message": "真实结果数据不足，禁止生成数据图表。",
        }
    ctx = context or f"为章节「{section_title}」生成数据可视化图表"
    ctx = f"{ctx}\n\nEvidence summary:\n{evidence}"
    result = generate_figure(ctx, nexent_client=_nexent_client, sandbox=_sandbox)
    return {
        "status": "success" if result.success else "fallback",
        "success": result.success,
        "figure_paths": result.figure_paths,
        "output": result.output[:500] if result.output else "",
        "error": result.error[:500] if result.error else "",
        "execution_time": round(result.execution_time, 2),
        "message": f"图表生成{'成功' if result.success else '降级为占位图'} ({len(result.figure_paths)} 个文件)",
    }


# ============================================================
# MCP Tool 5: 逐章写作（单章，含 Discriminator 审查）
# ============================================================
async def _write_section_step_pipeline(
    section_id: str = "",
    section_title: str = "",
    key_points: str = "",
    previous_sections_json: str = "",
    language: str = "",
    figure_descriptions: str = "",
    allow_degraded_writing: bool = False,
    user_acknowledgement: str = "",
    mcp_context: Context | None = None,
) -> dict:
    """
    写作单个章节（含 Discriminator 审查循环）。
    每次调用写一章，完成后返回章节内容，Nexent LLM 逐章调用并展示进度。
    收集全部章节后调用 render_final_paper 完成渲染。
    """
    policy_error = _set_material_policy_from_write_request(
        allow_degraded_writing=allow_degraded_writing,
        user_acknowledgement=user_acknowledgement,
    )
    if policy_error:
        return policy_error
    papers = DYNAMIC_KNOWLEDGE_BASE.get("current_references", [])
    if not papers:
        current_outline = DYNAMIC_KNOWLEDGE_BASE.get("current_outline") or {}
        papers = _build_reference_pool(
            current_outline.get("title", section_title),
            current_outline.get("sections", []),
        )
    policy = _material_policy()
    material_report = assess_material_sufficiency(
        topic=section_title,
        material_text=_material_text() or key_points,
        result_evidence=_result_evidence(),
        references=papers,
        approved_assets=_image_assets.approved_assets(),
        materials=_materials.all(),
        allow_degraded_writing=policy["allow_degraded_writing"],
        degraded_reason=policy["user_acknowledgement"],
    )
    DYNAMIC_KNOWLEDGE_BASE["material_sufficiency"] = material_report
    if not _material_report_allows_writing(material_report):
        return {
            **material_report,
            "section_id": section_id,
            "section_title": section_title,
            "message": _material_waiting_message("禁止开始章节正文写作"),
        }

    prev_sections = []
    if previous_sections_json:
        try:
            prev_sections = json.loads(previous_sections_json)
        except json.JSONDecodeError:
            prev_sections = []

    prev_ctx = get_previous_context(prev_sections, max_count=3)
    section_context = build_section_context(section_title, prev_sections, prev_context=prev_ctx)
    move_prompt = format_moves_prompt(section_title)
    if language:
        move_prompt += f"\n\n使用语言: {language}（全部内容用此语言撰写，包括标题、正文和专业术语）"

    # 如果有已生成的图表，注入到写作 prompt
    if figure_descriptions:
        move_prompt += f"\n\n已生成图表:\n{figure_descriptions}\n在正文中引用时使用 [Figure: 图片路径] 格式。"

    outline = {
        "id": section_id or f"sec_{len(prev_sections)+1}",
        "title": section_title,
        "key_points": [k.strip() for k in key_points.split(",") if k.strip()],
    }

    review_enabled = _APP_CONFIG.get("discriminator", {}).get("enabled", True)
    max_review = _APP_CONFIG.get("discriminator", {}).get("max_review_attempts", 3)

    try:
        await _notify_progress(
            mcp_context,
            progress=0,
            total=max_review + 2,
            message=f"开始起草章节「{section_title}」",
        )
        result = await _run_with_heartbeat(
            write_section_local,
            outline,
            papers,
            ctx=mcp_context,
            progress=1,
            total=max_review + 2,
            message=f"正在起草章节「{section_title}」",
            context=section_context,
            move_sequence=move_prompt,
        )

        if review_enabled:
            for ra in range(max_review):
                try:
                    review = await _run_with_heartbeat(
                        review_section,
                        result.get("content", ""),
                        section_title,
                        ctx=mcp_context,
                        progress=ra + 2,
                        total=max_review + 2,
                        message=f"正在审查章节「{section_title}」第 {ra + 1} 轮",
                    )
                    if review.get("status") == "PASS":
                        break
                except Exception as e:
                    logger.warning(f"Discriminator 失败 (第{ra+1}次): {e}")
                    break
                if ra < max_review - 1:
                    result = await _run_with_heartbeat(
                        write_section_local,
                        outline,
                        papers,
                        ctx=mcp_context,
                        progress=ra + 2.5,
                        total=max_review + 2,
                        message=f"正在重写章节「{section_title}」第 {ra + 1} 轮",
                        context=section_context,
                        move_sequence=move_prompt,
                    )

        section_content = result.get("content", "")
        figure_paths = []

        # 存入服务端缓存，供 render_final_paper 使用（避免 LLM JSON 截断问题）
        _SERVER_SECTIONS_CACHE.append({
            "section_id": outline["id"],
            "title": section_title,
            "content": section_content,
            "references_used": result.get("references_used", []),
        })
        await _notify_progress(
            mcp_context,
            progress=max_review + 2,
            total=max_review + 2,
            message=f"章节「{section_title}」写作完成",
        )

        return {
            "status": "success",
            "section_id": outline["id"],
            "section_title": section_title,
            "section_content": section_content,
            "references_used": result.get("references_used", []),
            "word_count": len(section_content),
            "figure_paths": figure_paths,
            "message": f"章节「{section_title}」写作完成（{len(section_content)} 字"
                      + ("，含" + str(len(figure_paths)) + "张图表)" if figure_paths else "）"),
        }

    except Exception as e:
        logger.error(f"章节写作失败: {e}")
        return {
            "status": "error",
            "section_id": outline["id"],
            "section_title": section_title,
            "section_content": f"【章节写作失败: {e}】",
            "references_used": [],
            "word_count": 0,
            "message": f"章节「{section_title}」写作失败: {e}",
        }


def _section_job_payload(
    *,
    section_id: str,
    section_title: str,
    key_points: str,
    previous_sections_json: str,
    language: str,
    figure_descriptions: str,
    allow_degraded_writing: bool,
    user_acknowledgement: str,
) -> dict:
    return {
        "section_id": section_id,
        "section_title": section_title,
        "key_points": key_points,
        "previous_sections_json": previous_sections_json,
        "language": language,
        "figure_descriptions": figure_descriptions,
        "allow_degraded_writing": allow_degraded_writing,
        "user_acknowledgement": user_acknowledgement,
    }


async def _run_background_section_job(job_id: str, payload: dict) -> None:
    _write_jobs.mark_running(job_id)
    try:
        result = await _write_section_step_pipeline(
            **payload,
            mcp_context=WriteJobProgressContext(_write_jobs, job_id),
        )
    except Exception as exc:
        logger.exception("background section job failed: %s", job_id)
        _write_jobs.fail(job_id, str(exc))
        return
    _write_jobs.finish(job_id, result)


@mcp.tool()
async def write_section_step(
    section_id: str = "",
    section_title: str = "",
    key_points: str = "",
    previous_sections_json: str = "",
    language: str = "",
    figure_descriptions: str = "",
    allow_degraded_writing: bool = False,
    user_acknowledgement: str = "",
    run_mode: str = "background",
    advanced_mode: bool = False,
    mcp_context: Context | None = None,
) -> dict:
    """
    Start one section write step.

    Background mode returns a section job immediately so Nexent does not wait
    while drafting and discriminator review run locally. Poll
    get_write_section_job_status for section_content. Blocking mode is kept
    for short local verification. Default Nexent writing must use write_paper;
    set advanced_mode only for explicit section-by-section review workflows.
    """
    if not advanced_mode:
        return {
            "status": "redirect_to_write_paper",
            "redirect_tool": "write_paper",
            "message": (
                "默认论文生成不再使用逐章工具。请调用 write_paper；"
                "若尚无大纲，可直接把 document_path/topic/requirements 传给 write_paper。"
            ),
        }
    payload = _section_job_payload(
        section_id=section_id,
        section_title=section_title,
        key_points=key_points,
        previous_sections_json=previous_sections_json,
        language=language,
        figure_descriptions=figure_descriptions,
        allow_degraded_writing=allow_degraded_writing,
        user_acknowledgement=user_acknowledgement,
    )
    if run_mode.strip().lower() in {"blocking", "sync", "synchronous"}:
        return await _write_section_step_pipeline(**payload, mcp_context=mcp_context)

    policy_error = _set_material_policy_from_write_request(
        allow_degraded_writing=allow_degraded_writing,
        user_acknowledgement=user_acknowledgement,
    )
    if policy_error:
        return policy_error
    title = section_title or section_id or "untitled section"
    job = _write_jobs.create(title=title, section_count=1, job_kind="section")
    task = asyncio.create_task(_run_background_section_job(job["job_id"], payload))
    task.add_done_callback(_cleanup_write_job_task(job["job_id"]))
    _write_job_tasks[job["job_id"]] = task
    return {
        "status": "accepted",
        "job_id": job["job_id"],
        "job_kind": "section",
        "section_id": section_id,
        "section_title": title,
        "poll_tool": "get_write_section_job_status",
        "message": (
            "章节写作已在本地后台启动。请轮询 get_write_section_job_status，"
            "完成后该工具会返回 section_content。"
        ),
    }


# 服务端累积的章节（由 write_section_step 写入，供 render_final_paper 使用）
_SERVER_SECTIONS_CACHE = []

# ============================================================
# MCP Tool 6: 最终渲染（引用解析 + Word 输出）
# ============================================================
@mcp.tool()
def render_final_paper(title: str = "", sections_json: str = "") -> dict:
    """
    传入所有已完成的章节，执行双趟引用解析 + Word 渲染，返回下载链接。
    在逐章调用 write_section_step 后调用此工具。

    如果 sections_json 内容不完整，会自动使用服务端缓存的章节数据。
    """
    sections = []
    if sections_json:
        try:
            sections = json.loads(sections_json)
        except json.JSONDecodeError:
            return {"status": "error", "message": "sections_json JSON 解析失败"}

    # 如果 LLM 传的章节内容量明显偏少，用服务端缓存的数据
    total_chars = sum(len(s.get("content") or s.get("section_content") or "") for s in sections)
    cached_total = sum(len(s.get("content") or "") for s in _SERVER_SECTIONS_CACHE)
    if cached_total > 0 and cached_total > total_chars:
        logger.warning(f"sections_json 内容不完整（{total_chars}字），使用服务端缓存（{cached_total}字）")
        sections = list(_SERVER_SECTIONS_CACHE)
        if not title and sections:
            title = sections[0].get("title", "")

    # 字段名归一化：LLM 传的 key 可能是 section_content (来自 write_section_step 返回值)
    for s in sections:
        if "section_content" in s and "content" not in s:
            s["content"] = s["section_content"]

    # 诊断：记录各章节的实际内容长度
    for s in sections:
        c = s.get("content") or s.get("section_content") or ""
        logger.info(f"  render 接收: 「{s.get('title','')}」 content_len={len(c)} 字")
        if not c.strip():
            logger.warning(f"  WARN: 章节「{s.get('title','')}」内容为空！")

    ast = {
        "title": title or "未命名论文",
        "document_profile": _document_profile(),
        "front_matter": _front_matter(title or "未命名论文"),
        "sections": sections,
        "references": list(DYNAMIC_KNOWLEDGE_BASE.get("current_references", [])),
        "material_policy": _material_policy(),
        "material_sufficiency": DYNAMIC_KNOWLEDGE_BASE.get("material_sufficiency"),
        "entity_registry": {
            "images": [],
            "figures": [],
            "tables": [],
        },
        "asset_plan": list(DYNAMIC_KNOWLEDGE_BASE.get("asset_plan", [])),
        "materials": _active_materials(),
    }
    ast["evidence_registry"] = build_evidence_registry(
        ast["sections"],
        ast["references"],
        source_material=DYNAMIC_KNOWLEDGE_BASE.get("source_material", ""),
        result_evidence=DYNAMIC_KNOWLEDGE_BASE.get("result_evidence", ""),
        approved_assets=ast["entity_registry"]["images"],
        materials=ast["materials"],
    )

    citation_warnings = []
    try:
        _citation_resolver.resolve(ast)
    except Exception as e:
        logger.error(f"引用解析失败: {e}")
        citation_warnings.append(f"引用解析失败: {e}")

    audit = audit_manuscript(
        ast["sections"],
        evidence_registry=ast["evidence_registry"],
        entity_registry=ast["entity_registry"],
    )
    validation = validate_renderable_ast(ast)
    if not validation["ok"]:
        return {
            "status": "render_blocked",
            "errors": validation["errors"],
            "warnings": validation["warnings"],
            "message": "渲染被阻止：章节、引用或图片资产尚未闭环。",
        }
    warnings = _warning_groups(
        material_report=ast.get("material_sufficiency") or {},
        audit=audit,
        validation=validation,
        citation_warnings=citation_warnings,
    )
    generation_mode = "degraded" if warnings["evidence_warnings"] or not audit["ok"] else "standard"
    _complete_english_front_matter(ast)
    ast["generation_notes"] = _generation_notes(warnings, generation_mode)
    ast["generation_notes"]["layout_warnings"] = list(ast["front_matter"].get("layout_warnings") or [])

    try:
        ast["_finalize_word_fields"] = True
        output_path = render_docx(ast)
    except Exception as e:
        return {"status": "error", "message": f"Word 渲染失败: {e}"}
    ast["generation_notes"]["layout_warnings"] = list(ast["front_matter"].get("layout_warnings") or [])

    abs_path = os.path.abspath(output_path)
    host = _get_advertised_host()
    port = _get_server_port()
    dl_url = f"http://{host}:{port}/download?path={abs_path}"
    output_ok, output_error = _verify_docx_output(abs_path, len(sections))
    if not output_ok:
        return {"status": "error", "message": output_error, "output_path": abs_path}

    save_json(ast, "data/checkpoints/final_ast.json")

    logger.info(f"论文渲染完成: {abs_path}")
    return {
        "status": "completed",
        "output_path": abs_path,
        "download_url": dl_url,
        "section_count": len(sections),
        "reference_count": len(ast.get("references", [])),
        "validation": validation,
        "quality_audit": audit,
        **warnings,
        **_asset_delivery_metrics(ast),
        "document_profile": ast["document_profile"],
        "layout_warning_count": len(ast["generation_notes"].get("layout_warnings", [])),
        "generation_mode": generation_mode,
        "message": f"论文已完成！下载链接: {dl_url}",
    }


@mcp.tool()
def render_paper_tool(
    output_format: str = "docx",
    title: str = "",
    sections_json: str = "",
) -> dict:
    """Compatibility alias for older Nexent agents; render the final Word paper."""
    requested = (output_format or "docx").strip().lower()
    if requested not in {"docx", "word"}:
        return {
            "status": "error",
            "message": "Only Word DOCX output is supported. Use output_format='docx'.",
        }
    result = render_final_paper(title=title, sections_json=sections_json)
    result.setdefault("compatibility_alias", "render_paper_tool")
    return result


# ============================================================
# MCP Tool 7: 全自动写作论文（旧版/短论文用）
# ============================================================
async def _write_paper_pipeline(
    confirmed_outline=None,
    context: Context | None = None,
    job_id: str = "",
    task_snapshot: dict | None = None,
) -> dict:
    """
    全自动写作：逐章写作 + Discriminator 审查 + 双趟引用解析 + Word 渲染。
    自动发送进度通知保持连接活跃，不会超时。
    """
    ctx = context
    outline_data = confirmed_outline
    if isinstance(outline_data, str):
        try:
            outline_data = json.loads(outline_data)
        except json.JSONDecodeError:
            return {"status": "error", "message": "大纲 JSON 解析失败"}
    if isinstance(outline_data, list):
        sections = outline_data
        title = "未命名论文"
        language = _resolve_language()
    elif isinstance(outline_data, dict):
        if "outline" in outline_data and isinstance(outline_data["outline"], dict):
            outline_data = outline_data["outline"]
        outline_data = _localize_standard_outline(outline_data, outline_data.get("language", ""))
        title = outline_data.get("title", "未命名论文")
        sections = outline_data.get("sections", [])
        language = _resolve_language(outline_data.get("language", ""))
    else:
        # 从内存获取大纲
        kb_o = DYNAMIC_KNOWLEDGE_BASE.get("current_outline")
        if kb_o and kb_o.get("sections"):
            sections = kb_o["sections"]
            title = kb_o.get("title", "未命名论文")
            language = _resolve_language(kb_o.get("language", ""))
        else:
            return {"status": "error", "message": "请提供大纲或先调用 generate_outline"}

    has_task_snapshot = task_snapshot is not None
    task_state = dict(task_snapshot or {})
    task_materials = list(task_state.get("materials", _active_materials()))
    task_source_material = str(
        task_state.get("source_material", DYNAMIC_KNOWLEDGE_BASE.get("source_material", "")) or ""
    )
    task_result_evidence = str(
        task_state.get("result_evidence", DYNAMIC_KNOWLEDGE_BASE.get("result_evidence", "")) or ""
    )
    task_policy = dict(task_state.get("material_policy") or _material_policy())
    task_reference_snapshot = list(
        task_state.get("references", DYNAMIC_KNOWLEDGE_BASE.get("current_references", []))
    )
    task_asset_plan = list(
        task_state.get("asset_plan", DYNAMIC_KNOWLEDGE_BASE.get("asset_plan", []))
    ) or build_asset_plan(sections, topic=title)
    if has_task_snapshot:
        material_report = assess_material_sufficiency(
            topic=title,
            material_text=task_source_material,
            result_evidence=task_result_evidence,
            references=task_reference_snapshot,
            approved_assets=[],
            materials=task_materials,
            allow_degraded_writing=bool(task_policy.get("allow_degraded_writing")),
            degraded_reason=str(task_policy.get("user_acknowledgement", "") or ""),
        )
    else:
        material_report = _material_sufficiency(topic=title)
    DYNAMIC_KNOWLEDGE_BASE["material_sufficiency"] = material_report
    if _material_report_blocks_output(material_report):
        return {
            **material_report,
            "message": "缺少论文主题或可读材料，无法进入全自动正文写作。",
        }

    ast = {
        "title": title,
        "document_profile": str(task_state.get("document_profile") or _document_profile()),
        "front_matter": dict(task_state.get("front_matter") or _front_matter(title)),
        "asset_root": _task_asset_root(job_id),
        "sections": [],
        "references": [],
        "asset_plan": task_asset_plan,
        "material_policy": task_policy,
        "material_sufficiency": material_report,
        "language": language,
        "_job_id": job_id,
    }
    if has_task_snapshot and "references" in task_state:
        task_references = task_reference_snapshot
    else:
        try:
            task_references = _build_reference_pool(title, sections)
        except Exception as exc:
            logger.warning("任务文献池建立失败，继续输出并记录警告: %s", exc)
            task_references = []
    result_figures, result_tables, figure_warnings, result_asset_summary = _auto_generate_result_figures(
        title=title,
        asset_plan=ast["asset_plan"],
        job_id=job_id,
        language=language,
        result_evidence=task_result_evidence if has_task_snapshot else None,
        material_policy=task_policy if has_task_snapshot else None,
    )
    simulated_results = any(figure.get("simulated") for figure in result_figures)
    if simulated_results:
        ast["simulation_registry"] = {
            "enabled": True,
            "replacement_required": True,
            "summary": result_asset_summary,
            "items": [figure.get("asset_id") for figure in result_figures]
            + [table.get("asset_id") for table in result_tables],
        }
    review_enabled = _APP_CONFIG.get("discriminator", {}).get("enabled", True)
    max_review = _APP_CONFIG.get("discriminator", {}).get("max_review_attempts", 3)

    logger.info(f"开始写作: '{title}' ({len(sections)} 章)")
    await _notify_progress(
        ctx,
        progress=0,
        total=len(sections),
        message=f"论文写作任务已启动，共 {len(sections)} 个章节",
    )

    for idx, section in enumerate(sections):
        sec_title = section.get("title", f"章节{idx+1}")
        logger.info(f"[{idx+1}/{len(sections)}] {sec_title}")

        try:
            papers = task_references

            prev_ctx = get_previous_context(ast["sections"], max_count=3)
            section_context = build_section_context(
                sec_title,
                ast["sections"],
                prev_context=prev_ctx,
                figure_summary=result_asset_summary,
            )
            move_prompt = format_moves_prompt(sec_title) + _language_instruction(language)
            if simulated_results:
                move_prompt += (
                    "\n\n## 降级结果边界（必须遵守）\n"
                    "本任务没有真实结果数据。任何数据集、实验、指标、对比或结论均须明确称为"
                    "模拟占位内容或待替换内容，不得声称来自真实观测、真实实验或真实数据集。"
                    if language == "中文"
                    else "\n\n## Degraded Results Boundary (mandatory)\n"
                    "No empirical result dataset is available. Describe every dataset, experiment, metric, "
                    "comparison and conclusion as simulated placeholder content requiring replacement; do not "
                    "claim real observations or real experiments."
                )

            # 写作
            result = await _run_with_heartbeat(
                write_section_local,
                section,
                papers,
                ctx=ctx,
                progress=idx,
                total=len(sections),
                message=f"正在起草章节「{sec_title}」 ({idx + 1}/{len(sections)})",
                context=section_context,
                move_sequence=move_prompt,
            )

            # Discriminator 审查（有异常保护）
            if review_enabled:
                for ra in range(max_review):
                    try:
                        review = await _run_with_heartbeat(
                            review_section,
                            result.get("content", ""),
                            sec_title,
                            ctx=ctx,
                            progress=idx + 0.4,
                            total=len(sections),
                            message=f"正在审查章节「{sec_title}」第 {ra + 1} 轮",
                        )
                        if review.get("status") == "PASS":
                            break
                    except Exception as e:
                        logger.warning(f"Discriminator 调用失败 (第{ra+1}次): {e}")
                        break  # 审查失败也放行，不阻塞写作
                    if ra < max_review - 1:
                        result = await _run_with_heartbeat(
                            write_section_local,
                            section,
                            papers,
                            ctx=ctx,
                            progress=idx + 0.6,
                            total=len(sections),
                            message=f"正在重写章节「{sec_title}」第 {ra + 1} 轮",
                            context=section_context,
                            move_sequence=move_prompt,
                        )

            if simulated_results:
                result = _enforce_simulated_results_disclosure(
                    result,
                    language,
                    include_section_notice=_results_section(sec_title),
                )
            ast["sections"].append(result)

            # 累积文献到 AST，供后续引用解析使用
            for p in papers:
                t = (p.get("title") or "").strip().lower()
                if t and not any(r.get("_title_lower") == t for r in ast.get("references", [])):
                    p["_title_lower"] = t
                    ast.setdefault("references", []).append(p)

            logger.info(f"  章节完成 ({len(result.get('content',''))} 字)")

        except Exception as e:
            logger.error(f"章节 '{sec_title}' 处理失败: {e}")
            ast["sections"].append({
                "section_id": section.get("id", f"sec{idx}"),
                "title": sec_title,
                "content": f"【该章节因错误跳过: {e}】",
                "references_used": [],
            })

        await _notify_progress(
            ctx,
            progress=idx + 1,
            total=len(sections),
            message=f"章节「{sec_title}」写作完成 ({idx + 1}/{len(sections)})",
        )

    # 通知：引用解析中
    await _notify_progress(ctx, progress=len(sections), total=len(sections), message="正在解析参考文献")

    # 清理内部字段
    for ref in ast.get("references", []):
        ref.pop("_title_lower", None)

    citation_warnings = []
    try:
        await _run_with_heartbeat(
            _citation_resolver.resolve,
            ast,
            ctx=ctx,
            progress=len(sections),
            total=len(sections),
            message="正在解析参考文献",
        )
    except Exception as e:
        logger.error(f"引用解析失败: {e}")
        citation_warnings.append(f"引用解析失败: {e}")

    auto_assets, auto_asset_warnings = _auto_attach_open_image_assets(
        title=title,
        asset_plan=ast["asset_plan"],
        sections=sections,
        job_id=job_id,
        source_fingerprint=str(
            task_state.get("source_request", {}).get("request_fingerprint", "")
            or request_fingerprint(dict(task_state.get("source_request") or {}))
        ),
    )
    method_figures, method_figure_warnings = _auto_generate_method_figures(
        title=title,
        asset_plan=ast["asset_plan"],
        language=language,
        job_id=job_id,
    )
    material_assets, material_asset_warnings = _auto_bind_active_material_images(
        materials=task_materials,
        sections=sections,
        asset_plan=ast["asset_plan"],
        job_id=job_id,
    )
    ast["entity_registry"] = {
        "images": _current_job_image_assets(
            job_id,
            auto_assets + material_assets,
            material_ids=[material.get("material_id", "") for material in task_materials],
        ),
        "figures": method_figures + result_figures,
        "tables": result_tables,
    }
    ast["materials"] = task_materials
    ast["evidence_registry"] = build_evidence_registry(
        ast["sections"],
        ast["references"],
        source_material=task_source_material,
        result_evidence=task_result_evidence,
        approved_assets=ast["entity_registry"]["images"],
        materials=ast["materials"],
    )
    audit = audit_manuscript(
        ast["sections"],
        evidence_registry=ast["evidence_registry"],
        entity_registry=ast["entity_registry"],
    )
    validation = validate_renderable_ast(ast)
    if not validation["ok"]:
        return {
            "status": "render_blocked",
            "errors": validation["errors"],
            "warnings": validation["warnings"],
            "section_count": len(ast["sections"]),
            "message": "论文写作完成，但渲染前校验失败。",
        }
    warnings = _warning_groups(
        material_report=material_report,
        audit=audit,
        validation=validation,
        citation_warnings=citation_warnings,
        asset_warnings=auto_asset_warnings + material_asset_warnings + method_figure_warnings + figure_warnings,
    )
    generation_mode = (
        "degraded"
        if material_report.get("missing_materials") or warnings["evidence_warnings"] or not audit["ok"]
        else "standard"
    )
    _complete_english_front_matter(ast)
    ast["generation_notes"] = _generation_notes(warnings, generation_mode)
    ast["generation_notes"]["layout_warnings"] = list(ast["front_matter"].get("layout_warnings") or [])

    await _notify_progress(ctx, progress=len(sections), total=len(sections), message="正在渲染 Word 文档")
    try:
        ast["_finalize_word_fields"] = True
        output_path = await _run_with_heartbeat(
            render_docx,
            ast,
            ctx=ctx,
            progress=len(sections),
            total=len(sections),
            message="正在渲染 Word 文档",
        )
    except Exception as e:
        # 即使渲染失败，也返回已有数据
        logger.error(f"Word 渲染失败: {e}")
        return {
            "status": "partial",
            "download_url": "",
            "section_count": len(ast["sections"]),
            "message": f"论文写作完成，但 Word 渲染失败: {e}。AST 数据已保存到 data/checkpoints/final_ast.json",
        }
    ast["generation_notes"]["layout_warnings"] = list(ast["front_matter"].get("layout_warnings") or [])

    DYNAMIC_KNOWLEDGE_BASE["current_ast"] = ast
    save_json(DYNAMIC_KNOWLEDGE_BASE, "data/checkpoints/knowledge_base.json")
    save_json(ast, "data/checkpoints/final_ast.json")
    ast_path = "data/checkpoints/final_ast.json"
    if job_id:
        ast_path = f"data/checkpoints/jobs/{job_id}/final_ast.json"
        save_json(ast, ast_path)

    abs_path = os.path.abspath(output_path)
    output_ok, output_error = _verify_docx_output(abs_path, len(ast["sections"]))
    if not output_ok:
        return {
            "status": "error",
            "output_path": abs_path,
            "section_count": len(ast["sections"]),
            "message": output_error,
        }
    host = _get_advertised_host()
    port = _get_server_port()
    dl_url = f"http://{host}:{port}/download?path={abs_path}"

    logger.info(f"=" * 60)
    logger.info(f"论文写作完成！")
    logger.info(f"本地文件路径: {abs_path}")
    logger.info(f"下载链接: {dl_url}")
    logger.info(f"即将返回工具结果到 Nexent...")
    logger.info(f"=" * 60)

    result = {
        "status": "completed",
        "output_path": abs_path,
        "ast_path": ast_path,
        "download_url": dl_url,
        "section_count": len(ast["sections"]),
        "reference_count": len(ast["references"]),
        "validation": validation,
        "render_warnings": validation.get("warnings", []),
        **warnings,
        **_asset_delivery_metrics(ast),
        "document_profile": ast["document_profile"],
        "layout_warning_count": len(ast["generation_notes"].get("layout_warnings", [])),
        "generation_mode": generation_mode,
        "message": f"论文写作完成！\n本地文件路径: {abs_path}\n下载链接: {dl_url}",
    }
    logger.info(f"write_paper 返回成功, download_url={dl_url}")
    return result


def _outline_job_summary(confirmed_outline) -> tuple[str, int]:
    outline_data = confirmed_outline
    if isinstance(outline_data, str):
        try:
            outline_data = json.loads(outline_data)
        except json.JSONDecodeError:
            return "invalid outline", 0
    if isinstance(outline_data, dict) and isinstance(outline_data.get("outline"), dict):
        outline_data = outline_data["outline"]
    if isinstance(outline_data, dict):
        return str(outline_data.get("title") or "untitled"), len(outline_data.get("sections") or [])
    if isinstance(outline_data, list):
        return "untitled", len(outline_data)
    outline = DYNAMIC_KNOWLEDGE_BASE.get("current_outline") or {}
    return str(outline.get("title") or "untitled"), len(outline.get("sections") or [])


async def _run_background_write_job(job_id: str, confirmed_outline) -> None:
    _write_jobs.mark_running(job_id)
    snapshot = (_write_jobs.get(job_id) or {}).get("snapshot", {})
    outline_id = str(snapshot.get("outline_id", "") or "")
    try:
        source_request = snapshot.get("source_request", {})
        if not confirmed_outline and any(source_request.values()):
            _write_jobs.add_event(job_id, "正在依据上传材料或主题生成大纲")
            outline_result = await asyncio.to_thread(
                generate_outline,
                document_path=str(source_request.get("document_path", "")),
                requirements=str(source_request.get("requirements", "")),
                topic=str(source_request.get("topic", "")),
                language=str(source_request.get("language", "")),
            )
            if not outline_result.get("outline"):
                _write_jobs.finish(job_id, {
                    "status": "error",
                    "job_id": job_id,
                    "message": outline_result.get("message", "大纲生成失败"),
                    "outline_result": outline_result,
                })
                return
            confirmed_outline = outline_result["outline"]
        result = await _write_paper_pipeline(
            confirmed_outline=confirmed_outline,
            context=WriteJobProgressContext(_write_jobs, job_id),
            job_id=job_id,
            task_snapshot=snapshot,
        )
    except Exception as exc:
        logger.exception("background write job failed: %s", job_id)
        _write_jobs.fail(job_id, str(exc))
        if outline_id:
            _pending_outlines.mark(outline_id, "failed", job_id=job_id)
        return
    result.setdefault("job_id", job_id)
    _write_jobs.finish(job_id, result)
    if outline_id:
        final_status = "completed" if result.get("status") == "completed" else "failed"
        _pending_outlines.mark(outline_id, final_status, job_id=job_id)


def _cleanup_write_job_task(job_id: str):
    def cleanup(_task):
        _write_job_tasks.pop(job_id, None)
    return cleanup


@mcp.tool()
async def write_paper(
    confirmed_outline=None,
    outline_confirmed: bool = False,
    outline_id: str = "",
    document_path: str = "",
    topic: str = "",
    requirements: str = "",
    language: str = "",
    run_mode: str = "background",
    allow_degraded_writing: bool = False,
    user_acknowledgement: str = "",
    context: Context | None = None,
) -> dict:
    """
    Start full paper writing.

    Background mode runs as a durable local job while keeping this MCP call open
    until the final download_url is available. Nexent emits transport keepalives
    while the tool is active, avoiding both idle timeout and repeated agent
    polling steps. Use run_mode="submit" or "detached" only when the caller
    explicitly wants an immediate job_id for later recovery.
    """
    source_request = {
        "document_path": document_path.strip(),
        "topic": topic.strip(),
        "requirements": requirements.strip(),
        "language": _resolve_language(language),
    }
    pending = _pending_outlines.get(outline_id) if outline_id else None
    if outline_id and not pending:
        return {
            "status": "error",
            "message": f"unknown outline_id: {outline_id}",
        }
    if pending and pending.get("status") == "writing" and pending.get("job_id"):
        return get_write_paper_job_status(str(pending["job_id"]), wait_seconds=0)
    if not confirmed_outline and not pending and not any(
        source_request[key] for key in ("document_path", "topic", "requirements")
    ):
        return {
            "status": "error",
            "message": (
                "write_paper requires this task's confirmed_outline, uploaded document_path, "
                "or explicit topic/requirements; stale server outline fallback is disabled."
            ),
        }
    policy_error = _set_material_policy_from_write_request(
        allow_degraded_writing=allow_degraded_writing,
        user_acknowledgement=user_acknowledgement,
    )
    if policy_error:
        return policy_error
    if not pending and outline_confirmed and source_request["topic"]:
        pending = _pending_outlines.find_open_by_title(source_request["topic"])
        if pending:
            outline_id = str(pending["outline_id"])
            confirmed_outline = None
    if pending and pending.get("status") == "writing" and pending.get("job_id"):
        return get_write_paper_job_status(str(pending["job_id"]), wait_seconds=0)
    if pending and not confirmed_outline:
        if outline_confirmed:
            confirmed_outline = _activate_pending_outline(pending)
            source_request = dict(pending.get("source_request") or source_request)
        else:
            outline_result = dict(pending.get("outline_result") or {})
            return {
                **outline_result,
                "status": "outline_confirmation_required",
                "outline_id": pending["outline_id"],
                "confirmation_tool": "confirm_outline_and_start_writing",
                "message": (
                    "待确认大纲已就绪，正文尚未开始。"
                    "用户确认后调用 confirm_outline_and_start_writing(outline_id=...)。"
                ),
            }
    if not confirmed_outline and any(source_request.values()):
        existing = _pending_outlines.find_open(source_request)
        if outline_confirmed and existing:
            outline_id = existing["outline_id"]
            confirmed_outline = _activate_pending_outline(existing)
        else:
            outline_result = generate_outline(**source_request)
            if not outline_result.get("outline"):
                return outline_result
            return {
                **outline_result,
                "status": "outline_confirmation_required",
                "next_action": "confirm_outline_and_start_writing",
                "message": (
                    "已生成并锁定论文大纲，正文尚未开始。请展示大纲与材料缺口；"
                    "用户确认后调用 confirm_outline_and_start_writing(outline_id=...)。"
                ),
            }
    if confirmed_outline and not outline_confirmed:
        return {
            "status": "outline_confirmation_required",
            "outline": confirmed_outline,
            "next_action": "obtain_explicit_user_outline_confirmation",
            "message": (
                "已收到大纲，但没有用户确认标记，正文尚未开始。"
                "请先展示大纲并获得用户明确确认，再设置 outline_confirmed=true。"
            ),
        }
    if isinstance(confirmed_outline, dict):
        confirmed_outline = _localize_standard_outline(
            confirmed_outline,
            language or confirmed_outline.get("language", ""),
        )
    title, section_count = _outline_job_summary(confirmed_outline)
    if source_request["topic"]:
        title = source_request["topic"]
    elif source_request["requirements"]:
        title = source_request["requirements"][:120]
    elif source_request["document_path"]:
        title = _extract_filename(source_request["document_path"])
    material_report = _material_sufficiency(topic=title)
    DYNAMIC_KNOWLEDGE_BASE["material_sufficiency"] = material_report
    if _material_report_blocks_output(material_report) and not source_request["document_path"]:
        return {
            **material_report,
            "message": "后台写作任务未启动：缺少论文主题或可读材料。",
        }
    if not _material_report_allows_writing(material_report):
        return {
            **material_report,
            "status": "material_confirmation_required",
            "outline": confirmed_outline,
            "message": (
                "大纲已确认，但支撑成稿的资料不足，正文尚未开始。"
                "请向用户列出缺失材料并询问：补充材料，或明确接受降级写作风险。"
                "只有用户选择后者时，才可使用 allow_degraded_writing=true 与 user_acknowledgement 继续。"
            ),
        }
    if run_mode.strip().lower() in {"blocking", "sync", "synchronous"}:
        result = await _write_paper_pipeline(confirmed_outline=confirmed_outline, context=context)
        if outline_id:
            final_status = "completed" if result.get("status") == "completed" else "failed"
            _pending_outlines.mark(outline_id, final_status)
        return result
    job = _write_jobs.create(
        title=title,
        section_count=section_count,
        snapshot=_job_snapshot(confirmed_outline, source_request, outline_id),
    )
    if outline_id:
        _pending_outlines.mark(outline_id, "writing", job_id=job["job_id"])
    task = asyncio.create_task(_run_background_write_job(job["job_id"], confirmed_outline))
    task.add_done_callback(_cleanup_write_job_task(job["job_id"]))
    _write_job_tasks[job["job_id"]] = task
    accepted_result = {
        "status": "accepted",
        "job_id": job["job_id"],
        "title": title,
        "section_count": section_count,
        "material_sufficiency": material_report,
        "generation_mode": "degraded" if material_report.get("missing_materials") else "standard",
        "evidence_warnings": _warnings_from_material_report(material_report),
        "poll_tool": "get_write_paper_job_status",
        "message": (
            "论文写作已在本地后台启动。请调用一次 "
            "get_write_paper_job_status(job_id, wait_seconds=600) 等待最终 download_url；"
            "平台运行层将在等待期间保持连接活跃。"
        ),
    }
    if run_mode.strip().lower() in {"submit", "detached", "async_submit"}:
        return accepted_result
    try:
        await asyncio.shield(task)
    except asyncio.CancelledError:
        logger.warning("write_paper request disconnected while background job continues: %s", job["job_id"])
        return accepted_result
    return get_write_paper_job_status(job["job_id"], wait_seconds=0, include_details=False)


@mcp.tool()
async def confirm_outline_and_start_writing(
    outline_id: str = "",
    run_mode: str = "background",
    allow_degraded_writing: bool = False,
    user_acknowledgement: str = "",
    context: Context | None = None,
) -> dict:
    """
    Confirm the displayed outline and start one whole-paper writing job.

    Call this tool immediately after the user confirms a previously displayed
    outline. outline_id is optional: if it was not surfaced in the visible
    conversation, omit it and the service confirms the most recently waiting
    outline. Never call generate_outline again merely to recover an outline_id.
    """
    pending = _pending_outlines.get(outline_id) if outline_id else _pending_outlines.latest_waiting()
    if not pending:
        return {
            "status": "error",
            "message": "No waiting outline found. Generate an outline before confirming writing.",
        }
    if pending.get("status") == "writing" and pending.get("job_id"):
        return get_write_paper_job_status(str(pending["job_id"]), wait_seconds=0)
    source_request = dict(pending.get("source_request") or {})
    return await write_paper(
        outline_id=pending["outline_id"],
        outline_confirmed=True,
        document_path=str(source_request.get("document_path", "")),
        topic=str(source_request.get("topic", "")),
        requirements=str(source_request.get("requirements", "")),
        language=str(source_request.get("language", "")),
        run_mode=run_mode,
        allow_degraded_writing=allow_degraded_writing,
        user_acknowledgement=user_acknowledgement,
        context=context,
    )


@mcp.tool()
def get_write_paper_job_status(
    job_id: str,
    event_limit: int = 12,
    wait_seconds: int = 600,
    include_details: bool = False,
) -> dict:
    """Long-poll one paper job; return a compact final link unless details are requested."""
    wait_until = time.monotonic() + max(0, min(wait_seconds, 900))
    job = _write_jobs.get(job_id, event_limit=max(0, min(event_limit, 50)))
    while job and job.get("status") in {"accepted", "running"} and time.monotonic() < wait_until:
        time.sleep(1)
        job = _write_jobs.get(job_id, event_limit=max(0, min(event_limit, 50)))
    if not job:
        return {"status": "not_found", "job_id": job_id, "message": "unknown write job"}
    result = job.get("result") or {}
    response = {
        "status": job["status"],
        "job_id": job_id,
        "title": job.get("title"),
        "progress": job.get("progress"),
        "events": job.get("events", []),
        "result_status": job.get("result_status", ""),
        "download_url": result.get("download_url", ""),
        "output_path": result.get("output_path", ""),
        "section_count": result.get("section_count", job.get("section_count", 0)),
        "reference_count": result.get("reference_count", 0),
        "generation_mode": result.get("generation_mode", ""),
        "warning_counts": {
            "quality": len(result.get("quality_warnings", [])),
            "evidence": len(result.get("evidence_warnings", [])),
            "asset": len(result.get("asset_warnings", [])),
        },
        "quality_warnings": result.get("quality_warnings", [])[:3],
        "evidence_warnings": result.get("evidence_warnings", [])[:3],
        "asset_warnings": result.get("asset_warnings", [])[:3],
        "image_count": result.get("image_count", 0),
        "figure_count": result.get("figure_count", 0),
        "table_count": result.get("table_count", 0),
        "network_image_count": result.get("network_image_count", 0),
        "simulated_asset_count": result.get("simulated_asset_count", 0),
        "document_profile": result.get("document_profile", ""),
        "layout_warning_count": result.get("layout_warning_count", 0),
        "message": (
            "后台写作仍在进行中。"
            if job["status"] in {"accepted", "running"}
            else "后台写作已结束；如 download_url 非空，请直接返回下载链接。"
        ),
    }
    if include_details:
        response["result"] = result
        response["quality_warnings"] = result.get("quality_warnings", [])
        response["evidence_warnings"] = result.get("evidence_warnings", [])
        response["asset_warnings"] = result.get("asset_warnings", [])
    return response


@mcp.tool()
def get_write_section_job_status(job_id: str, event_limit: int = 12) -> dict:
    """Return one background section job and its section_content when finished."""
    job = _write_jobs.get(job_id, event_limit=max(0, min(event_limit, 50)))
    if not job or job.get("job_kind") != "section":
        return {"status": "not_found", "job_id": job_id, "message": "unknown section write job"}
    result = job.get("result") or {}
    return {
        "status": job["status"],
        "job_id": job_id,
        "job_kind": "section",
        "section_title": job.get("title"),
        "progress": job.get("progress"),
        "events": job.get("events", []),
        "result_status": job.get("result_status", ""),
        "result": result,
        "section_content": result.get("section_content", ""),
        "message": (
            "后台章节写作仍在进行中。"
            if job["status"] in {"accepted", "running"}
            else "后台章节写作已结束，请检查 result。"
        ),
    }


@mcp.tool()
def list_write_paper_jobs() -> dict:
    """List recent background paper and section jobs for reconnect and recovery."""
    return {"status": "success", "jobs": _write_jobs.summaries()}


# ============================================================
# MCP Tool 6: 渲染输出
# ============================================================
# ============================================================
# MCP Tool 7: Word 文档后处理
# ============================================================
@mcp.tool()
def edit_document(
    file_path: str = "",
    operation: str = "",
    target_text: str = "",
    new_text: str = "",
    paragraph_index: int = -1,
    heading_level: int = 0,
    font_name: str = "",
    font_size: int = 0,
    line_spacing: float = 0,
    alignment: str = "",
    comment_text: str = "",
) -> dict:
    """对 Word 文档进行后处理。支持 replace/set_heading/edit_paragraph/set_font/set_line_spacing/set_alignment/set_margins/add_page_numbers/add_comment/accept_revisions/reject_revisions"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not os.path.exists(file_path):
        return {"status": "error", "message": f"文件不存在: {file_path}"}
    try:
        doc = Document(file_path)
    except Exception as e:
        return {"status": "error", "message": f"无法打开文档: {e}"}

    if operation == "replace":
        count = 0
        for para in doc.paragraphs:
            if target_text in para.text:
                for run in para.runs:
                    if target_text in run.text:
                        run.text = run.text.replace(target_text, new_text)
                        count += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if target_text in para.text:
                            for run in para.runs:
                                if target_text in run.text:
                                    run.text = run.text.replace(target_text, new_text)
                                    count += 1
        doc.save(file_path)
        return {"status": "success", "message": f"替换完成，共修改 {count} 处"}

    elif operation == "set_heading":
        if heading_level < 1 or heading_level > 9:
            return {"status": "error", "message": "heading_level 范围 1-9"}
        sname = f"Heading {heading_level}"
        count = 0
        for para in doc.paragraphs:
            if target_text and target_text in para.text:
                para.style = doc.styles[sname]; count += 1
            elif not target_text and paragraph_index >= 0:
                if paragraph_index < len(doc.paragraphs):
                    doc.paragraphs[paragraph_index].style = doc.styles[sname]; count += 1; break
        doc.save(file_path)
        return {"status": "success" if count > 0 else "warning",
                "message": f"已设置 {count} 个段落为 {sname}" if count > 0 else "未找到匹配段落"}

    elif operation == "edit_paragraph":
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return {"status": "error", "message": f"段落索引 {paragraph_index} 超出"}
        para = doc.paragraphs[paragraph_index]
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        else:
            run_elem = OxmlElement("w:r")
            text_elem = OxmlElement("w:t")
            text_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            text_elem.text = new_text
            run_elem.append(text_elem)
            para._element.append(run_elem)
        doc.save(file_path)
        return {"status": "success", "message": f"段落 {paragraph_index} 已更新"}

    elif operation == "set_font":
        count = 0
        for para in doc.paragraphs:
            if target_text and target_text not in para.text:
                continue
            for run in para.runs:
                if font_name:
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
                if font_size > 0:
                    run.font.size = Pt(font_size)
                count += 1
            if not target_text:
                break
        doc.save(file_path)
        return {"status": "success", "message": f"已修改 {count} 处字体"}

    elif operation == "set_line_spacing":
        if line_spacing <= 0:
            return {"status": "error", "message": "line_spacing 必须大于 0"}
        for para in doc.paragraphs:
            para.paragraph_format.line_spacing = line_spacing
        doc.save(file_path)
        return {"status": "success", "message": f"行距已设为 {line_spacing}"}

    elif operation == "set_alignment":
        amap = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
        al = amap.get(alignment.lower()) if alignment else None
        if not al:
            return {"status": "error", "message": "alignment: left/center/right/justify"}
        count = 0
        for para in doc.paragraphs:
            if target_text and target_text in para.text:
                para.paragraph_format.alignment = al; count += 1
            elif not target_text:
                para.paragraph_format.alignment = al; count += 1
        doc.save(file_path)
        return {"status": "success", "message": f"已设置 {count} 个段落对齐"}

    elif operation == "set_margins":
        parts = new_text.replace("，", ",").split(",")
        if len(parts) != 4:
            return {"status": "error", "message": "格式: 上,右,下,左（厘米）"}
        try:
            margins = [Cm(float(p.strip())) for p in parts]
        except Exception:
            return {"status": "error", "message": "页边距值必须为数字"}
        sec = doc.sections[0]
        sec.top_margin, sec.right_margin, sec.bottom_margin, sec.left_margin = margins
        doc.save(file_path)
        return {"status": "success", "message": "页边距已设置"}

    elif operation == "add_page_numbers":
        for sec in doc.sections:
            footer = sec.footer
            footer.is_linked_to_previous = False
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
            run._element.append(f1)
            r2 = para.add_run()
            instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            r2._element.append(instr)
            r3 = para.add_run()
            f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
            r3._element.append(f2)
        doc.save(file_path)
        return {"status": "success", "message": "页码已添加"}

    elif operation == "add_comment":
        import datetime
        found = False
        for para in doc.paragraphs:
            if target_text in para.text:
                cid = str(len(list(doc.element.body.iterchildren())) + 1)
                comment = OxmlElement("w:comment")
                comment.set(qn("w:id"), cid)
                comment.set(qn("w:author"), "Paper Agent")
                comment.set(qn("w:date"), datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"))
                cr = OxmlElement("w:r")
                ct = OxmlElement("w:t")
                ct.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                ct.text = comment_text; cr.append(ct); comment.append(cr)
                doc.element.body.append(comment)
                for run in para.runs:
                    if target_text in run.text:
                        cr2 = OxmlElement("w:commentReference")
                        cr2.set(qn("w:id"), cid)
                        run._element.append(cr2); found = True; break
                if found:
                    break
        doc.save(file_path)
        if found:
            return {"status": "success", "message": "批注已添加"}
        return {"status": "warning", "message": "未找到匹配文本"}

    elif operation in ("accept_revisions", "reject_revisions"):
        body = doc.element.body
        count = 0
        accept = operation == "accept_revisions"
        for tag in ("w:ins", "w:del"):
            for el in list(body.iter(qn(tag))):
                if tag == "w:ins" and accept:
                    el.tag = qn("w:r"); count += 1
                elif tag == "w:ins" and not accept:
                    el.getparent().remove(el); count += 1
                elif tag == "w:del" and accept:
                    el.getparent().remove(el); count += 1
                elif tag == "w:del" and not accept:
                    el.tag = qn("w:r"); count += 1
        doc.save(file_path)
        return {"status": "success", "message": f"修订处理完成，影响 {count} 处"}

    return {"status": "error", "message": f"不支持的 operation: {operation}"}


# ============================================================
# MCP Tool 8: Word 表格操作
# ============================================================
@mcp.tool()
def table_operation(
    file_path: str = "",
    operation: str = "",
    rows: int = 0,
    cols: int = 0,
    row_index: int = -1,
    col_index: int = -1,
    cell_text: str = "",
    table_index: int = 0,
) -> dict:
    """对 Word 文档中的表格进行操作。支持 insert_table/edit_cell/merge_cells/delete_table"""
    from docx import Document
    from docx.oxml.ns import qn

    if not os.path.exists(file_path):
        return {"status": "error", "message": f"文件不存在: {file_path}"}
    try:
        doc = Document(file_path)
    except Exception as e:
        return {"status": "error", "message": f"无法打开文档: {e}"}

    if operation == "insert_table":
        if rows <= 0 or cols <= 0:
            return {"status": "error", "message": "rows 和 cols 必须大于 0"}
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        doc.save(file_path)
        return {"status": "success", "message": f"已插入 {rows}x{cols} 表格"}

    elif operation == "edit_cell":
        tables = doc.tables
        if table_index < 0 or table_index >= len(tables):
            return {"status": "error", "message": f"表格索引 {table_index} 超出"}
        t = tables[table_index]
        if row_index < 0 or row_index >= len(t.rows):
            return {"status": "error", "message": f"行索引 {row_index} 超出"}
        if col_index < 0 or col_index >= len(t.rows[row_index].cells):
            return {"status": "error", "message": f"列索引 {col_index} 超出"}
        cell = t.rows[row_index].cells[col_index]
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ""
            p.text = ""
        cell.text = cell_text
        doc.save(file_path)
        return {"status": "success", "message": f"表格 {table_index} ({row_index},{col_index}) 已更新"}

    elif operation == "merge_cells":
        tables = doc.tables
        if table_index < 0 or table_index >= len(tables):
            return {"status": "error", "message": "表格索引超出"}
        t = tables[table_index]
        if row_index < 0 or row_index >= len(t.rows) - 1:
            return {"status": "error", "message": "合并需指定起始行（不可为最后一行）"}
        if col_index < 0 or col_index >= len(t.rows[row_index].cells):
            return {"status": "error", "message": "列索引超出"}
        start = t.rows[row_index].cells[col_index]
        end = t.rows[row_index + 1].cells[col_index]
        start.merge(end)
        doc.save(file_path)
        return {"status": "success", "message": f"表格 {table_index} 第{col_index}列 {row_index}~{row_index+1} 行已合并"}

    elif operation == "delete_table":
        tables = doc.tables
        if table_index < 0 or table_index >= len(tables):
            return {"status": "error", "message": "表格索引超出"}
        t = tables[table_index]
        t._element.getparent().remove(t._element)
        doc.save(file_path)
        return {"status": "success", "message": f"已删除表格 {table_index}"}

    return {"status": "error", "message": f"不支持的 operation: {operation}"}


# ============================================================
# 入口
# ============================================================
if __name__ == "__main__":
    listen_host = _get_listen_host()
    port = _get_server_port()
    advertised = _get_advertised_host()

    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    result = sock.connect_ex(("127.0.0.1", port))
    sock.close()
    if result == 0:
        print(f"错误: 端口 {port} 已被占用。运行 taskkill /F /IM python.exe 杀死旧进程。")
        sys.exit(1)

    from starlette.applications import Starlette
    from starlette.routing import Mount, Route
    from starlette.responses import FileResponse, JSONResponse
    import uvicorn

    async def _download_file(request):
        fp = request.query_params.get("path", "")
        if not fp or not os.path.exists(fp):
            return JSONResponse({"error": "文件不存在"}, status_code=404)
        return FileResponse(fp, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            filename=os.path.basename(fp))

    mcp_app = mcp.http_app(transport="sse")
    app = Starlette(routes=[
        Route("/download", endpoint=_download_file),
        Mount("/", app=mcp_app),
    ])
    logger.info(f"启动 MCP 服务: {listen_host}:{port}")
    logger.info(f"Nexent 注册地址: http://{advertised}:{port}  (advertised_host 为空时填实际 IP)")
    logger.info(f"文件下载: http://{advertised}:{port}/download?path=...")
    uvicorn.run(app, host=listen_host, port=port, log_level="info")
