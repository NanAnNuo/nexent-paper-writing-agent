import base64
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document

from execution.docx_renderer import CQU_PROFESSIONAL_MASTER, _resolved_front_matter, render_docx
from execution.simulated_results import build_simulated_result_assets
from orchestrator.pending_outlines import PendingOutlineStore
from main import (
    _asset_delivery_metrics,
    _auto_attach_open_image_assets,
    _auto_bind_active_material_images,
    _auto_generate_result_figures,
    _auto_generate_method_figures,
    _complete_english_front_matter,
    _current_job_image_assets,
    _enforce_simulated_results_disclosure,
    _literature_queries,
    _verify_docx_output,
    _write_paper_pipeline,
    confirm_outline_and_start_writing,
    generate_outline,
    get_write_paper_job_status,
    write_paper,
    write_section_step,
)


class OutputFirstRenderingTests(unittest.TestCase):
    def test_graph_air_quality_queries_do_not_use_molecular_pool(self):
        queries = _literature_queries("基于图神经网络的城市空气质量预测方法研究")

        self.assertIn("spatiotemporal graph neural network air quality forecasting", queries)
        self.assertIn("air quality prediction machine learning PM2.5", queries)
        self.assertNotIn("graph neural network molecular property prediction", queries)

    def test_degraded_word_contains_generation_notes_and_passes_output_check(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "draft.docx"
            render_docx(
                {
                    "title": "Output-first Draft",
                    "sections": [
                        {"title": "Introduction", "content": "## Context\n\nParagraph one.\n\nParagraph two."},
                        {"title": "Conclusion", "content": "The draft remains reviewable."},
                    ],
                    "references": [],
                    "generation_notes": {
                        "generation_mode": "degraded",
                        "evidence_warnings": ["材料缺口: missing_result_dataset"],
                    },
                },
                str(output),
            )

            ok, error = _verify_docx_output(str(output), expected_sections=2)
            text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)

            self.assertTrue(ok, error)
            self.assertIn("生成说明与待核验项", text)
            self.assertIn("missing_result_dataset", text)

    def test_legacy_sec_number_image_binding_renders_into_canonical_section(self):
        png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        with tempfile.TemporaryDirectory() as tmp:
            image = Path(tmp) / "device.png"
            output = Path(tmp) / "with-image.docx"
            image.write_bytes(png)
            render_docx(
                {
                    "title": "含图论文",
                    "language": "中文",
                    "sections": [
                        {"section_id": "abstract", "title": "摘要", "content": "文本。"},
                        {"section_id": "introduction", "title": "引言", "content": "文本。"},
                        {"section_id": "methods", "title": "方法", "content": "文本。"},
                    ],
                    "references": [],
                    "entity_registry": {
                        "images": [{
                            "approved": True,
                            "local_path": str(image),
                            "section_id": "sec3",
                            "caption": "系统实物图",
                        }]
                    },
                },
                str(output),
            )
            document = Document(output)
            self.assertEqual(len(document.inline_shapes), 1)
            self.assertIn("图: 系统实物图", "\n".join(p.text for p in document.paragraphs))

    def test_english_generation_notes_and_reference_heading_are_localized(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "english.docx"
            render_docx(
                {
                    "title": "English Paper",
                    "language": "English",
                    "sections": [{"title": "Conclusion", "content": "Text."}],
                    "references": [],
                    "generation_notes": {
                        "generation_mode": "degraded",
                        "evidence_warnings": ["missing dataset"],
                    },
                },
                str(output),
            )
            text = "\n".join(p.text for p in Document(output).paragraphs)
            self.assertIn("References", text)
            self.assertIn("Generation Notes and Items Requiring Verification", text)
            self.assertNotIn("生成说明与待核验项", text)

    def test_missing_inline_figure_does_not_leave_visible_placeholder(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "missing-figure.docx"
            render_docx(
                {
                    "title": "Figure fallback",
                    "sections": [{
                        "section_id": "results",
                        "title": "Results",
                        "content": "Result narrative.\n\n[Figure: unavailable.png]",
                    }],
                    "references": [],
                },
                str(output),
            )
            text = "\n".join(p.text for p in Document(output).paragraphs)
            self.assertNotIn("[图片:", text)
            self.assertNotIn("unavailable.png", text)

    def test_required_methods_visual_gets_truthful_concept_diagram(self):
        with tempfile.TemporaryDirectory() as tmp:
            figures, warnings = _auto_generate_method_figures(
                title="基于脑电运动想象的机械臂控制系统",
                asset_plan=[{
                    "section_id": "methods",
                    "purpose": "system_or_device_context",
                    "required": True,
                }],
                language="中文",
                job_id="visual-test",
                output_dir=tmp,
            )
            output = Path(tmp) / "concept.docx"
            render_docx(
                {
                    "title": "概念图稿件",
                    "language": "中文",
                    "sections": [{"section_id": "methods", "title": "方法", "content": "正文。"}],
                    "references": [],
                    "entity_registry": {"figures": figures, "images": []},
                },
                str(output),
            )
            document = Document(output)
            self.assertEqual(warnings, [])
            self.assertTrue(Path(figures[0]["local_path"]).exists())
            self.assertEqual(len(document.inline_shapes), 2)
            self.assertIn("概念图，不表示实验结果", "\n".join(p.text for p in document.paragraphs))

    def test_air_quality_methods_visual_does_not_reuse_bci_diagram(self):
        with tempfile.TemporaryDirectory() as tmp:
            figures, warnings = _auto_generate_method_figures(
                title="基于图神经网络的城市空气质量预测方法研究",
                asset_plan=[{
                    "section_id": "methods",
                    "purpose": "system_or_device_context",
                    "required": True,
                }],
                language="中文",
                job_id="air-workflow",
                output_dir=tmp,
            )

            self.assertEqual(warnings, [])
            self.assertIn("空气质量", figures[0]["caption"])
            self.assertNotIn("脑电", figures[0]["caption"])
            self.assertNotIn("机械臂", figures[0]["caption"])
            self.assertEqual(len(figures), 2)

    def test_cqu_professional_master_profile_renders_front_matter_and_numbered_bilingual_captions(self):
        png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        with tempfile.TemporaryDirectory() as tmp:
            image = Path(tmp) / "workflow.png"
            image.write_bytes(png)
            output = Path(tmp) / "cqu-master.docx"
            render_docx(
                {
                    "title": "基于深度学习的测试论文",
                    "language": "中文",
                    "document_profile": CQU_PROFESSIONAL_MASTER,
                    "front_matter": {"title_en": "Test Thesis Based on Deep Learning"},
                    "sections": [
                        {"section_id": "abstract", "title": "摘要", "content": "摘要正文。"},
                        {"section_id": "methods", "title": "方法", "content": "方法正文。"},
                    ],
                    "references": [{"title": "Reference", "authors": ["A"], "year": 2024, "venue": "Journal"}],
                    "entity_registry": {
                        "figures": [{
                            "approved": True,
                            "local_path": str(image),
                            "section_id": "methods",
                            "caption": "研究流程示意图",
                            "english_caption": "Research workflow",
                        }],
                        "tables": [{
                            "section_id": "methods",
                            "caption": "指标汇总",
                            "english_caption": "Metric summary",
                            "headers": ["指标", "值"],
                            "rows": [["A", "1"]],
                        }],
                    },
                },
                str(output),
            )
            document = Document(output)
            text = "\n".join(p.text for p in document.paragraphs)
            header_text = "\n".join(p.text for section in document.sections for p in section.header.paragraphs)
            self.assertIn("重庆大学硕士学位论文", text)
            self.assertIn("专业学位", text)
            self.assertIn("A Thesis Submitted to Chongqing University", text)
            self.assertIn("Abstract", text)
            self.assertIn("目录", text)
            self.assertNotIn("请在 Word 中更新目录域", text)
            self.assertIn("图1.1 研究流程示意图", text)
            self.assertIn("Fig.1.1 Research workflow", text)
            self.assertIn("表1.1 指标汇总", text)
            self.assertIn("Table 1.1 Metric summary", text)
            self.assertIn("致谢", text)
            self.assertIn("重庆大学硕士学位论文", header_text)
            toc_title = next(p for p in document.paragraphs if p.text == "目录")
            self.assertNotEqual(toc_title.style.name, "Heading 1")
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                settings_xml = archive.read("word/settings.xml").decode("utf-8")
                header_xml = "\n".join(
                    archive.read(name).decode("utf-8")
                    for name in archive.namelist()
                    if name.startswith("word/header")
                )
            self.assertIn("TOC", document_xml)
            self.assertIn(r"\h", document_xml)
            self.assertNotIn("updateFields", settings_xml)
            self.assertNotIn("STYLEREF", header_xml)
            self.assertIn("upperRoman", document_xml)

    def test_cqu_toc_finalization_is_opt_in_and_does_not_persist_internal_flag(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "precomputed-toc.docx"
            ast = {
                "title": "目录预生成",
                "language": "中文",
                "document_profile": CQU_PROFESSIONAL_MASTER,
                "_finalize_word_fields": True,
                "sections": [{"section_id": "introduction", "title": "引言", "content": "正文。"}],
                "references": [],
            }
            with patch("execution.docx_renderer._finalize_word_toc", return_value=[]) as finalize:
                render_docx(ast, str(output))

            finalize.assert_called_once_with(str(output))
            self.assertNotIn("_finalize_word_fields", ast)

    def test_cqu_markdown_tables_are_real_tables_and_internal_headings_remain_subordinate(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "structured.docx"
            render_docx(
                {
                    "title": "结构化成稿",
                    "language": "中文",
                    "document_profile": CQU_PROFESSIONAL_MASTER,
                    "sections": [{
                        "section_id": "results",
                        "title": "结果",
                        "content": (
                            "## 消融实验\n\n"
                            "表1 消融实验结果\n"
                            "| 方法 | mAP |\n"
                            "|------|-----|\n"
                            "| Baseline | 78.3 |\n"
                            "| Proposed | 84.7 |"
                        ),
                    }],
                    "references": [],
                },
                str(output),
            )
            document = Document(output)
            headings = [(p.style.name, p.text) for p in document.paragraphs if p.style.name.startswith("Heading")]
            text = "\n".join(p.text for p in document.paragraphs)
            self.assertEqual(len(document.tables), 1)
            self.assertNotIn("| 方法 |", text)
            self.assertIn(("Heading 2", "消融实验"), headings)

    def test_cqu_section_does_not_repeat_its_own_leading_markdown_heading(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "no-duplicate-heading.docx"
            render_docx(
                {
                    "title": "去重目录结构",
                    "language": "中文",
                    "document_profile": CQU_PROFESSIONAL_MASTER,
                    "sections": [{
                        "section_id": "introduction",
                        "title": "引言",
                        "content": (
                            "## 说明\n\n前置说明。\n\n"
                            "## 引言\n\n正文内容。\n\n### 研究背景\n\n背景内容。"
                        ),
                    }],
                    "references": [],
                },
                str(output),
            )
            headings = [p.text for p in Document(output).paragraphs if p.style.name.startswith("Heading")]
            self.assertEqual(headings.count("1 引言"), 1)
            self.assertNotIn("引言", headings)
            self.assertIn("说明", headings)
            self.assertIn("研究背景", headings)

    def test_cqu_visual_paragraphs_are_not_clipped_by_fixed_body_line_spacing_or_reused_by_position(self):
        png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        with tempfile.TemporaryDirectory() as tmp:
            image = Path(tmp) / "diagram.png"
            image.write_bytes(png)
            output = Path(tmp) / "visible-image.docx"
            render_docx(
                {
                    "title": "图像测试",
                    "language": "中文",
                    "document_profile": CQU_PROFESSIONAL_MASTER,
                    "sections": [
                        {"section_id": "sec2", "title": "引言", "content": "正文。"},
                        {"section_id": "sec3", "title": "相关研究", "content": "正文。"},
                    ],
                    "references": [],
                    "entity_registry": {"figures": [{
                        "approved": True,
                        "local_path": str(image),
                        "section_id": "sec2",
                        "caption": "研究场景图",
                    }]},
                },
                str(output),
            )
            document = Document(output)
            picture_paragraphs = [p for p in document.paragraphs if "w:drawing" in p._p.xml]
            self.assertEqual(len(document.inline_shapes), 1)
            self.assertEqual(len(picture_paragraphs), 1)
            self.assertEqual(picture_paragraphs[0].paragraph_format.line_spacing, 1.0)

    def test_complete_english_front_matter_replaces_placeholder_with_generated_text(self):
        ast = {
            "title": "中文标题",
            "front_matter": {
                "title_en": "English Title To Be Completed",
                "abstract_en": "English abstract to be completed after the Chinese manuscript is finalized.",
                "layout_warnings": ["英文题目、英文摘要及封面身份信息为待填写占位内容，正式提交前需补全。"],
            },
            "sections": [{"title": "摘要", "content": "本文研究X光安检图像违禁品检测算法，并说明结果为模拟占位。"}],
        }
        response = {
            "title_en": "Prohibited Item Detection in X-Ray Security Images",
            "abstract_en": "This study investigates prohibited item detection in X-ray security images and reports only explicitly labelled simulated placeholder outcomes pending validation with real experimental data.",
            "keywords_en": "X-ray security; object detection; simulated data",
        }
        with patch("main.get_llm_client") as client:
            client.return_value.call.return_value = __import__("json").dumps(response)
            _complete_english_front_matter(ast)
        self.assertEqual(ast["front_matter"]["title_en"], response["title_en"])
        self.assertEqual(ast["front_matter"]["abstract_en"], response["abstract_en"])
        self.assertNotIn("英文摘要", ast["front_matter"]["layout_warnings"][0])

    def test_complete_english_front_matter_retries_malformed_json_once(self):
        ast = {
            "title": "中文标题",
            "front_matter": {
                "abstract_en": "English abstract to be completed after the Chinese manuscript is finalized.",
            },
            "sections": [{"title": "摘要", "content": "本文研究脑电运动想象控制机械臂，并明确说明模拟数据边界。"}],
        }
        response = {
            "title_en": "Motor Imagery EEG Control for a Robotic Arm",
            "abstract_en": "This study examines robotic-arm control using motor imagery electroencephalography and preserves the explicit limitation that simulated evidence must be replaced before formal reporting.",
            "keywords_en": "EEG; motor imagery; robotic arm",
        }
        with patch("main.get_llm_client") as client:
            client.return_value.call.side_effect = ["{malformed json", __import__("json").dumps(response)]
            _complete_english_front_matter(ast)

        self.assertEqual(client.return_value.call.call_count, 2)
        self.assertEqual(ast["front_matter"]["abstract_en"], response["abstract_en"])

    def test_cqu_front_matter_omits_reference_tail_from_abstract_and_normalizes_english_keywords(self):
        front, _body = _resolved_front_matter(
            {},
            [{"title": "摘要", "content": "## 摘要\n\n摘要正文。\n\n参考文献\n[1] 不应进入摘要页。"}],
            "测试论文",
        )
        self.assertEqual(front["abstract_cn"], "摘要正文。")
        self.assertNotIn("参考文献", front["abstract_cn"])
        ast = {
            "title": "中文标题",
            "front_matter": {"abstract_en": "English abstract to be completed after the Chinese manuscript is finalized."},
            "sections": [{"title": "摘要", "content": "摘要正文。\n\n参考文献\n[1] 不应进入翻译。"}],
        }
        response = {
            "title_en": "English Title",
            "abstract_en": "This abstract is intentionally long enough for validation and contains a faithful summary without reference entries or unsupported evidence in the front matter.",
            "keywords_en": ["X-ray security", "object detection", "simulated data"],
        }
        with patch("main.get_llm_client") as client:
            client.return_value.call.return_value = __import__("json").dumps(response)
            _complete_english_front_matter(ast)
            prompt = client.return_value.call.call_args.args[0]
        self.assertNotIn("不应进入翻译", prompt)
        self.assertEqual(ast["front_matter"]["keywords_en"], "X-ray security; object detection; simulated data")

    def test_simulated_result_assets_are_labelled_and_rendered_with_table(self):
        with tempfile.TemporaryDirectory() as tmp:
            figures, tables, summary = build_simulated_result_assets(
                "降级稿",
                "results",
                job_id="sim-test",
                output_dir=tmp,
                language="中文",
            )
            output = Path(tmp) / "simulated.docx"
            render_docx(
                {
                    "title": "降级稿",
                    "language": "中文",
                    "sections": [{"section_id": "results", "title": "结果", "content": "## 模拟数据说明\n\n正文。"}],
                    "references": [],
                    "entity_registry": {"figures": figures, "images": [], "tables": tables},
                },
                str(output),
            )
            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertEqual(len(document.inline_shapes), 2)
            self.assertEqual(len(document.tables), 1)
            self.assertTrue(all(figure["simulated"] for figure in figures))
            self.assertIn("模拟结果数据", summary)
            self.assertIn("模拟数据表", text)
            self.assertIn("必须替换", text)

    def test_air_quality_simulated_assets_use_topic_specific_metrics(self):
        with tempfile.TemporaryDirectory() as tmp:
            figures, tables, summary = build_simulated_result_assets(
                "基于图神经网络的城市空气质量预测方法研究",
                "results",
                job_id="air-quality",
                output_dir=tmp,
                language="中文",
            )

            self.assertIn("air-quality", figures[0]["local_path"])
            self.assertEqual(figures[0]["job_id"], "air-quality")
        self.assertIn("PM2.5", figures[0]["caption"])
        self.assertTrue(any("RMSE" in header for header in tables[0]["headers"]))
        self.assertTrue(any("R²" in header for header in tables[0]["headers"]))
        self.assertIn("ST-GNN", summary)
        self.assertNotIn("在线成功率", str(tables))
        self.assertNotIn("feedback lock", summary)

    def test_simulated_section_removes_real_data_claims(self):
        normalized = _enforce_simulated_results_disclosure({
            "content": (
                "本文实验采用真实空气质量监测数据集，在真实观测数据上进行了迁移测试。"
                "该数据集为非模拟占位数据，图中比较预测值与真实值。"
            )
        }, "中文")

        self.assertIn("模拟数据说明", normalized["content"])
        self.assertIn("模拟空气质量占位数据集", normalized["content"])
        self.assertIn("属于模拟占位数据", normalized["content"])
        self.assertIn("预测值与模拟目标值", normalized["content"])
        self.assertNotIn("真实空气质量监测数据集", normalized["content"])
        self.assertNotIn("真实观测数据", normalized["content"])
        self.assertNotIn("真实值", normalized["content"])
        self.assertNotIn("非模拟占位数据", normalized["content"])

    @patch("main._result_evidence", return_value="")
    @patch("main._material_policy", return_value={"allow_degraded_writing": True})
    def test_degraded_result_figure_path_produces_simulation_assets(self, _policy, _evidence):
        figures, tables, warnings, summary = _auto_generate_result_figures(
            title="降级结果稿",
            asset_plan=[{"asset_type": "data_figure", "section_id": "results", "purpose": "performance"}],
            job_id="degraded-result-test",
            language="中文",
        )
        self.assertEqual(len(figures), 2)
        self.assertEqual(len(tables), 1)
        self.assertTrue(figures[0]["simulated"])
        self.assertIn("模拟", warnings[0])
        self.assertIn("非真实实验观测", summary)

    @patch("main._image_assets.approved_assets")
    def test_job_assets_do_not_inherit_previous_paper_images(self, approved_assets):
        with tempfile.TemporaryDirectory() as tmp:
            current_file = Path(tmp) / "new.png"
            current_file.write_bytes(b"asset")
            approved_assets.return_value = [
                {"asset_id": "old", "approved": True, "material_id": "mat-old"},
                {"asset_id": "new", "approved": True, "job_id": "write-new", "local_path": str(current_file)},
            ]
            selected = _current_job_image_assets("write-new")
            self.assertEqual([asset["asset_id"] for asset in selected], ["new"])

    def test_active_uploaded_and_embedded_images_are_copied_into_current_job_assets(self):
        png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            result_image = root / "result.png"
            embedded_image = root / "embedded.png"
            result_image.write_bytes(png)
            embedded_image.write_bytes(png)
            from orchestrator.asset_manager import ImageAssetStore
            store = ImageAssetStore(root / "assets.json")
            with patch("main._image_assets", store), patch("main._task_asset_root", return_value=str(root / "write-job")):
                assets, warnings = _auto_bind_active_material_images(
                    materials=[
                        {
                            "material_id": "standalone",
                            "parse_status": "parsed",
                            "material_role": "image_asset",
                            "local_path": str(result_image),
                            "source_path": str(result_image),
                            "metadata": {"filename": "accuracy.png", "source_kind": "uploaded_image"},
                        },
                        {
                            "material_id": "embedded",
                            "parse_status": "parsed",
                            "material_role": "image_asset",
                            "local_path": str(embedded_image),
                            "source_path": str(embedded_image),
                            "metadata": {"filename": "device.png", "source_kind": "docx_embedded_image"},
                        },
                    ],
                    sections=[{"id": "sec2"}, {"id": "sec3"}, {"id": "sec4"}],
                    asset_plan=[
                        {"section_id": "sec3", "purpose": "system_or_device_context"},
                        {"section_id": "sec4", "asset_type": "data_figure"},
                    ],
                    job_id="write-job",
                )

            self.assertEqual([asset["section_id"] for asset in assets], ["sec4", "sec3"])
            self.assertTrue(all("write-job" in asset["local_path"] for asset in assets))
            self.assertEqual(len(warnings), 1)
            self.assertIn("题注与章节归属需核验", warnings[0])

    @patch("main._image_assets.approved_assets")
    def test_failed_current_job_image_is_not_rendered_or_counted(self, approved_assets):
        approved_assets.return_value = [
            {
                "asset_id": "failed",
                "approved": True,
                "job_id": "write-new",
                "source": "wikimedia_commons",
                "download_error": "429",
            }
        ]
        selected = _current_job_image_assets("write-new")
        metrics = _asset_delivery_metrics({"entity_registry": {"images": approved_assets.return_value}})
        self.assertEqual(selected, [])
        self.assertEqual(metrics["network_image_count"], 0)

    @patch("main.search_planned_commons_candidates")
    def test_open_image_download_tries_second_provider_after_commons_failure(self, search):
        search.return_value = [
            {
                "section_id": "intro",
                "purpose": "application_context",
                "source": "wikimedia_commons",
                "image_url": "https://commons.test/image.jpg",
                "query": "xray",
            },
            {
                "section_id": "intro",
                "purpose": "application_context",
                "source": "openverse",
                "image_url": "https://openverse.test/image.jpg",
                "query": "xray",
            },
        ]
        with tempfile.TemporaryDirectory() as tmp:
            from orchestrator.asset_manager import ImageAssetStore

            store = ImageAssetStore(Path(tmp) / "assets.json")

            def download(_output_dir, *, asset_ids):
                selected = next(asset for asset in store._assets if asset["asset_id"] in asset_ids)
                if selected["source"] == "wikimedia_commons":
                    selected["download_error"] = "429"
                else:
                    target = Path(tmp) / "openverse.jpg"
                    target.write_bytes(b"asset")
                    selected["local_path"] = str(target)
                store._save()
                return store.approved_assets()

            with patch("main._image_assets", store), patch.object(store, "download_approved", side_effect=download):
                assets, warnings = _auto_attach_open_image_assets(
                    title="X光安检",
                    asset_plan=[{"asset_type": "searched_image", "section_id": "intro"}],
                    sections=[],
                    job_id="write-new",
                )
        self.assertEqual([asset["source"] for asset in assets], ["openverse"])
        self.assertEqual(warnings, [])


class OutputFirstPipelineTests(unittest.IsolatedAsyncioTestCase):
    def setUp(self):
        self.save_json_patch = patch("main.save_json")
        self.save_json_patch.start()
        self.addCleanup(self.save_json_patch.stop)

    async def test_material_gaps_become_degraded_completed_word(self):
        def fake_write_section(section, papers, **_kwargs):
            return {
                "section_id": section["id"],
                "title": section["title"],
                "content": "## Evidence gap\n\nDraft text with a result claim that needs verification.",
                "references_used": [],
            }

        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "pipeline.docx"
            with (
                patch("main._material_sufficiency", return_value={
                    "status": "WAITING_REQUIRED_USER_MATERIALS",
                    "missing_materials": ["missing_result_dataset", "references"],
                    "blocking_missing_materials": [],
                }),
                patch("main.write_section_local", side_effect=fake_write_section),
                patch("main.review_section", return_value={"status": "PASS"}),
                patch("main.search_paper_pool", return_value=[]),
                patch("main._auto_attach_open_image_assets", return_value=([], ["image search empty"])),
                patch("main._auto_generate_method_figures", return_value=([], [])),
                patch("main.render_docx", side_effect=lambda ast: render_docx(ast, str(output))),
            ):
                result = await _write_paper_pipeline({
                    "title": "Output-first paper",
                    "sections": [
                        {"id": "sec1", "title": "Introduction", "key_points": ["problem"]},
                        {"id": "sec2", "title": "Conclusion", "key_points": ["boundary"]},
                    ],
                }, job_id="write-test")

        self.assertEqual(result["status"], "completed")
        self.assertEqual(result["generation_mode"], "degraded")
        self.assertIn("missing_result_dataset", "\n".join(result["evidence_warnings"]))
        self.assertIn("image search empty", result["asset_warnings"])

    async def test_background_delivery_preserves_configured_review_rewrites(self):
        write_calls = []

        def fake_write_section(section, papers, **_kwargs):
            write_calls.append(section["id"])
            return {
                "section_id": section["id"],
                "title": section["title"],
                "content": "## Draft\n\nContent remains deliverable with a review warning.",
                "references_used": [],
            }

        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "background-fast.docx"
            with (
                patch("main._material_sufficiency", return_value={
                    "status": "DEGRADED_WRITING_ALLOWED",
                    "missing_materials": ["missing_result_dataset"],
                    "blocking_missing_materials": [],
                }),
                patch("main.write_section_local", side_effect=fake_write_section),
                patch("main.review_section", return_value={"status": "FAIL", "reason": "weak support"}),
                patch("main.search_paper_pool", return_value=[]),
                patch("main._auto_attach_open_image_assets", return_value=([], [])),
                patch("main._auto_generate_method_figures", return_value=([], [])),
                patch("main._auto_generate_result_figures", return_value=([], [], [], "")),
                patch("main.render_docx", side_effect=lambda ast: render_docx(ast, str(output))),
            ):
                result = await _write_paper_pipeline({
                    "title": "Background delivery",
                    "sections": [{"id": "sec1", "title": "Results", "key_points": ["claim"]}],
                }, job_id="write-fast")

        self.assertEqual(result["status"], "completed")
        self.assertEqual(write_calls, ["sec1", "sec1", "sec1"])

    async def test_chinese_task_freezes_language_into_outline_and_section_generation(self):
        move_prompts = []
        captured_asts = []

        def fake_write_section(section, papers, **kwargs):
            move_prompts.append(kwargs.get("move_sequence", ""))
            return {
                "section_id": section["id"],
                "title": section["title"],
                "content": "## 结果说明\n\n正文内容。",
                "references_used": [],
            }

        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "language.docx"
            with (
                patch("main._material_sufficiency", return_value={
                    "status": "DEGRADED_WRITING_ALLOWED",
                    "missing_materials": [],
                    "blocking_missing_materials": [],
                }),
                patch("main.write_section_local", side_effect=fake_write_section),
                patch("main.review_section", return_value={"status": "PASS"}),
                patch("main.search_paper_pool", return_value=[]),
                patch("main._auto_attach_open_image_assets", return_value=([], [])),
                patch("main._auto_generate_method_figures", return_value=([], [])),
                patch("main._auto_generate_result_figures", return_value=([], [], [], "")),
                patch(
                    "main.render_docx",
                    side_effect=lambda ast: captured_asts.append(ast) or render_docx(ast, str(output)),
                ),
            ):
                result = await _write_paper_pipeline({
                    "title": "EEG Study",
                    "language": "中文",
                    "sections": [{"id": "methods", "title": "Methods", "key_points": ["pipeline"]}],
                }, job_id="write-zh")

        self.assertEqual(result["status"], "completed")
        self.assertEqual(captured_asts[0]["language"], "中文")
        self.assertEqual(captured_asts[0]["document_profile"], CQU_PROFESSIONAL_MASTER)
        self.assertIn("front_matter", captured_asts[0])
        self.assertEqual(result["document_profile"], CQU_PROFESSIONAL_MASTER)
        self.assertEqual(captured_asts[0]["sections"][0]["title"], "方法")
        self.assertIn("整篇论文必须使用中文撰写", move_prompts[0])

    async def test_default_section_tool_redirects_to_whole_paper_tool(self):
        result = await write_section_step(section_id="sec1", section_title="Introduction")

        self.assertEqual(result["status"], "redirect_to_write_paper")
        self.assertEqual(result["redirect_tool"], "write_paper")

    async def test_source_request_requires_outline_display_before_job_start(self):
        with tempfile.TemporaryDirectory() as tmp:
            with (
                patch("main._pending_outlines", PendingOutlineStore(Path(tmp) / "pending.json")),
                patch("main.generate_outline", return_value={
                    "status": "waiting_materials",
                    "outline": {"title": "Frozen uploaded subject", "sections": []},
                    "material_sufficiency": {
                        "missing_materials": ["missing_result_dataset"],
                    },
                }),
                patch("main.asyncio.create_task") as create_task,
            ):
                result = await write_paper(
                    document_path="s3://nexent/attachments/project.docx",
                    topic="Frozen uploaded subject",
                    requirements="生成完整论文",
                    run_mode="submit",
                )

        self.assertEqual(result["status"], "outline_confirmation_required")
        self.assertEqual(result["outline"]["title"], "Frozen uploaded subject")
        create_task.assert_not_called()

    def test_generate_outline_reuses_pending_ticket_without_regeneration(self):
        with tempfile.TemporaryDirectory() as tmp:
            store = PendingOutlineStore(Path(tmp) / "pending.json")
            source_request = {
                "document_path": "",
                "topic": "Frozen topic",
                "requirements": "完整论文",
                "language": "中文",
            }
            ticket = store.create(source_request, {
                "status": "waiting_confirmation",
                "outline": {"title": "Frozen topic", "sections": []},
            })
            with (
                patch("main._pending_outlines", store),
                patch("main.get_llm_client") as llm_client,
            ):
                result = generate_outline(
                    topic="Frozen topic",
                    requirements="完整论文",
                    language="中文",
                )

        self.assertEqual(result["status"], "outline_confirmation_required")
        self.assertEqual(result["outline_id"], ticket["outline_id"])
        self.assertTrue(result["confirmation_without_id_supported"])
        self.assertTrue(result["reused_pending_outline"])
        llm_client.assert_not_called()

    def test_new_topic_outline_drops_materials_bound_to_previous_request(self):
        captured_prompts = []

        class FakeClient:
            def call(self, prompt, **_kwargs):
                captured_prompts.append(prompt)
                return '{"title": "空气质量预测论文", "sections": []}'

        with tempfile.TemporaryDirectory() as tmp:
            store = PendingOutlineStore(Path(tmp) / "pending.json")
            old_material = {
                "material_id": "mat-old",
                "parse_status": "parsed",
                "material_role": "source_text",
                "extracted_text": "脑电机械臂旧论文材料",
            }
            with (
                patch("main._pending_outlines", store),
                patch("main._materials.all", return_value=[old_material]),
                patch("main.get_llm_client", return_value=FakeClient()),
                patch("main._build_reference_pool", return_value=[]),
                patch("main._prepare_planned_image_candidates", return_value={"status": "not_required"}),
                patch("main._material_sufficiency", return_value={
                    "status": "WAITING_REQUIRED_USER_MATERIALS",
                    "missing_materials": ["missing_result_dataset"],
                    "blocking_missing_materials": [],
                }),
                patch.dict("main.DYNAMIC_KNOWLEDGE_BASE", {
                    "active_material_ids": ["mat-old"],
                    "active_material_request_fingerprint": "previous-request",
                    "source_material": "脑电机械臂旧论文材料",
                    "result_evidence": "旧实验结果",
                    "material_policy": {
                        "allow_degraded_writing": True,
                        "user_acknowledgement": "旧任务确认",
                    },
                }),
            ):
                result = generate_outline(topic="基于图神经网络的城市空气质量预测方法研究", language="中文")
                from main import DYNAMIC_KNOWLEDGE_BASE
                self.assertEqual(DYNAMIC_KNOWLEDGE_BASE["active_material_ids"], [])
                self.assertFalse(DYNAMIC_KNOWLEDGE_BASE["material_policy"]["allow_degraded_writing"])

        self.assertEqual(result["outline"]["title"], "空气质量预测论文")
        self.assertIn("基于图神经网络的城市空气质量预测方法研究", captured_prompts[0])
        self.assertNotIn("脑电机械臂旧论文材料", captured_prompts[0])

    def test_generate_outline_keeps_staged_upload_when_document_is_added(self):
        class FakeClient:
            def call(self, _prompt, **_kwargs):
                return '{"title": "联合材料论文", "sections": []}'

        staged_image = {
            "material_id": "mat-image",
            "parse_status": "parsed",
            "material_role": "image_asset",
            "local_path": "uploaded-result.png",
        }
        document_material = {
            "material_id": "mat-doc",
            "parse_status": "parsed",
            "material_role": "source_text",
            "extracted_text": "正文资料",
        }
        with tempfile.TemporaryDirectory() as tmp:
            with (
                patch("main._pending_outlines", PendingOutlineStore(Path(tmp) / "pending.json")),
                patch("main._materials.all", return_value=[staged_image, document_material]),
                patch("main._ingest_material_input", return_value={
                    "materials": [document_material],
                    "failed": [],
                    "unsupported": [],
                }),
                patch("main.get_llm_client", return_value=FakeClient()),
                patch("main._build_reference_pool", return_value=[]),
                patch("main._prepare_planned_image_candidates", return_value={"status": "not_required"}),
                patch("main._material_sufficiency", return_value={
                    "status": "WAITING_REQUIRED_USER_MATERIALS",
                    "missing_materials": [],
                    "blocking_missing_materials": [],
                }),
                patch.dict("main.DYNAMIC_KNOWLEDGE_BASE", {
                    "active_material_ids": ["mat-image"],
                    "active_material_request_fingerprint": "staged",
                }),
            ):
                generate_outline(document_path="uploaded.docx", topic="联合材料论文", language="中文")
                from main import DYNAMIC_KNOWLEDGE_BASE
                active_ids = DYNAMIC_KNOWLEDGE_BASE["active_material_ids"]

        self.assertEqual(active_ids, ["mat-image", "mat-doc"])

    async def test_confirmation_ticket_starts_existing_outline_without_regeneration(self):
        async def complete_background_job(job_id, _outline):
            from main import _write_jobs

            _write_jobs.finish(job_id, {
                "status": "completed",
                "job_id": job_id,
                "download_url": "http://paper/confirmed.docx",
                "asset_warnings": [],
                "quality_warnings": [],
                "evidence_warnings": [],
            })

        with tempfile.TemporaryDirectory() as tmp:
            store = PendingOutlineStore(Path(tmp) / "pending.json")
            ticket = store.create({
                "document_path": "",
                "topic": "Confirmed topic",
                "requirements": "完整论文",
                "language": "中文",
            }, {
                "status": "waiting_confirmation",
                "outline": {"title": "Confirmed topic", "sections": []},
                "asset_plan": [{"section_id": "methods", "purpose": "confirmed-device"}],
                "references": [{"title": "Confirmed reference"}],
            })
            with (
                patch("main._pending_outlines", store),
                patch.dict("main.DYNAMIC_KNOWLEDGE_BASE", {
                    "current_outline": {"title": "Wrong later topic", "sections": []},
                    "asset_plan": [{"section_id": "results", "purpose": "wrong"}],
                    "current_references": [{"title": "Wrong reference"}],
                    "current_reference_topic": "Wrong later topic",
                }),
                patch("main.generate_outline") as regenerate,
                patch("main._material_sufficiency", return_value={
                    "status": "DEGRADED_WRITING_ALLOWED",
                    "missing_materials": [],
                    "blocking_missing_materials": [],
                }),
                patch("main._run_background_write_job", side_effect=complete_background_job),
            ):
                result = await confirm_outline_and_start_writing(
                    outline_id=ticket["outline_id"],
                    allow_degraded_writing=True,
                    user_acknowledgement="用户确认降级写作",
                )
                from main import DYNAMIC_KNOWLEDGE_BASE
                self.assertEqual(DYNAMIC_KNOWLEDGE_BASE["current_outline"]["title"], "Confirmed topic")
                self.assertEqual(DYNAMIC_KNOWLEDGE_BASE["asset_plan"][0]["purpose"], "confirmed-device")
                self.assertEqual(DYNAMIC_KNOWLEDGE_BASE["current_reference_topic"], "Confirmed topic")

        self.assertEqual(result["status"], "completed")
        self.assertEqual(result["download_url"], "http://paper/confirmed.docx")
        regenerate.assert_not_called()

    async def test_confirmation_keeps_material_uploaded_after_outline_was_displayed(self):
        async def complete_background_job(job_id, _outline):
            from main import _write_jobs

            _write_jobs.finish(job_id, {"status": "completed", "job_id": job_id, "download_url": "http://paper/with-image.docx"})

        with tempfile.TemporaryDirectory() as tmp:
            store = PendingOutlineStore(Path(tmp) / "pending.json")
            ticket = store.create({"topic": "Confirmed topic", "language": "中文"}, {
                "status": "waiting_confirmation",
                "outline": {"title": "Confirmed topic", "sections": []},
                "active_material_ids": ["mat-doc"],
            })
            with (
                patch("main._pending_outlines", store),
                patch("main._materials.all", return_value=[
                    {"material_id": "mat-doc", "parse_status": "parsed", "material_role": "source_text"},
                    {"material_id": "mat-image", "parse_status": "parsed", "material_role": "image_asset"},
                ]),
                patch.dict("main.DYNAMIC_KNOWLEDGE_BASE", {
                    "active_material_ids": ["mat-image"],
                    "active_material_request_fingerprint": "staged",
                }),
                patch("main._material_sufficiency", return_value={
                    "status": "DEGRADED_WRITING_ALLOWED",
                    "missing_materials": [],
                    "blocking_missing_materials": [],
                }),
                patch("main._run_background_write_job", side_effect=complete_background_job),
            ):
                await confirm_outline_and_start_writing(outline_id=ticket["outline_id"])
                from main import DYNAMIC_KNOWLEDGE_BASE
                active_ids = DYNAMIC_KNOWLEDGE_BASE["active_material_ids"]

        self.assertEqual(active_ids, ["mat-doc", "mat-image"])

    async def test_confirmation_without_visible_id_uses_waiting_outline(self):
        async def complete_background_job(job_id, _outline):
            from main import _write_jobs

            _write_jobs.finish(job_id, {
                "status": "completed",
                "job_id": job_id,
                "download_url": "http://paper/latest.docx",
                "asset_warnings": [],
                "quality_warnings": [],
                "evidence_warnings": [],
            })

        with tempfile.TemporaryDirectory() as tmp:
            store = PendingOutlineStore(Path(tmp) / "pending.json")
            store.create({
                "document_path": "",
                "topic": "Displayed topic",
                "requirements": "完整论文",
                "language": "中文",
            }, {
                "status": "waiting_materials",
                "outline": {"title": "Displayed topic", "sections": []},
            })
            with (
                patch("main._pending_outlines", store),
                patch("main._material_sufficiency", return_value={
                    "status": "DEGRADED_WRITING_ALLOWED",
                    "missing_materials": ["missing_result_dataset"],
                    "blocking_missing_materials": [],
                }),
                patch("main._run_background_write_job", side_effect=complete_background_job),
            ):
                result = await confirm_outline_and_start_writing(
                    allow_degraded_writing=True,
                    user_acknowledgement="用户确认降级写作",
                )

        self.assertEqual(result["download_url"], "http://paper/latest.docx")

    async def test_source_confirmation_flag_uses_pending_outline_without_regeneration(self):
        async def complete_background_job(job_id, _outline):
            from main import _write_jobs

            _write_jobs.finish(job_id, {
                "status": "completed",
                "job_id": job_id,
                "download_url": "http://paper/source-confirmed.docx",
                "asset_warnings": [],
                "quality_warnings": [],
                "evidence_warnings": [],
            })

        with tempfile.TemporaryDirectory() as tmp:
            store = PendingOutlineStore(Path(tmp) / "pending.json")
            store.create({
                "document_path": "",
                "topic": "Same source",
                "requirements": "完整论文",
                "language": "中文",
            }, {
                "status": "waiting_confirmation",
                "outline": {"title": "Same source", "sections": []},
            })
            with (
                patch("main._pending_outlines", store),
                patch("main.generate_outline") as regenerate,
                patch("main._material_sufficiency", return_value={
                    "status": "DEGRADED_WRITING_ALLOWED",
                    "missing_materials": [],
                    "blocking_missing_materials": [],
                }),
                patch("main._run_background_write_job", side_effect=complete_background_job),
            ):
                result = await write_paper(
                    topic="Same source",
                    requirements="完整论文",
                    language="中文",
                    outline_confirmed=True,
                    allow_degraded_writing=True,
                    user_acknowledgement="用户确认降级写作",
                )

        self.assertEqual(result["download_url"], "http://paper/source-confirmed.docx")
        regenerate.assert_not_called()

    async def test_confirmed_display_text_recovers_outline_by_title_without_regeneration(self):
        async def complete_background_job(job_id, _outline):
            from main import _write_jobs

            _write_jobs.finish(job_id, {
                "status": "completed",
                "job_id": job_id,
                "download_url": "http://paper/recovered-confirmation.docx",
                "asset_warnings": [],
                "quality_warnings": [],
                "evidence_warnings": [],
            })

        with tempfile.TemporaryDirectory() as tmp:
            store = PendingOutlineStore(Path(tmp) / "pending.json")
            store.create({
                "document_path": "",
                "topic": "Displayed title",
                "requirements": "最初要求",
                "language": "中文",
            }, {
                "status": "waiting_confirmation",
                "outline": {"title": "Displayed title", "sections": []},
            })
            with (
                patch("main._pending_outlines", store),
                patch("main.generate_outline") as regenerate,
                patch("main._material_sufficiency", return_value={
                    "status": "DEGRADED_WRITING_ALLOWED",
                    "missing_materials": [],
                    "blocking_missing_materials": [],
                }),
                patch("main._run_background_write_job", side_effect=complete_background_job),
            ):
                result = await write_paper(
                    confirmed_outline="# Displayed title\n\n## 引言",
                    topic="Displayed title",
                    requirements="模型在确认轮次改写后的要求",
                    language="中文",
                    outline_confirmed=True,
                    allow_degraded_writing=True,
                    user_acknowledgement="用户已确认",
                )

        self.assertEqual(result["download_url"], "http://paper/recovered-confirmation.docx")
        regenerate.assert_not_called()

    async def test_background_pipeline_uses_frozen_task_snapshot(self):
        captured_asts = []

        def fake_write_section(section, papers, **_kwargs):
            self.assertEqual([paper["title"] for paper in papers], ["Frozen reference"])
            return {
                "section_id": section["id"],
                "title": section["title"],
                "content": "## 固定主题\n\n正文内容。",
                "references_used": [],
            }

        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "frozen.docx"
            with (
                patch.dict("main.DYNAMIC_KNOWLEDGE_BASE", {
                    "asset_plan": [{"purpose": "wrong-plan"}],
                    "current_references": [{"title": "Wrong reference"}],
                    "source_material": "wrong source",
                    "result_evidence": "wrong result",
                }),
                patch("main.assess_material_sufficiency", return_value={
                    "status": "DEGRADED_WRITING_ALLOWED",
                    "missing_materials": [],
                    "blocking_missing_materials": [],
                }),
                patch("main.write_section_local", side_effect=fake_write_section),
                patch("main.review_section", return_value={"status": "PASS"}),
                patch("main.search_paper_pool") as search_pool,
                patch("main._auto_attach_open_image_assets", return_value=([], [])),
                patch("main._auto_generate_method_figures", return_value=([], [])),
                patch("main._auto_generate_result_figures", return_value=([], [], [], "frozen result")),
                patch(
                    "main.render_docx",
                    side_effect=lambda ast: captured_asts.append(ast) or render_docx(ast, str(output)),
                ),
            ):
                result = await _write_paper_pipeline(
                    {"title": "Frozen topic", "sections": [{"id": "methods", "title": "方法"}]},
                    job_id="write-frozen",
                    task_snapshot={
                        "references": [{"title": "Frozen reference"}],
                        "asset_plan": [{"purpose": "frozen-plan"}],
                        "materials": [{"material_id": "mat-frozen"}],
                        "source_material": "frozen source",
                        "result_evidence": "frozen result",
                        "material_policy": {"allow_degraded_writing": True, "user_acknowledgement": "ok"},
                    },
                )

        self.assertEqual(result["status"], "completed")
        self.assertEqual(captured_asts[0]["asset_plan"], [{"purpose": "frozen-plan"}])
        self.assertEqual(captured_asts[0]["materials"], [{"material_id": "mat-frozen"}])
        search_pool.assert_not_called()

    async def test_confirmed_outline_with_evidence_gap_requires_degradation_choice(self):
        with (
            patch("main._material_sufficiency", return_value={
                "status": "WAITING_REQUIRED_USER_MATERIALS",
                "missing_materials": ["missing_result_dataset"],
                "blocking_missing_materials": [],
            }),
            patch("main.asyncio.create_task") as create_task,
        ):
            result = await write_paper(
                confirmed_outline={"title": "Paper", "sections": []},
                outline_confirmed=True,
                run_mode="submit",
            )

        self.assertEqual(result["status"], "material_confirmation_required")
        create_task.assert_not_called()

    async def test_default_background_mode_waits_for_final_download_in_one_tool_call(self):
        async def complete_background_job(job_id, _outline):
            from main import _write_jobs

            _write_jobs.finish(job_id, {
                "status": "completed",
                "job_id": job_id,
                "download_url": "http://paper/final.docx",
                "asset_warnings": [],
                "quality_warnings": [],
                "evidence_warnings": [],
            })

        with (
            patch("main._material_sufficiency", return_value={
                "status": "DEGRADED_WRITING_ALLOWED",
                "missing_materials": [],
                "blocking_missing_materials": [],
            }),
            patch("main._run_background_write_job", side_effect=complete_background_job),
        ):
            result = await write_paper(
                confirmed_outline={"title": "Single active MCP request", "sections": []},
                outline_confirmed=True,
                allow_degraded_writing=True,
                user_acknowledgement="用户确认接受降级稿风险",
            )

        self.assertEqual(result["status"], "completed")
        self.assertEqual(result["download_url"], "http://paper/final.docx")

    async def test_whole_paper_job_rejects_stale_outline_fallback(self):
        result = await write_paper()

        self.assertEqual(result["status"], "error")
        self.assertIn("stale server outline fallback is disabled", result["message"])


class OutputFirstStatusTests(unittest.TestCase):
    def test_job_status_defaults_to_single_long_wait(self):
        from inspect import signature

        self.assertEqual(signature(get_write_paper_job_status).parameters["wait_seconds"].default, 600)

    def test_job_status_long_poll_returns_completed_result(self):
        running = {
            "status": "running",
            "title": "Draft",
            "progress": {"current": 1, "total": 6},
            "events": [],
            "result_status": "",
            "result": None,
        }
        completed = {
            **running,
            "status": "completed",
            "result_status": "completed",
            "result": {
                "status": "completed",
                "download_url": "http://paper",
                "asset_warnings": [f"asset-{idx}" for idx in range(8)],
            },
        }

        with (
            patch("main._write_jobs.get", side_effect=[running, completed]) as get_job,
            patch("main.time.monotonic", side_effect=[10.0, 10.0]),
            patch("main.time.sleep") as sleep,
        ):
            result = get_write_paper_job_status("write-1", wait_seconds=60)

        self.assertEqual(result["status"], "completed")
        self.assertEqual(result["download_url"], "http://paper")
        self.assertEqual(result["warning_counts"]["asset"], 8)
        self.assertEqual(result["asset_warnings"], ["asset-0", "asset-1", "asset-2"])
        self.assertNotIn("result", result)
        self.assertEqual(get_job.call_count, 2)
        sleep.assert_called_once_with(1)


if __name__ == "__main__":
    unittest.main()
