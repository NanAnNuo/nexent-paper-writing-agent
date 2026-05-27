import tempfile
import unittest
import zipfile
import base64
from pathlib import Path
from unittest.mock import patch

from docx import Document

from execution.render_validator import validate_renderable_ast
from orchestrator.asset_manager import ImageAssetStore
from orchestrator.material_registry import MaterialRegistry
from orchestrator.material_sufficiency import assess_material_sufficiency, build_asset_plan


class MaterialRegistryTests(unittest.TestCase):
    def test_ingests_docx_txt_and_result_csv(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            txt = root / "notes.txt"
            txt.write_text("Method notes for a classifier experiment.", encoding="utf-8")
            docx = root / "brief.docx"
            document = Document()
            document.add_paragraph("Research problem and study design.")
            document.save(docx)
            csv = root / "results.csv"
            csv.write_text("accuracy,latency\n0.91,120\n0.93,110\n", encoding="utf-8")

            registry = MaterialRegistry(root / "materials.json")
            report = registry.ingest_paths([txt, docx, csv])

            self.assertEqual(report["failed"], [])
            self.assertEqual(len(registry.result_materials()), 1)
            self.assertIn("accuracy", registry.result_summaries()[0])
            self.assertIn("Research problem", registry.readable_text())

    def test_docx_embedded_raster_images_become_image_materials(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            png = root / "device.png"
            png.write_bytes(base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII="
            ))
            docx = root / "brief.docx"
            document = Document()
            document.add_paragraph("Methods material.")
            document.add_picture(str(png))
            document.save(docx)

            registry = MaterialRegistry(root / "materials.json")
            report = registry.ingest_paths([docx])
            images = [item for item in report["materials"] if item.get("material_role") == "image_asset"]

            self.assertEqual(len(images), 1)
            self.assertEqual(images[0]["metadata"]["source_kind"], "docx_embedded_image")
            self.assertTrue(Path(images[0]["local_path"]).exists())

    def test_pdf_without_extractable_text_is_recorded_as_failed(self):
        with tempfile.TemporaryDirectory() as tmp:
            pdf = Path(tmp) / "scan.pdf"
            pdf.write_bytes(b"%PDF-1.4 fake")
            registry = MaterialRegistry(Path(tmp) / "materials.json")
            with patch(
                "orchestrator.material_registry._read_pdf",
                side_effect=ValueError("pdf_has_no_extractable_text"),
            ):
                report = registry.ingest_paths([pdf])

            self.assertEqual(report["materials"], [])
            self.assertEqual(len(report["failed"]), 1)
            self.assertEqual(registry.all()[0]["parse_status"], "failed")

    def test_zip_ingests_supported_files_and_lists_unsupported_members(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            bundle = root / "materials.zip"
            with zipfile.ZipFile(bundle, "w") as archive:
                archive.writestr("source.txt", "Method and project background.")
                archive.writestr("ignored.bin", b"\x00\x01")

            registry = MaterialRegistry(root / "materials.json")
            report = registry.ingest_paths([bundle])

            self.assertEqual(len(report["materials"]), 1)
            self.assertEqual(len(report["unsupported"]), 1)
            self.assertIn("project background", registry.readable_text())


class MaterialGateTests(unittest.TestCase):
    def test_literary_analysis_does_not_plan_engineering_figures_from_section_names(self):
        plan = build_asset_plan(
            [
                {"id": "intro", "title": "引言"},
                {"id": "method", "title": "方法：文本分析与历史语境"},
                {"id": "results", "title": "结果：独白意蕴分析"},
            ],
            topic="在行动的边缘颤栗：《哈姆雷特》生存还是毁灭之谜",
        )

        self.assertEqual(plan, [])

    def test_technical_topic_still_plans_workflow_and_result_assets(self):
        plan = build_asset_plan(
            [
                {"id": "methods", "title": "方法"},
                {"id": "results", "title": "结果"},
            ],
            topic="基于脑电运动想象的意念控制机械臂系统",
        )

        self.assertTrue(any(item["purpose"] == "system_or_device_context" for item in plan))
        self.assertTrue(any(item["purpose"] == "result_evidence" for item in plan))

    def test_registry_result_dataset_satisfies_result_material_gate(self):
        report = assess_material_sufficiency(
            topic="EEG experiment result evaluation",
            material_text="Method: EEG signals are collected for classifier evaluation.",
            references=[{"title": "Paper"}],
            approved_assets=[{"asset_id": "approved", "approved": True}],
            materials=[{
                "material_id": "mat-result",
                "material_role": "result_dataset",
                "parse_status": "parsed",
                "data_summary": "rows=10, accuracy mean=0.9",
            }],
        )

        self.assertEqual(report["status"], "sufficient")
        self.assertNotIn("missing_result_dataset", report["missing_materials"])

    def test_user_acknowledged_degraded_writing_can_continue_without_result_dataset(self):
        report = assess_material_sufficiency(
            topic="EEG experiment result evaluation",
            material_text="Method: a classifier evaluation will be discussed from available sources.",
            references=[{"title": "Paper"}],
            allow_degraded_writing=True,
            degraded_reason="User cannot provide the result dataset and accepts quality risk.",
        )

        self.assertEqual(report["status"], "DEGRADED_WRITING_ALLOWED")
        self.assertIn("missing_result_dataset", report["missing_materials"])
        self.assertTrue(report["quality_risk_acknowledged"])
        self.assertEqual(report["next_action"], "write_sections_with_quality_risk_disclosure")

    def test_waiting_gate_exposes_degraded_writing_retry_parameters(self):
        report = assess_material_sufficiency(
            topic="EEG experiment result evaluation",
            material_text="Method: EEG data will be evaluated with a classifier.",
            references=[{"title": "Paper"}],
        )

        self.assertEqual(report["status"], "WAITING_REQUIRED_USER_MATERIALS")
        self.assertTrue(report["degraded_writing_available"])
        self.assertTrue(report["degraded_writing_requirements"]["allow_degraded_writing"])

    def test_degraded_writing_does_not_guess_a_missing_research_problem(self):
        report = assess_material_sufficiency(
            references=[{"title": "Paper"}],
            allow_degraded_writing=True,
            degraded_reason="User accepts quality risk.",
        )

        self.assertEqual(report["status"], "WAITING_REQUIRED_USER_MATERIALS")
        self.assertIn("research_problem", report["blocking_missing_materials"])

    def test_uploaded_image_unbound_warns_without_blocking_render(self):
        ast = {
            "title": "Draft",
            "sections": [{"section_id": "intro", "title": "Intro", "content": "Text."}],
            "references": [],
            "materials": [{
                "material_id": "mat-image",
                "material_role": "image_asset",
                "parse_status": "parsed",
            }],
            "entity_registry": {"images": []},
        }

        result = validate_renderable_ast(ast)

        self.assertTrue(result["ok"])
        self.assertTrue(any("unbound uploaded image asset" in warning for warning in result["warnings"]))

    def test_uploaded_image_binding_creates_renderable_asset_metadata(self):
        with tempfile.TemporaryDirectory() as tmp:
            image = Path(tmp) / "device.png"
            image.write_bytes(b"not-a-real-image-for-metadata-test")
            store = ImageAssetStore(Path(tmp) / "assets.json")
            asset = store.bind_uploaded_material(
                {
                    "material_id": "mat-image",
                    "local_path": str(image),
                    "source_path": str(image),
                    "metadata": {"filename": "device.png"},
                },
                section_id="methods",
                purpose="device_photo",
                caption="Device used in the experiment.",
            )

            self.assertTrue(asset["approved"])
            self.assertEqual(asset["material_id"], "mat-image")
            self.assertEqual(asset["license"], "user_provided")

    def test_uploaded_image_binding_copies_asset_into_job_directory(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            image = root / "accuracy.png"
            image.write_bytes(b"image")
            store = ImageAssetStore(root / "assets.json")
            asset = store.bind_uploaded_material(
                {
                    "material_id": "mat-result",
                    "local_path": str(image),
                    "source_path": str(image),
                    "metadata": {"filename": "accuracy.png"},
                },
                section_id="results",
                purpose="user_provided_result_image",
                caption="Uploaded accuracy result.",
                job_id="write-one",
                output_dir=root / "write-one" / "images",
            )

            self.assertEqual(asset["job_id"], "write-one")
            self.assertIn("write-one", asset["local_path"])
            self.assertTrue(Path(asset["local_path"]).exists())


if __name__ == "__main__":
    unittest.main()
