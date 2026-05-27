"""Structured ingestion and storage for user-provided paper materials."""

from __future__ import annotations

import json
import csv
import shutil
import uuid
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document


TEXT_TYPES = {".docx", ".txt", ".pdf"}
DATA_TYPES = {".csv", ".tsv", ".xlsx"}
IMAGE_TYPES = {".png", ".jpg", ".jpeg", ".webp"}
SUPPORTED_TYPES = TEXT_TYPES | DATA_TYPES | IMAGE_TYPES | {".zip"}
ZIP_MAX_FILES = 100
ZIP_MAX_BYTES = 200 * 1024 * 1024


class MaterialRegistry:
    """Persist parsed source materials and data evidence for one paper project."""

    def __init__(self, path: str | Path = "data/checkpoints/materials.json"):
        self.path = Path(path)
        self.path.parent.mkdir(parents=True, exist_ok=True)
        self._materials = self._load()

    def ingest_paths(self, paths: Iterable[str | Path]) -> dict:
        materials = []
        failed = []
        unsupported = []
        for raw_path in paths:
            report = self.ingest_path(raw_path)
            materials.extend(report["materials"])
            failed.extend(report["failed"])
            unsupported.extend(report["unsupported"])
        self._save()
        return {
            "materials": materials,
            "failed": failed,
            "unsupported": unsupported,
        }

    def ingest_path(self, raw_path: str | Path) -> dict:
        path = Path(raw_path)
        if not path.exists():
            return {
                "materials": [],
                "failed": [{"source_path": str(path), "error": "file_not_found"}],
                "unsupported": [],
            }
        if path.suffix.lower() == ".zip":
            return self._ingest_zip(path)

        material = self._parse_material(path)
        self._upsert(material)
        unsupported = [material] if material["material_role"] == "unsupported" else []
        failed = [material] if material["parse_status"] == "failed" else []
        materials = [] if unsupported or failed else [material]
        if path.suffix.lower() == ".docx" and material["parse_status"] == "parsed":
            embedded = self._extract_docx_image_materials(path, parent_material_id=material["material_id"])
            for image_material in embedded:
                self._upsert(image_material)
            materials.extend(embedded)
        return {"materials": materials, "failed": failed, "unsupported": unsupported}

    def all(self) -> list[dict]:
        return [dict(material) for material in self._materials]

    def summaries(self) -> list[dict]:
        fields = (
            "material_id",
            "source_path",
            "local_path",
            "file_type",
            "material_role",
            "parse_status",
            "metadata",
            "errors",
        )
        return [{field: material.get(field) for field in fields} for material in self._materials]

    def readable_text(self) -> str:
        return "\n\n".join(
            material["extracted_text"]
            for material in self._materials
            if material.get("parse_status") == "parsed"
            and material.get("material_role") in {"source_text", "reference_material"}
            and material.get("extracted_text")
        )

    def result_summaries(self) -> list[str]:
        return [
            material["data_summary"]
            for material in self._materials
            if material.get("parse_status") == "parsed"
            and material.get("material_role") == "result_dataset"
            and material.get("data_summary")
        ]

    def result_materials(self) -> list[dict]:
        return [
            dict(material)
            for material in self._materials
            if material.get("parse_status") == "parsed"
            and material.get("material_role") == "result_dataset"
        ]

    def image_materials(self) -> list[dict]:
        return [
            dict(material)
            for material in self._materials
            if material.get("parse_status") == "parsed"
            and material.get("material_role") == "image_asset"
        ]

    def has_failed_source_material(self) -> bool:
        return any(
            material.get("parse_status") == "failed"
            and material.get("file_type") in TEXT_TYPES
            for material in self._materials
        )

    def get(self, material_id: str) -> dict | None:
        for material in self._materials:
            if material.get("material_id") == material_id:
                return dict(material)
        return None

    def _ingest_zip(self, path: Path) -> dict:
        materials = []
        failed = []
        unsupported = []
        target_root = self.path.parent / "materials" / f"zip-{uuid.uuid4().hex[:10]}"
        target_root.mkdir(parents=True, exist_ok=True)

        try:
            with zipfile.ZipFile(path) as archive:
                infos = [info for info in archive.infolist() if not info.is_dir()]
                if len(infos) > ZIP_MAX_FILES:
                    raise ValueError(f"zip_file_limit_exceeded:{ZIP_MAX_FILES}")
                if sum(info.file_size for info in infos) > ZIP_MAX_BYTES:
                    raise ValueError(f"zip_size_limit_exceeded:{ZIP_MAX_BYTES}")

                for info in infos:
                    relative = Path(info.filename)
                    if relative.is_absolute() or ".." in relative.parts:
                        failed.append({
                            "source_path": f"{path}!{info.filename}",
                            "error": "unsafe_zip_member",
                        })
                        continue
                    extracted = target_root / relative
                    extracted.parent.mkdir(parents=True, exist_ok=True)
                    with archive.open(info) as source, extracted.open("wb") as target:
                        shutil.copyfileobj(source, target)
                    report = self.ingest_path(extracted)
                    materials.extend(report["materials"])
                    failed.extend(report["failed"])
                    unsupported.extend(report["unsupported"])
        except Exception as exc:
            failed.append({"source_path": str(path), "error": str(exc)})
        return {"materials": materials, "failed": failed, "unsupported": unsupported}

    def _parse_material(self, path: Path) -> dict:
        suffix = path.suffix.lower()
        base = {
            "material_id": f"mat-{uuid.uuid4().hex[:10]}",
            "source_path": str(path),
            "local_path": str(path.resolve()),
            "file_type": suffix,
            "material_role": "unsupported",
            "parse_status": "unsupported",
            "extracted_text": "",
            "data_summary": "",
            "metadata": {"filename": path.name, "size_bytes": path.stat().st_size},
            "errors": [],
        }
        if suffix not in SUPPORTED_TYPES:
            base["errors"].append("unsupported_file_type")
            return base

        try:
            if suffix == ".txt":
                base["material_role"] = "source_text"
                base["extracted_text"] = path.read_text(encoding="utf-8")
                base["parse_status"] = "parsed"
            elif suffix == ".docx":
                base["material_role"] = "source_text"
                base["extracted_text"] = _read_docx(path)
                base["parse_status"] = "parsed"
            elif suffix == ".pdf":
                base["material_role"] = "reference_material"
                base["extracted_text"] = _read_pdf(path)
                base["parse_status"] = "parsed"
            elif suffix in {".csv", ".tsv", ".xlsx"}:
                base["material_role"] = "result_dataset"
                base["data_summary"], metadata = _summarize_dataset(path)
                base["metadata"].update(metadata)
                base["parse_status"] = "parsed"
            elif suffix in IMAGE_TYPES:
                base["material_role"] = "image_asset"
                image_dir = self.path.parent / "materials" / "images"
                image_dir.mkdir(parents=True, exist_ok=True)
                target = image_dir / f"{base['material_id']}{suffix}"
                if path.resolve() != target.resolve():
                    shutil.copy2(path, target)
                base["local_path"] = str(target.resolve())
                base["parse_status"] = "parsed"
        except Exception as exc:
            base["parse_status"] = "failed"
            base["errors"].append(str(exc))
        return base

    def _extract_docx_image_materials(self, path: Path, *, parent_material_id: str) -> list[dict]:
        """Extract embedded DOCX raster images into the material asset store."""
        image_dir = self.path.parent / "materials" / "docx_images" / parent_material_id
        image_dir.mkdir(parents=True, exist_ok=True)
        materials = []
        try:
            with zipfile.ZipFile(path) as archive:
                members = [
                    member for member in archive.namelist()
                    if member.startswith("word/media/")
                    and Path(member).suffix.lower() in IMAGE_TYPES
                ]
                for index, member in enumerate(members, 1):
                    suffix = Path(member).suffix.lower()
                    material_id = f"mat-{uuid.uuid4().hex[:10]}"
                    target = image_dir / f"{material_id}{suffix}"
                    with archive.open(member) as source, target.open("wb") as output:
                        shutil.copyfileobj(source, output)
                    materials.append({
                        "material_id": material_id,
                        "source_path": f"{path}!{member}",
                        "local_path": str(target.resolve()),
                        "file_type": suffix,
                        "material_role": "image_asset",
                        "parse_status": "parsed",
                        "extracted_text": "",
                        "data_summary": "",
                        "metadata": {
                            "filename": Path(member).name,
                            "size_bytes": target.stat().st_size,
                            "source_kind": "docx_embedded_image",
                            "parent_material_id": parent_material_id,
                            "embedded_index": index,
                        },
                        "errors": [],
                    })
        except Exception as exc:
            # Embedded assets are optional; the parent DOCX material stays usable.
            return [{
                "material_id": f"mat-{uuid.uuid4().hex[:10]}",
                "source_path": str(path),
                "local_path": "",
                "file_type": ".docx",
                "material_role": "image_asset",
                "parse_status": "failed",
                "extracted_text": "",
                "data_summary": "",
                "metadata": {
                    "source_kind": "docx_embedded_image",
                    "parent_material_id": parent_material_id,
                },
                "errors": [f"docx_image_extract_failed:{exc}"],
            }]
        return materials

    def _upsert(self, material: dict):
        local_path = material.get("local_path")
        existing = {
            item.get("local_path"): item
            for item in self._materials
            if item.get("local_path")
        }
        if local_path in existing:
            material["material_id"] = existing[local_path]["material_id"]
            self._materials = [
                material if item.get("local_path") == local_path else item
                for item in self._materials
            ]
        else:
            self._materials.append(material)

    def _load(self) -> list[dict]:
        if not self.path.exists():
            return []
        return json.loads(self.path.read_text(encoding="utf-8"))

    def _save(self):
        self.path.write_text(json.dumps(self._materials, ensure_ascii=False, indent=2), encoding="utf-8")


def _read_docx(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pdf_parser_unavailable:pypdf") from exc

    reader = PdfReader(str(path))
    text = "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
    if len(text) < 40:
        raise ValueError("pdf_has_no_extractable_text")
    return text


def _summarize_dataset(path: Path) -> tuple[str, dict]:
    suffix = path.suffix.lower()
    if suffix in {".csv", ".tsv"}:
        delimiter = "," if suffix == ".csv" else "\t"
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            reader = csv.DictReader(handle, delimiter=delimiter)
            rows = list(reader)
        return _summarize_row_dicts("data", rows, list(reader.fieldnames or []))

    try:
        import pandas as pd
    except ImportError as exc:
        raise RuntimeError("xlsx_parser_unavailable:pandas") from exc

    sheet_summaries = []
    metadata = {"sheets": []}
    for sheet_name, frame in pd.read_excel(path, sheet_name=None).items():
        summary, sheet_metadata = _summarize_row_dicts(
            str(sheet_name),
            frame.to_dict(orient="records"),
            [str(column) for column in frame.columns],
        )
        sheet_summaries.append(summary)
        metadata["sheets"].extend(sheet_metadata["sheets"])
    return "\n\n".join(sheet_summaries), metadata


def _summarize_row_dicts(sheet_name: str, rows: list[dict], columns: list[str]) -> tuple[str, dict]:
    missing = 0
    numeric_values: dict[str, list[float]] = {column: [] for column in columns}
    for row in rows:
        for column in columns:
            value = row.get(column, "")
            if value in ("", None):
                missing += 1
                continue
            try:
                numeric_values[column].append(float(value))
            except (TypeError, ValueError):
                continue
    numeric_columns = [column for column, values in numeric_values.items() if values]
    metrics = []
    for column in numeric_columns[:8]:
        values = numeric_values[column]
        metrics.append(
            f"{column}: min={min(values):.4g}, mean={sum(values) / len(values):.4g}, max={max(values):.4g}"
        )
    summary = "\n".join([
        f"Sheet {sheet_name}: rows={len(rows)}, columns={len(columns)}, missing_cells={missing}",
        f"Columns: {', '.join(columns[:20])}",
        f"Numeric columns: {', '.join(numeric_columns[:20]) or 'none'}",
        f"Numeric overview: {'; '.join(metrics) or 'no numeric overview available'}",
    ])
    return summary, {
        "sheets": [{
            "name": sheet_name,
            "rows": len(rows),
            "columns": columns[:50],
            "numeric_columns": numeric_columns[:50],
            "missing_cells": missing,
        }]
    }
